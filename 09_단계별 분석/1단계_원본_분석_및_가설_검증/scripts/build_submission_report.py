#!/usr/bin/env python3
"""Create timestamped submission workbook, six charts, and a Word report."""
from __future__ import annotations
import argparse, json, shutil
from datetime import datetime
from pathlib import Path
from zoneinfo import ZoneInfo
import matplotlib.pyplot as plt
from matplotlib import font_manager
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def args():
    p=argparse.ArgumentParser(); p.add_argument('--analysis-dir',type=Path,required=True); p.add_argument('--output-dir',type=Path,required=True); return p.parse_args()
def load(p,n): return pd.read_csv(p/f'{n}.csv')
def setup_font():
    candidates=['Malgun Gothic','NanumGothic','Noto Sans CJK KR']
    available={f.name for f in font_manager.fontManager.ttflist}
    font=next((x for x in candidates if x in available),None)
    if not font: raise RuntimeError('Korean font not found')
    plt.rcParams.update({'font.family':font,'axes.unicode_minus':False,'font.size':11,'figure.dpi':140})
    return font
def savefig(path):
    plt.tight_layout(); plt.savefig(path,dpi=180,bbox_inches='tight',facecolor='white'); plt.close()
def label_bars(ax,bars,nums,dens=None):
    for i,(b,n) in enumerate(zip(bars,nums)):
        detail=f'{n}/{dens[i]}' if dens is not None else f'n={n}'
        ax.text(b.get_x()+b.get_width()/2,b.get_height()+2,f'{b.get_height():.1f}%\n({detail})',ha='center',va='bottom',fontsize=9)
def percent_axis(ax,title):
    ax.set_ylim(0,100); ax.set_ylabel('비율 (%)'); ax.set_title(title,fontweight='bold',pad=14); ax.grid(axis='y',alpha=.25); ax.spines[['top','right']].set_visible(False)
def add_table(doc,headers,rows,widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Light Shading Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers): t.rows[0].cells[i].text=str(h)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row): cells[i].text=str(v); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs: r.font.name='맑은 고딕'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); r.font.size=Pt(8.5)
    return t
def section(doc,title,text=None):
    doc.add_heading(title,level=1)
    if text: doc.add_paragraph(text)
def add_picture(doc,path,caption):
    doc.add_picture(str(path),width=Inches(6.35)); p=doc.paragraphs[-1]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    c=doc.add_paragraph(caption); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; c.style='Caption'
def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def main():
    a=args(); src=a.analysis_dir.resolve(); out=a.output_dir.resolve(); out.mkdir(parents=True,exist_ok=True)
    now=datetime.now(ZoneInfo('Asia/Seoul')); stamp=now.strftime('%y%m%d_%H%M')
    workbook=out/f'호텔검색_원본분석_분석결과_{stamp}_01.xlsx'
    report=out/f'호텔검색_원본분석_분석보고서_{stamp}_01.docx'
    topics=['핵심지표','A1_필터별_결과없음률','지역별_결과없음률','검색의도별_결과없음률','세션_결과_세그먼트','재검색유형별_성공_상세진입']
    charts=[out/f'호텔검색_원본분석_{t}_{stamp}_01.png' for t in topics]
    for p in [workbook,report,*charts]:
        if p.exists(): raise FileExistsError(f'Existing output protected: {p}')
    shutil.copy2(src/'original_296_analysis.xlsx',workbook)
    ov=load(src,'overview'); a1=load(src,'A1_filters'); reg=load(src,'A2_region'); intent=load(src,'A2_intent'); b1=load(src,'B1'); b2=load(src,'B2'); b3=load(src,'B3'); seg=load(src,'segments'); h3=load(src,'H3_transitions')
    setup_font(); blue='#2F6BFF'; orange='#F59E0B'; teal='#0F9D8A'; gray='#667085'
    # 1 core metrics
    names=['결과 없음률','0건 후\n후속검색률','비0건 후\n후속검색률','즉시 회복률','세션 최종\n회복률','상세진입률\n(전체 검색)']
    keys=['zero_result_rate','followup_rate_after_zero','followup_rate_after_nonzero','immediate_recovery_rate','session_final_recovery_rate','hotel_click_detail_entry_rate_all_searches']
    rr=ov.set_index('metric').loc[keys]; vals=rr.rate.mul(100).tolist(); nums=rr.numerator.astype(int).tolist(); dens=rr.denominator.astype(int).tolist()
    fig,ax=plt.subplots(figsize=(10.5,5.4)); bars=ax.bar(names,vals,color=[gray,blue,gray,teal,teal,orange]); label_bars(ax,bars,nums,dens); percent_axis(ax,'원본 296건 핵심 지표'); ax.tick_params(axis='x',labelsize=9); savefig(charts[0])
    # 2 A1 paired bars
    piv=a1.copy(); labels=['편의시설 3개 이상','최소평점 설정','가격 설정']; exp=piv.iloc[[0,2,4]]; ref=piv.iloc[[1,3,5]]; x=np.arange(3); w=.36
    fig,ax=plt.subplots(figsize=(9.6,5.5)); bA=ax.bar(x-w/2,exp.rate*100,w,label='설정/제한군',color=blue); bB=ax.bar(x+w/2,ref.rate*100,w,label='미설정/비교군',color=gray); label_bars(ax,bA,exp.numerator.astype(int).tolist(),exp.denominator.astype(int).tolist()); label_bars(ax,bB,ref.numerator.astype(int).tolist(),ref.denominator.astype(int).tolist()); ax.set_xticks(x,labels); ax.legend(); percent_axis(ax,'A1 — 필터 제한과 결과 없음'); savefig(charts[1])
    # 3/4 simple bars
    for df,path,title,color in [(reg,charts[2],'지역별 결과 없음률',blue),(intent,charts[3],'검색의도별 결과 없음률',teal)]:
        fig,ax=plt.subplots(figsize=(9.5,5.4)); bars=ax.bar(df.group,df.rate*100,color=color); label_bars(ax,bars,df.numerator.astype(int).tolist(),df.denominator.astype(int).tolist()); percent_axis(ax,title); ax.tick_params(axis='x',rotation=15)
        sparse=df[df.denominator<10]
        if len(sparse): ax.text(.99,.97,'[주의] 희소 셀: '+', '.join(f'{r.group} n={int(r.denominator)}' for _,r in sparse.iterrows())+'\n기술통계로만 해석',transform=ax.transAxes,ha='right',va='top',fontsize=9,color='#B54708',bbox=dict(facecolor='#FFF7E6',edgecolor='#F79009',boxstyle='round,pad=.35'))
        savefig(path)
    # 5 segments
    fig,ax=plt.subplots(figsize=(9.5,5.4)); bars=ax.bar(seg.group,seg.rate*100,color=[blue,gray,teal,orange]); label_bars(ax,bars,seg.n.astype(int).tolist()); percent_axis(ax,'4개 세션 결과 세그먼트 (n=43)'); savefig(charts[4])
    # 6 H3 grouped
    pr=h3.pivot(index='group',columns='metric',values=['rate','numerator','denominator']).reindex(['동일조건 반복','조건 완화','검색어 수정','지역 변경','조건 강화','혼합 변경']); x=np.arange(len(pr)); w=.36
    m1='next_search_positive_result_rate'; m2='next_search_hotel_click_rate'
    fig,ax=plt.subplots(figsize=(11,5.6)); b1x=ax.bar(x-w/2,pr[('rate',m1)]*100,w,label='다음 검색 결과 발생률',color=teal); b2x=ax.bar(x+w/2,pr[('rate',m2)]*100,w,label='다음 검색 hotel_click률',color=orange); label_bars(ax,b1x,pr[('numerator',m1)].astype(int).tolist(),pr[('denominator',m1)].astype(int).tolist()); label_bars(ax,b2x,pr[('numerator',m2)].astype(int).tolist(),pr[('denominator',m2)].astype(int).tolist()); ax.set_xticks(x,pr.index); ax.legend(); percent_axis(ax,'H3 — 재검색 유형별 다음 검색 결과'); ax.text(.99,.97,'[주의] 혼합 변경 n=2 · 탐색적 결과',transform=ax.transAxes,ha='right',va='top',fontsize=9,color='#B54708'); savefig(charts[5])
    # Word report
    doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.65); sec.bottom_margin=Inches(.65); sec.left_margin=Inches(.75); sec.right_margin=Inches(.75)
    styles=doc.styles
    for name in ['Normal','Title','Heading 1','Heading 2','Caption']:
        s=styles[name]; s.font.name='맑은 고딕'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕')
    styles['Normal'].font.size=Pt(10); styles['Normal'].paragraph_format.space_after=Pt(6); styles['Heading 1'].font.color.rgb=RGBColor(31,78,121)
    title=doc.add_paragraph(); title.style='Title'; title.alignment=WD_ALIGN_PARAGRAPH.CENTER; title.add_run('일본 호텔 검색\n원본 296건 분석 보고서')
    p=doc.add_paragraph(f'산출 시각: {now:%Y-%m-%d %H:%M} KST  |  분석 범위: 실제 관측 원본'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    box=doc.add_table(rows=1,cols=1); box.alignment=WD_TABLE_ALIGNMENT.CENTER; shade(box.cell(0,0),'EAF2F8'); box.cell(0,0).text='핵심 결론  제한적 필터는 결과 없음과 강하게 관련되고, 0건 경험 후 대부분은 재검색하지만 바로 다음 검색의 회복률은 17.1%에 그쳤다. 이는 관찰 관련이며 제품 개선 효과를 입증하지 않는다.'
    doc.add_heading('목차',level=1)
    for i,t in enumerate(['프로젝트 문제와 1단계 목적','입력 파일·버전·근거 우선순위','분석 단위와 지표 사전','데이터 구조 및 품질 진단','원본 296건 핵심 결과','A1·A2·B1~B3 검증 결과','세션 결과 세그먼트','H3 재검색 유형 분석','제품 시사점','해석 한계','2단계 진입 판단','재현 방법과 산출물 목록'],1): doc.add_paragraph(f'{i}. {t}')
    section(doc,'1. 프로젝트 문제와 1단계 목적','검색 의도와 실제 노출 결과가 맞지 않을 때 결과 없음·반복 검색·이탈이 발생할 수 있다. 1단계는 원본 296건에서 검색 조건–결과–후속 행동을 재현 가능하게 연결하고 후속 실험 가설의 근거를 정리하는 데 목적이 있다.')
    section(doc,'2. 입력 파일·버전·근거 우선순위'); add_table(doc,['구분','파일/기준'],[['원본 DB','travel_data_filtered_complete_2026-09-03_v02_비식별.sqlite'],['원본 SHA-256','a0cbf893663b99f1a2e4bb8f5e1c202f0a2467f7baccf01f9e858ff54d955571'],['판단 우선순위','현재 DB 재계산 → 교수 회신 → 가설 v08 → 통합 계획 → 초기 계획 → 증강 v06'],['해석 기준','수치는 DB 재계산, 해석은 교수 판단을 우선']])
    section(doc,'3. 분석 단위와 지표 사전'); add_table(doc,['단위','분모','주요 지표'],[['검색',296,'0건률, 필터별·지역별·의도별 지표'],['0건→다음 검색 전이',140,'즉시 회복, H3'],['세션',43,'세션 최종 회복, 결과 세그먼트'],['0건 경험 세션',28,'B2, B3 최종 회복']]); doc.add_paragraph('hotel_click을 호텔 상세진입 KPI로 사용하며 hotel_detail_view를 별도 KPI로 더하지 않는다.')
    section(doc,'4. 데이터 구조 및 품질 진단','8개 업무 테이블의 행 수는 USER 89, HOTEL 1,000, ROOM 3,000, SEARCH 296, SEARCH_FILTER 296, SEARCH_RESULT 8,555, EVENT 10,432, BOOKING 36으로 기준값과 일치했다. SEARCH–SEARCH_FILTER 1:1, 결과 수 합계, 고아키는 통과했다. 다만 숙박일 역전 9건, 검색결과에 없는 클릭 2건, BOOKING–ROOM 호텔 불일치 2건을 파생 플래그로 격리했다. ROOM은 객실 상품/타입이므로 동일 room_id 기간 중첩을 자동 오류로 판정하지 않았다.')
    section(doc,'5. 원본 296건 핵심 결과'); add_picture(doc,charts[0],'그림 1. 원본 296건 핵심 지표'); doc.add_paragraph('결과 없음률은 147/296=49.7%였다. 0건 후 후속검색률은 140/147=95.2%이지만 즉시 회복은 24/140=17.1%였다. 0건을 경험한 28개 세션 중 21개(75.0%)는 세션 마지막에 비0건으로 회복했다. 전체 검색의 hotel_click 기준 상세진입률은 45/296=15.2%였다.')
    section(doc,'6. A1·A2·B1~B3 검증 결과'); add_picture(doc,charts[1],'그림 2. A1 필터별 결과 없음률'); doc.add_paragraph('A1은 부분 채택이다. 편의시설 3개 이상 OR=36.94(95% CI 18.98–71.90), 최소평점 설정 OR=10.88(6.34–18.67), 가격 설정 OR=7.05(4.23–11.75)였고 모두 Fisher 양측 p<.001이었다. 입력 상태를 구분할 컬럼이 없어 조건 제한 효과와 입력 품질 효과를 분리하지 못했다. A3는 제외하며 빈 query_text를 오타로 분류하지 않았다.'); add_picture(doc,charts[2],'그림 3. A2 지역별 결과 없음률'); add_picture(doc,charts[3],'그림 4. A2 검색의도별 결과 없음률'); doc.add_paragraph('A2는 조건부 진행이다. QUALITY_FILTER n=3, PRICE n=9와 같은 희소 셀은 기술통계로만 제시하며 지역·의도의 우열이나 인과를 단정하지 않는다. B1은 0건 후 후속검색 95.2%, 비0건 후 75.8%, Fisher OR=6.37(95% CI 2.73–14.86), p<.001이다. B2는 0건 경험 세션 평균 9.00회(중앙 7, IQR 5.25), 미경험 2.93회(중앙 2, IQR 1.50), U=382, p<.001이다. 탐색지속성 교란 가능성 때문에 인과로 해석하지 않는다. B3의 즉시 회복과 세션 최종 회복은 서로 다른 분모로 구분했다.')
    section(doc,'7. 세션 결과 세그먼트'); add_picture(doc,charts[4],'그림 5. 상호배타 4개 세션 결과 세그먼트'); add_table(doc,['세그먼트','n','비율'],[[r.group,int(r.n),f'{r.rate*100:.1f}%'] for _,r in seg.iterrows()])
    section(doc,'8. H3 재검색 유형 분석'); add_picture(doc,charts[5],'그림 6. H3 재검색 유형별 다음 검색 성공·상세진입'); doc.add_paragraph('지역 변경의 다음 검색 결과 발생률이 10/24=41.7%로 가장 높았고, 검색어 수정 3/10=30.0%, 조건 완화 11/41=26.8%의 순이었다. hotel_click률은 검색어 수정 30.0%, 조건 완화 19.5%, 지역 변경 12.5%였다. H3는 탐색적 결과이며 혼합 변경 n=2 등 소표본과 분류 규칙에 민감하다. 제품 개선의 인과효과로 단정하지 않는다.')
    section(doc,'9. 제품 시사점','후속 실험 가설은 “0건 사용자에게 같은 도시의 인접 지역과 단계적으로 완화 가능한 조건을 제시하면 동일조건 반복을 줄이고 다음 검색의 결과 회복과 호텔 상세진입을 높일 수 있다”이다. 현재 결과는 실험 후보를 지정할 관찰 근거일 뿐, 해당 기능의 효과를 입증하지 않았다.')
    section(doc,'10. 해석 한계');
    for text in ['원본은 세션 43개로 작고 지역×의도 희소 셀이 있다.','typo_flag/입력완료 상태가 없어 A1의 필터 효과와 입력 품질을 분리할 수 없다. A3는 제외했다.','B2와 H3는 자기선택·탐색지속성의 교란 가능성이 있다.','hotel_click과 hotel_detail_view는 동반 로그로 보고 중복 KPI로 사용하지 않았다.','BOOKING 36건은 참고·무결성 점검용이며 실제 예약 전환율로 확대 해석하지 않았다.']: doc.add_paragraph(text,style='List Bullet')
    section(doc,'11. 2단계 진입 판단','조건부 진입이다. 원본 관찰 결과와 품질 이슈를 기준선으로 동결하고, 후속 단계에서는 관찰값과 가정을 물리적·표시상 분리해야 한다. 1,000명은 구조·논리 검증 단계, 10,000명은 희소집단·시나리오 비교 단계의 명칭으로만 기록하며 이번 작업에서 생성·분석하지 않았다.')
    section(doc,'12. 재현 방법과 산출물 목록'); doc.add_paragraph('python "1단계_원본_분석_및_가설_검증/scripts/analyze_original_296.py" --db "03_데이터모델링/이전버전/데이터셋/2026-09-03_v02/travel_data_filtered_complete_2026-09-03_v02_비식별.sqlite" --output-dir "1단계_원본_분석_및_가설_검증/local_outputs/20260903_original_296_analysis"',style=None); add_table(doc,['산출물','설명'],[[workbook.name,'9개 정규화 분석 시트'],[report.name,'제출·발표용 Word 보고서'],*[[p.name,f'차트 {i+1}'] for i,p in enumerate(charts)]])
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; footer.add_run('일본 호텔 검색 원본 296건 분석 | 실제 관측과 후속 시나리오 계획 분리')
    doc.save(report)
    manifest={'timestamp_kst':now.isoformat(),'workbook':workbook.name,'report':report.name,'charts':[p.name for p in charts],'synthetic_generated':False}
    (out/f'호텔검색_원본분석_산출물명세_{stamp}_01.json').write_text(json.dumps(manifest,ensure_ascii=False,indent=2),encoding='utf-8')
    print(json.dumps(manifest,ensure_ascii=False,indent=2))
if __name__=='__main__': main()
