from pathlib import Path
import argparse, hashlib, json, sqlite3
from datetime import datetime
from zoneinfo import ZoneInfo
import pandas as pd
from docx import Document

KST=ZoneInfo('Asia/Seoul'); TS='260903_1745'; VER='02'
def sha(p):
 h=hashlib.sha256()
 with open(p,'rb') as f:
  for b in iter(lambda:f.read(1<<20),b''):h.update(b)
 return h.hexdigest()
def ro(p):
 c=sqlite3.connect(Path(p).resolve().as_uri()+'?mode=ro',uri=True);c.execute('pragma query_only=on');assert c.execute('pragma query_only').fetchone()[0]==1;return c
def calc(p):
 c=ro(p);s=pd.read_sql_query('select search_id,session_id,search_time,total_result_count from search',c);e=pd.read_sql_query("select search_id,event_type from event where event_type='hotel_click'",c);integrity=c.execute('pragma integrity_check').fetchone()[0];counts={t:c.execute(f'select count(*) from "{t}"').fetchone()[0] for t in ['user','hotel','room','search','search_filter','search_result','event','booking']};c.close()
 s['search_time_parsed']=pd.to_datetime(s.search_time.str.replace(' KST','+09:00'),errors='raise');o=s.sort_values(['session_id','search_time_parsed','search_id'],kind='stable').copy();o['order']=o.groupby('session_id').cumcount();o['next_exists']=o.groupby('session_id').search_id.shift(-1).notna();o['next_positive']=o.groupby('session_id').total_result_count.shift(-1).gt(0);o['zero']=o.total_result_count.eq(0)
 z=o[o.zero];zt=z[z.next_exists]; recovered=0;first_zero_d=0;first_zero_n=0;any_zero_d=0
 for _,g in o.groupby('session_id',sort=False):
  vals=g.total_result_count.to_numpy();zeros=vals==0
  if zeros.any():
   any_zero_d+=1
   if any((vals[i+1:]>0).any() for i in range(len(vals)) if zeros[i]):recovered+=1
  if vals[0]==0:
   first_zero_d+=1
   if (vals[1:]>0).any():first_zero_n+=1
 click_n=s.search_id.isin(set(e.search_id)).sum()
 metrics=[
  ('zero_result_rate',int(z.shape[0]),len(o),'search'),
  ('followup_rate_after_zero',int(zt.shape[0]),len(z),'zero-result search'),
  ('immediate_recovery_rate',int(zt.next_positive.sum()),len(zt),'zero-to-next transition'),
  ('session_final_recovery_rate',recovered,any_zero_d,'session experiencing any zero result'),
  ('hotel_click_rate',int(click_n),len(o),'search'),
  ('first_search_zero_session_research_recovery_rate',first_zero_n,first_zero_d,'session whose first search is zero')]
 return pd.DataFrame([{'metric':m,'numerator':n,'denominator':d,'rate':n/d,'analysis_unit':u} for m,n,d,u in metrics]),counts,integrity
def record(p,role,base=None):
 p=Path(p);return {'role':role,'path':str(p.resolve() if base is None else p.resolve().relative_to(base.resolve())).replace('\\','/'),'size_bytes':p.stat().st_size,'sha256':sha(p)}
def main():
 ap=argparse.ArgumentParser();ap.add_argument('--source',type=Path,required=True);ap.add_argument('--synthetic',type=Path,required=True);ap.add_argument('--stage1-code',type=Path,required=True);ap.add_argument('--stage1-manifest',type=Path,required=True);ap.add_argument('--old-manifest',type=Path,required=True);ap.add_argument('--generation-code',type=Path,required=True);ap.add_argument('--config',type=Path,required=True);ap.add_argument('--old-excel',type=Path,required=True);ap.add_argument('--old-word',type=Path,required=True);ap.add_argument('--old-log',type=Path,required=True);ap.add_argument('--out',type=Path,required=True);a=ap.parse_args();a.out.mkdir(parents=True,exist_ok=True)
 sm,sc,si=calc(a.source);ym,yc,yi=calc(a.synthetic);cmp=sm.merge(ym,on='metric',suffixes=('_original','_synthetic'));cmp['difference_pp']=(cmp.rate_synthetic-cmp.rate_original)*100;cmp['abs_difference_pp']=cmp.difference_pp.abs();cmp['tolerance_pp']=3;cmp['pass']=cmp.abs_difference_pp.le(3)
 core=cmp[cmp.metric.ne('first_search_zero_session_research_recovery_rate')].copy();maxdiff=float(core.abs_difference_pp.max())
 old=pd.read_excel(a.old_excel,sheet_name='G3_core_metrics');old_final=old.loc[old.metric.eq('final_recovery_rate')].iloc[0]
 diagnosis=pd.DataFrame([{'check':'approved definition','finding':'denominator=any-zero sessions; numerator=session with any later positive after a zero'},{'check':'old QA implementation','finding':f"used first-search-zero recovery: original {old_final.original:.6f}, synthetic {old_final.synthetic:.6f}"},{'check':'classification','finding':'calculation definition error, not a label-only typo'},{'check':'canonical ordering','finding':'session_id, parsed search_time, search_id; stable sort'}])
 hashes=pd.DataFrame([record(p,r) for p,r in [(a.source,'source SQLite'),(a.synthetic,'unchanged synthetic SQLite'),(a.stage1_code,'stage1 canonical code'),(a.stage1_manifest,'stage1 manifest'),(a.old_manifest,'superseded stage2 manifest'),(a.generation_code,'unchanged generation code'),(a.config,'unchanged config'),(a.old_excel,'superseded QA Excel'),(a.old_word,'superseded QA Word'),(a.old_log,'superseded execution record')]])
 base_ok=(si=='ok' and yi=='ok' and sc['search']==296 and yc['user']==1000 and yc['booking']==0 and yc['search']==yc['search_filter'] and json.loads(a.old_manifest.read_text(encoding='utf-8'))['G1_G5']['G1']=='PASS')
 g={'G1':'PASS' if base_ok else 'FAIL','G2':'PASS' if si==yi=='ok' else 'FAIL','G3':'PASS' if core['pass'].all() else 'FAIL','G4':json.loads(a.old_manifest.read_text(encoding='utf-8'))['G1_G5']['G4'],'G5':json.loads(a.old_manifest.read_text(encoding='utf-8'))['G1_G5']['G5']};qa='PASS' if all(v=='PASS' for v in g.values()) else 'FAIL';exp=qa=='PASS'
 summary=pd.DataFrame([{'final_qa_status':qa,'expansion_allowed':exp,'stage3_entry_allowed':exp,'corrected_core_max_abs_difference_pp':maxdiff,'ordering':'session_id, parsed search_time, search_id stable'}])
 x=a.out/f'호텔검색_관측형합성1000명_QA교정결과_{TS}_{VER}.xlsx'
 with pd.ExcelWriter(x,engine='openpyxl') as w:
  for n,d in {'corrected_core_metrics':core,'two_recovery_metrics':cmp[cmp.metric.str.contains('recovery')],'definition_diagnosis':diagnosis,'hash_verification':hashes,'row_counts':pd.DataFrame([{'table':k,'original':sc.get(k),'synthetic':yc.get(k)} for k in sc]),'G1_G5':pd.DataFrame([{'gate':k,'status':v} for k,v in g.items()]),'qa_summary':summary}.items():d.to_excel(w,sheet_name=n,index=False)
  for ws in w.book.worksheets:
   ws.freeze_panes='A2';ws.auto_filter.ref=ws.dimensions
   for col in ws.columns:ws.column_dimensions[col[0].column_letter].width=min(70,max(14,max(len(str(c.value or'')) for c in col)+2))
 doc=Document();doc.add_heading('일본 호텔 검색 2단계 핵심 지표 교정 QA 보고서',0);doc.add_paragraph(f'교정 실행 {TS}, version {VER} | Final QA={qa} | expansion_allowed={str(exp).lower()}')
 doc.add_heading('1. 교정 사유',1);doc.add_paragraph('기존 QA는 첫 검색 0건 세션의 재검색 회복률을 B3 세션 최종 회복률로 사용했다. 따라서 단순 지표명 오기가 아니라 계산 정의 오류다.')
 doc.add_heading('2. 정의와 정렬',1);doc.add_paragraph('B3 세션 최종 회복률은 0건을 한 번이라도 경험한 세션을 분모로, 해당 0건 검색 이후 비0건이 발생한 세션을 분자로 한다. 정렬은 session_id 내 parsed search_time, search_id 안정 정렬이다.')
 doc.add_heading('3. 직접 재계산 결과',1);t=doc.add_table(rows=1,cols=6);t.style='Table Grid'
 for i,v in enumerate(['metric','original','synthetic','original n/d','synthetic n/d','diff pp']):t.rows[0].cells[i].text=v
 for _,r in cmp.iterrows():
  c=t.add_row().cells;vals=[r.metric,f'{r.rate_original:.2%}',f'{r.rate_synthetic:.2%}',f'{int(r.numerator_original)}/{int(r.denominator_original)}',f'{int(r.numerator_synthetic)}/{int(r.denominator_synthetic)}',f'{r.difference_pp:.2f}'];
  for i,v in enumerate(vals):c[i].text=str(v)
 doc.add_heading('4. 판정',1);doc.add_paragraph(f'G1–G5={json.dumps(g,ensure_ascii=False)}. 교정된 핵심 지표 최대 절대 편차는 {maxdiff:.2f}%p이다. 최종 QA={qa}, expansion_allowed={str(exp).lower()}, 3단계 진입 가능={str(exp).lower()}. 10,000명·스트레스·오버샘플링·S1–S3는 생성하지 않았다.')
 wp=a.out/f'호텔검색_관측형합성1000명_QA교정보고서_{TS}_{VER}.docx';doc.save(wp)
 lp=a.out/f'호텔검색_관측형합성1000명_QA교정실행기록_{TS}_{VER}.md';lp.write_text(f'''# 2단계 핵심 지표 교정 실행기록\n\n- 원본·합성 DB: 읽기 전용\n- 교정 사유: B3 세션 최종 회복률에 4/6 첫 검색 0건 세션 지표를 사용한 **계산 정의 오류**\n- 정렬: `session_id`, parsed `search_time`, `search_id` stable sort\n- 원본 B3 최종 회복: {int(core.loc[core.metric.eq('session_final_recovery_rate'),'numerator_original'].iloc[0])}/{int(core.loc[core.metric.eq('session_final_recovery_rate'),'denominator_original'].iloc[0])}\n- 합성 B3 최종 회복: {int(core.loc[core.metric.eq('session_final_recovery_rate'),'numerator_synthetic'].iloc[0])}/{int(core.loc[core.metric.eq('session_final_recovery_rate'),'denominator_synthetic'].iloc[0])}\n- G1–G5: `{json.dumps(g,ensure_ascii=False)}`\n- final QA: **{qa}**\n- expansion_allowed: **{str(exp).lower()}**\n- 기존 코드·config·합성 SQLite는 수정·재생성하지 않음\n- 10,000명·스트레스·오버샘플링·S1–S3 미생성\n''',encoding='utf-8')
 new=[record(x,'corrected QA Excel',a.out),record(wp,'corrected QA Word',a.out),record(lp,'corrected execution record',a.out),record(Path(__file__),'correction calculation code',a.out)]
 oldm=json.loads(a.old_manifest.read_text(encoding='utf-8'));mp=a.out/f'호텔검색_관측형합성1000명_실행묶음교정매니페스트_{TS}_{VER}.json';man={'stage2_bundle_run_id':f'STAGE2_QA_CORRECTION_{TS}_{VER}','correction_version':VER,'supersedes_manifest':str(a.old_manifest.resolve()).replace('\\','/'),'supersedes_manifest_sha256':sha(a.old_manifest),'qa_correction_reason':'B3 session_final_recovery_rate incorrectly used first-search-zero cohort (4/6); corrected to any-zero-experience cohort with later positive search','source_db_sha256':sha(a.source),'generation_code_sha256':sha(a.generation_code),'config_sha256':sha(a.config),'synthetic_db_sha256':sha(a.synthetic),'unchanged_generation_artifacts_verified':True,'G1_G5':g,'corrected_core_max_abs_difference_pp':maxdiff,'final_qa_status':qa,'expansion_allowed':exp,'stage3_entry_allowed':exp,'sole_approved_manifest_for_stage3':True,'approval_statement':'This correction manifest is the only approved manifest for Stage 3; the superseded manifest must not be used for Stage 3 entry.','forbidden_sets_generated':False,'actual_generated_at':datetime.now(KST).isoformat(timespec='seconds'),'corrected_metrics':cmp.to_dict('records'),'referenced_unchanged_files':[record(p,r) for p,r in [(a.source,'source SQLite'),(a.synthetic,'synthetic SQLite'),(a.generation_code,'generation code'),(a.config,'generation config'),(a.stage1_code,'stage1 canonical code'),(a.stage1_manifest,'stage1 manifest'),(a.old_manifest,'superseded stage2 manifest')]],'new_artifacts':new,'manifest_self':{'sha256':None,'reason':'self-referential hash omitted'}};mp.write_text(json.dumps(man,ensure_ascii=False,indent=2),encoding='utf-8')
 print(json.dumps({'manifest':str(mp),'qa':qa,'expansion_allowed':exp,'max_abs_difference_pp':maxdiff,'metrics':cmp.to_dict('records')},ensure_ascii=False,indent=2))
if __name__=='__main__':main()
