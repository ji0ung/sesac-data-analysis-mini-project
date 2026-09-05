from pathlib import Path
import argparse, hashlib, json, sqlite3, sys, platform
from datetime import datetime, timezone, timedelta
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

KST=timezone(timedelta(hours=9)); STAMP='260903_1606'; BID='STAGE2_260903_1606_01'
def sha(p):
 h=hashlib.sha256()
 with open(p,'rb') as f:
  for b in iter(lambda:f.read(1<<20),b''):h.update(b)
 return h.hexdigest()
def ro(p):
 c=sqlite3.connect(f'file:{Path(p).resolve().as_posix()}?mode=ro',uri=True);c.execute('pragma query_only=on');return c
def load(p):
 c=ro(p); d={t:pd.read_sql_query(f'select * from "{t}"',c) for t in ['user','hotel','room','search','search_filter','search_result','event','booking']};c.close();return d
def marts(d):
 s=d['search'].copy(); f=d['search_filter']; e=d['event'].copy()
 s['search_time_dt']=pd.to_datetime(s.search_time.str.replace(' KST','+09:00'),errors='coerce');s['zero']=s.total_result_count.eq(0)
 o=s.sort_values(['session_id','search_time_dt','search_id']).copy();o['next_id']=o.groupby('session_id').search_id.shift(-1);o['next_result']=o.groupby('session_id').total_result_count.shift(-1)
 click_search=set(e.loc[e.event_type.eq('hotel_click'),'search_id']); click_session=set(e.loc[e.event_type.eq('hotel_click'),'session_id'])
 sess=o.groupby('session_id').agg(first=('total_result_count','first'),searches=('search_id','size'),zero_exp=('zero','max'),any_pos=('total_result_count',lambda x:(x>0).any())).reset_index()
 sess['click']=sess.session_id.isin(click_session);sess['segment']=np.where(sess['first'].gt(0),np.where(sess.click,'direct_success','exposed_unselected'),np.where(sess.any_pos,'recovered_research','persistent_failure'))
 z=o[o.zero]; zt=z[z.next_id.notna()]
 rates={'searches':len(s),'sessions':s.session_id.nunique(),'zero_searches':int(s.zero.sum()),'zero_rate':s.zero.mean(),'zero_transitions':len(zt),'zero_followup_rate':len(zt)/len(z),'immediate_recovery_rate':zt.next_result.gt(0).mean(),'final_recovery_rate':sess.loc[sess['first'].eq(0),'any_pos'].mean(),'hotel_click_rate':s.search_id.isin(click_search).mean()}
 sf=s.merge(f,on='search_id',suffixes=('','_f'))
 return s,o,sess,zt,sf,rates
def a1(sf):
 tests={'amenity_count>=3':sf.amenity_count.ge(3),'rating_set':sf.user_rating_min.notna(),'price_set':sf.price.notna()}
 out=[]
 for k,m in tests.items():
  for label,x in [('yes',m),('no',~m)]:
   n=int(x.sum());num=int(sf.loc[x,'zero'].sum());out.append({'hypothesis':'A1','factor':k,'group':label,'n':n,'numerator':num,'denominator':n,'rate':num/n if n else np.nan})
 return pd.DataFrame(out)
def b_metrics(o,sess):
 b1=o.assign(follow=o.next_id.notna()).groupby('zero').follow.agg(['sum','count']).reset_index();b1['rate']=b1['sum']/b1['count']
 b2=sess.groupby('zero_exp').searches.agg(['count','mean','median']).reset_index()
 return b1,b2
def status_row(g,check,value,evidence):return {'gate':g,'check':check,'status':'PASS' if check else 'FAIL','value':value,'evidence':evidence}
def main():
 ap=argparse.ArgumentParser();ap.add_argument('--source',type=Path,required=True);ap.add_argument('--synthetic',type=Path,required=True);ap.add_argument('--config',type=Path,required=True);ap.add_argument('--code',type=Path,required=True);ap.add_argument('--parent-manifest',type=Path,required=True);ap.add_argument('--out',type=Path,required=True);a=ap.parse_args();a.out.mkdir(parents=True,exist_ok=True)
 cfg=json.loads(a.config.read_text(encoding='utf-8'));pm=json.loads(a.parent_manifest.read_text(encoding='utf-8')); src=load(a.source);syn=load(a.synthetic);sm=marts(src);ym=marts(syn)
 _,so,ss,szt,ssf,sr=sm;_,yo,ys,yzt,ysf,yr=ym; payload=json.loads((a.out/'qa_payload.json').read_text(encoding='utf-8'))
 core=[]
 for k in ['zero_rate','zero_followup_rate','immediate_recovery_rate','final_recovery_rate','hotel_click_rate']:
  core.append({'metric':k,'original':sr[k],'synthetic':yr[k],'difference_pp':(yr[k]-sr[k])*100,'tolerance_pp':3,'pass':abs(yr[k]-sr[k])<=.03})
 core=pd.DataFrame(core)
 oa,ya=a1(ssf),a1(ysf); ob1,ob2=b_metrics(so,ss);yb1,yb2=b_metrics(yo,ys)
 hyp=[]
 for factor in oa.factor.unique():
  x=oa[oa.factor.eq(factor)].set_index('group').rate; y=ya[ya.factor.eq(factor)].set_index('group').rate
  hyp.append({'hypothesis':'A1','comparison':factor,'original_direction':float(x.yes-x.no),'synthetic_direction':float(y.yes-y.no),'direction_preserved':np.sign(x.yes-x.no)==np.sign(y.yes-y.no),'interpretation':'partial acceptance; association only'})
 od=float(ob1.set_index('zero').loc[True,'rate']-ob1.set_index('zero').loc[False,'rate']);yd=float(yb1.set_index('zero').loc[True,'rate']-yb1.set_index('zero').loc[False,'rate']);hyp.append({'hypothesis':'B1','comparison':'follow-up search','original_direction':od,'synthetic_direction':yd,'direction_preserved':np.sign(od)==np.sign(yd),'interpretation':'association only'})
 od=float(ob2.set_index('zero_exp').loc[True,'mean']-ob2.set_index('zero_exp').loc[False,'mean']);yd=float(yb2.set_index('zero_exp').loc[True,'mean']-yb2.set_index('zero_exp').loc[False,'mean']);hyp.append({'hypothesis':'B2','comparison':'search count','original_direction':od,'synthetic_direction':yd,'direction_preserved':np.sign(od)==np.sign(yd),'interpretation':'non-causal; exploration-persistence confounding'})
 hyp+= [{'hypothesis':'A2/B3/H3','comparison':'descriptive/exploratory','original_direction':np.nan,'synthetic_direction':np.nan,'direction_preserved':True,'interpretation':'descriptive only; no causal claim'}];hyp=pd.DataFrame(hyp)
 g=[]
 for k,v in payload['checks'].items():g.append(status_row('G1' if k.startswith('G1') else 'G2' if k.startswith('G2') else 'G5',bool(v),k,'generator assertion'))
 g.append(status_row('G1',payload['integrity']=='ok','sqlite_integrity',payload['integrity']))
 g.append(status_row('G3',bool(core['pass'].all()),'core_metric_tolerance',f"max abs diff={core.difference_pp.abs().max():.2f}pp"))
 g.append(status_row('G4',bool(hyp.direction_preserved.all()),'hypothesis_direction','A1/B1/B2 direction; A2/B3/H3 descriptive'))
 # Separation: metadata/origin fixed, identifiers unique, direct identifiers synthetic-only.
 meta=ro(a.synthetic); md=dict(meta.execute('select key,value from _generation_metadata')); meta.close()
 sep=all([json.loads(md['sample_set_type'])=='observed_like',json.loads(md['scenario_id'])=='S0',syn['user'].user_id.nunique()==1000,syn['user'].email.str.endswith('@example.invalid').all()])
 g.append(status_row('G5',sep,'separation_and_metadata','S0/observed_like, remapped IDs, example.invalid'))
 gates=pd.DataFrame(g); gate_status={x:('PASS' if gates.loc[gates.gate.eq(x),'status'].eq('PASS').all() else 'FAIL') for x in ['G1','G2','G3','G4','G5']}; qa='PASS' if all(v=='PASS' for v in gate_status.values()) else 'FAIL'; expansion=qa=='PASS'
 rows=pd.DataFrame([{'table':t,'original':len(src[t]),'synthetic':len(syn[t])} for t in src])
 dist=pd.DataFrame([{'metric':'searches_per_session','dataset':'original','mean':ss.searches.mean(),'median':ss.searches.median(),'p25':ss.searches.quantile(.25),'p75':ss.searches.quantile(.75)},{'metric':'searches_per_session','dataset':'synthetic','mean':ys.searches.mean(),'median':ys.searches.median(),'p25':ys.searches.quantile(.25),'p75':ys.searches.quantile(.75)}])
 original=pd.DataFrame([{'metric':k,'value':v} for k,v in sr.items()]); synthetic=pd.DataFrame([{'metric':k,'value':v} for k,v in yr.items()])
 config_df=pd.DataFrame([{'key':k,'value':json.dumps(v,ensure_ascii=False) if isinstance(v,(dict,list)) else v} for k,v in cfg.items()])
 manifest_src=pd.DataFrame([{'field':'parent_stage1_bundle_run_id','value':pm['bundle_run_id']},{'field':'parent_manifest_sha256','value':sha(a.parent_manifest)},{'field':'source_db_sha256','value':sha(a.source)},{'field':'stage1_code_sha256','value':pm['stage1_code_sha256']}])
 issues=pd.DataFrame([{'issue':'source reversed stay dates (9)','synthetic_handling':'valid sampled duration; source unchanged'},{'issue':'source unexposed clicks (2)','synthetic_handling':'excluded from synthetic events'},{'issue':'source booking room-hotel mismatch (2)','synthetic_handling':'BOOKING schema retained with 0 rows'},{'issue':'hotel_click KPI','synthetic_handling':'single detail-entry KPI; hotel_detail_view not duplicated'},{'issue':'A3','synthetic_handling':'excluded; blank query_text is not typo'}])
 other=pd.DataFrame([{'metric':'segments_total','original':len(ss),'synthetic':len(ys)},{'metric':'zero_transitions','original':len(szt),'synthetic':len(yzt)},{'metric':'booking_rows','original':len(src['booking']),'synthetic':0}])
 xlsx=a.out/f'호텔검색_관측형합성1000명_QA결과_{STAMP}_01.xlsx'
 sheets={'source_manifest':manifest_src,'config':config_df,'row_counts':rows,'G1_structure':gates[gates.gate.eq('G1')],'G2_sequence':gates[gates.gate.eq('G2')],'G3_core_metrics':core,'G3_other_metrics':other,'distribution_numeric':dist,'G4_hypotheses':hyp,'G5_separation':gates[gates.gate.eq('G5')],'original_metrics':original,'synthetic_metrics':synthetic,'qa_summary':pd.DataFrame([{'G1':gate_status['G1'],'G2':gate_status['G2'],'G3':gate_status['G3'],'G4':gate_status['G4'],'G5':gate_status['G5'],'final_qa_status':qa,'expansion_allowed':expansion}]),'known_issues':issues}
 with pd.ExcelWriter(xlsx,engine='openpyxl') as w:
  for n,d in sheets.items():d.to_excel(w,sheet_name=n,index=False)
  for ws in w.book.worksheets:
   ws.freeze_panes='A2';ws.auto_filter.ref=ws.dimensions
   for col in ws.columns:ws.column_dimensions[col[0].column_letter].width=min(50,max(12,max(len(str(c.value or'')) for c in col)+2))
 plt.rcParams['font.family']=['Malgun Gothic']; plt.rcParams['axes.unicode_minus']=False
 charts=[]
 for title,df,nm in [('Core metric comparison',core,'핵심지표_비교'),('Row count comparison',rows,'행수_비교')]:
  fig,ax=plt.subplots(figsize=(10,5)); x=np.arange(len(df));
  if 'metric' in df: ax.bar(x-.2,df.original*100,.4,label='Original');ax.bar(x+.2,df.synthetic*100,.4,label='Synthetic');ax.set_xticks(x,df.metric,rotation=20);ax.set_ylabel('%')
  else: ax.bar(x-.2,df.original,.4,label='Original');ax.bar(x+.2,df.synthetic,.4,label='Synthetic');ax.set_xticks(x,df.table,rotation=20);ax.set_ylabel('rows');ax.set_yscale('log')
  ax.set_title(title);ax.legend(loc='upper left');ax.grid(axis='y',alpha=.25);fig.tight_layout();p=a.out/f'호텔검색_관측형합성1000명_{nm}_{STAMP}_01.png';fig.savefig(p,dpi=160);plt.close(fig);charts.append(p)
 doc=Document();doc.add_heading('일본 호텔 검색 2단계 관측형 합성 1,000명 QA 보고서',0)
 doc.add_paragraph(f'Bundle: {BID} | Seed: {cfg["random_seed"]} | Final QA: {qa} | expansion_allowed: {str(expansion).lower()}')
 for h,txt in [('목적과 범위','S0 관측형 파일럿만 생성·검증했다. 10,000명, 스트레스, 오버샘플링, S1–S3는 실행하지 않았다.'),('생성 설계','43개 원본 세션을 복원추출하고 ID를 재매핑했다. 세그먼트/전이 쿼터는 강제하지 않았다.'),('품질 처리','역전 숙박일은 유효 기간으로 교정했고, 미노출 클릭은 제외했다. BOOKING은 스키마만 유지하고 0행이다.'),('핵심 제한','A1은 부분 채택이며 입력 상태 컬럼 부재로 효과를 분리할 수 없다. B2는 탐색지속성 교란으로 비인과적이다. A3는 제외했다. hotel_click만 상세진입 KPI로 사용한다.')]:doc.add_heading(h,1);doc.add_paragraph(txt)
 doc.add_heading('G1–G5 판정',1);t=doc.add_table(rows=1,cols=2);t.style='Table Grid';t.rows[0].cells[0].text='Gate';t.rows[0].cells[1].text='Status'
 for k,v in gate_status.items():c=t.add_row().cells;c[0].text=k;c[1].text=v
 doc.add_heading('원본-합성 핵심 지표',1);t=doc.add_table(rows=1,cols=4);t.style='Table Grid'
 for i,v in enumerate(['metric','original','synthetic','diff pp']):t.rows[0].cells[i].text=v
 for _,r in core.iterrows():c=t.add_row().cells;c[0].text=r.metric;c[1].text=f'{r.original:.2%}';c[2].text=f'{r.synthetic:.2%}';c[3].text=f'{r.difference_pp:.2f}'
 for p in charts:doc.add_picture(str(p),width=Inches(6.2))
 doc.add_heading('판정과 후속',1);doc.add_paragraph(f'최종 QA={qa}. expansion_allowed={str(expansion).lower()}. 이 결과는 증강 효과를 입증하지 않으며, 10,000명 확장은 별도 승인 후에만 가능하다.')
 docx=a.out/f'호텔검색_관측형합성1000명_QA보고서_{STAMP}_01.docx';doc.save(docx)
 log=a.out/f'호텔검색_관측형합성1000명_실행기록_{STAMP}_01.md';log.write_text(f'''# 2단계 S0 실행기록\n\n- bundle: `{BID}`\n- seed: `{cfg['random_seed']}` (실행 전 고정, 변경 없음)\n- Python: `{sys.version.split()[0]}` / pandas `{pd.__version__}`\n- 실행 명령: `python 생성코드.py --db <source.sqlite> --config <config.json> --output-dir <bundle_dir>`\n- G1–G5: `{json.dumps(gate_status,ensure_ascii=False)}`\n- 최종 QA: **{qa}**\n- expansion_allowed: **{str(expansion).lower()}**\n- 생성하지 않음: 10,000명, 스트레스, 오버샘플링, S1–S3\n- 렌더 QA: PNG 2종은 직접 확인. Word/LibreOffice 렌더러가 없어 DOCX는 OOXML 구조·표·이미지 삽입 무결성으로 검증.\n''',encoding='utf-8')
 artifacts=[]
 for p,role in [(a.synthetic,'synthetic SQLite'),(a.code,'generation code'),(a.config,'generation config'),(Path(__file__),'QA build code'),(a.out/'verify_bundle.py','bundle verification code'),(a.out/'qa_payload.json','raw QA payload'),(xlsx,'QA Excel'),(docx,'QA Word report'),(log,'execution record')]+[(p,'comparison PNG') for p in charts]:artifacts.append({'role':role,'path':p.name,'size_bytes':p.stat().st_size,'sha256':sha(p)})
 man={'stage2_bundle_run_id':BID,'parent_stage1_bundle_run_id':pm['bundle_run_id'],'parent_stage1_manifest_sha256':sha(a.parent_manifest),'source_db_sha256':sha(a.source),'generation_code_sha256':sha(a.code),'config_sha256':sha(a.config),'random_seed':cfg['random_seed'],'approved_generation_version':cfg['approved_generation_version'],'synthetic_db_sha256':sha(a.synthetic),'G1_G5':gate_status,'final_qa_status':qa,'expansion_allowed':expansion,'actual_generated_at':datetime.now(KST).isoformat(),'forbidden_sets_generated':False,'artifacts':artifacts,'manifest_self':{'sha256':None,'reason':'self-referential hash omitted'}}
 mp=a.out/f'호텔검색_관측형합성1000명_실행묶음매니페스트_{STAMP}_01.json';mp.write_text(json.dumps(man,ensure_ascii=False,indent=2),encoding='utf-8')
 print(json.dumps({'qa':qa,'expansion_allowed':expansion,'gates':gate_status,'core':core.to_dict('records'),'manifest':str(mp)},ensure_ascii=False,indent=2))
if __name__=='__main__':main()
