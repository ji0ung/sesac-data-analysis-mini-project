from pathlib import Path
import hashlib, json, sys, zipfile
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

docx=Path(sys.argv[1]);manifest=Path(sys.argv[2]);log=Path(sys.argv[3]);repo=Path(sys.argv[4])
def sh(p):return hashlib.sha256(Path(p).read_bytes()).hexdigest()
def setfont(run,size=8,bold=False,color=None,name="Malgun Gothic"):
    run.font.name=name;run._element.rPr.rFonts.set(qn("w:eastAsia"),"맑은 고딕");run.font.size=Pt(size);run.bold=bold
    if color:run.font.color.rgb=RGBColor(*color)
def shade(cell,color):
    e=OxmlElement("w:shd");e.set(qn("w:fill"),color);cell._tc.get_or_add_tcPr().append(e)
def heading(d,text,level=2):
    p=d.add_heading(text,level=level)
    for r in p.runs:setfont(r,12 if level==2 else 10,True,(31,78,120))
def para(d,text,size=8.8):
    p=d.add_paragraph();p.paragraph_format.space_after=Pt(4);p.paragraph_format.line_spacing=1.08;r=p.add_run(text);setfont(r,size);return p
def bullets(d,items):
    for x in items:
        p=d.add_paragraph(style="List Bullet");p.paragraph_format.space_after=Pt(2);r=p.add_run(x);setfont(r,8.3)
def numbered(d,items):
    for x in items:
        p=d.add_paragraph(style="List Number");p.paragraph_format.space_after=Pt(3);r=p.add_run(x);setfont(r,8.5)
def table(d,headers,rows,size=7):
    t=d.add_table(rows=1,cols=len(headers));t.style="Table Grid";t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];c.text=str(h);shade(c,"1F4E78")
        for r in c.paragraphs[0].runs:setfont(r,size,True,(255,255,255))
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v);cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(0)
                for r in p.runs:setfont(r,size)
    return t
def code(d,text):
    p=d.add_paragraph();p.paragraph_format.space_after=Pt(5);shade_para=OxmlElement("w:shd");shade_para.set(qn("w:fill"),"F2F2F2");p._p.get_or_add_pPr().append(shade_para)
    r=p.add_run(text);setfont(r,7.5,name="Consolas");return p

d=Document(docx)
if any(p.text=="17.4 작업자가 그대로 실행할 수 있는 증강 재현 절차" for p in d.paragraphs):
    raise RuntimeError("Detailed augmentation procedure already exists")
d.add_page_break();heading(d,"17.4 작업자가 그대로 실행할 수 있는 증강 재현 절차",1)
para(d,"아래 절차는 승인된 생성 코드의 실제 동작을 사람이 다시 구현하거나 동일 코드로 재실행할 수 있도록 풀어 쓴 작업 명세다. 이 보고서 보완 과정에서는 생성기를 실행하지 않았다. 재실행 시에는 반드시 새 빈 출력 폴더를 사용하며 승인된 원본·코드·설정의 해시를 먼저 확인한다.")

heading(d,"17.4.1 준비물과 고정값")
table(d,["구분","파일 또는 값","검증 기준"],[
 ["원본 SQLite","travel_data_filtered_complete_2026-09-03_v02_비식별.sqlite","SHA-256 a0cbf893663b99f1a2e4bb8f5e1c202f0a2467f7baccf01f9e858ff54d955571"],
 ["생성 코드","02_관측형합성1000명_실행묶음_260903_1606_01/호텔검색_관측형합성1000명_생성코드_260903_1606_01.py","현재 CRLF 파일은 LF 정규화 후 승인 SHA-256 b7e8cea58cc3fa653d9bcafe781a1d0be9f910cd0c5d1bc27ac04a2bddd1cf2a"],
 ["생성 설정","02_관측형합성1000명_실행묶음_260903_1606_01/호텔검색_관측형합성1000명_생성설정_260903_1606_01.json","LF 정규화 승인 SHA-256 59e588293184a0a9c7c704752a72f276e1e1184660767cf77c0179fee4987512"],
 ["난수","random_seed=20260903","한 번 고정하며 결과를 맞추기 위해 변경하지 않음"],
 ["목표","n_users=1,000, n_sessions=1,000","사용자 1명당 세션 1개"],
 ["Python 패키지","Python, numpy, pandas, openpyxl, python-docx","생성 핵심은 sqlite3·numpy·pandas"],
],6.6)
para(d,"원본 DB는 SQLite URI mode=ro와 PRAGMA query_only=ON으로 열어야 한다. 실행 직전과 직후에 원본 파일의 SHA-256·크기·수정시각을 비교하여 원본이 바뀌지 않았음을 확인한다.")

heading(d,"17.4.2 실행 전 원본 점검")
numbered(d,[
 "원본 DB의 USER 89, HOTEL 1,000, ROOM 3,000, SEARCH 296, SEARCH_FILTER 296, SEARCH_RESULT 8,555, EVENT 10,432, BOOKING 36행을 확인한다.",
 "PRAGMA integrity_check 결과가 ok인지 확인한다.",
 "SEARCH.search_id와 SEARCH_FILTER.search_id가 1:1이고, SEARCH.total_result_count 합계가 SEARCH_RESULT 8,555행과 일치하는지 확인한다.",
 "세션 수가 43개인지 확인하고, 각 세션 내부 검색을 search_time, search_id 순으로 안정 정렬할 수 있는지 확인한다.",
 "원본의 숙박일 역전 9건과 미노출 클릭 2건은 원본을 수정하지 않고 합성 복제 과정에서만 아래 규칙으로 처리한다.",
])

heading(d,"17.4.3 새 SQLite 스키마 생성")
para(d,"새 출력 SQLite 파일을 만들고 원본 sqlite_master의 CREATE TABLE SQL을 읽어 USER, HOTEL, ROOM, SEARCH, SEARCH_FILTER, SEARCH_RESULT, EVENT, BOOKING의 동일 스키마를 생성한다. 이어서 생성 계보 저장용 _generation_metadata(key TEXT PRIMARY KEY, value TEXT)를 추가한다. 출력 파일이 이미 있으면 덮어쓰지 않고 즉시 중단한다.")
code(d,"for table in [user, hotel, room, search, search_filter, search_result, event, booking]:\n    ddl = source.sqlite_master[table].sql\n    destination.execute(ddl)\ndestination.execute('CREATE TABLE _generation_metadata(key TEXT PRIMARY KEY, value TEXT)')")

heading(d,"17.4.4 HOTEL·ROOM 기준정보 복사")
para(d,"HOTEL 1,000행과 ROOM 3,000행은 새로 합성하지 않는다. 원본의 모든 컬럼과 값을 그대로 새 DB에 복사한다. 따라서 이 두 테이블은 합성 행동의 참조 기준정보이며, ROOM은 물리 객실 재고가 아닌 객실 상품/타입으로 해석한다.")

heading(d,"17.4.5 43개 원본 세션을 1,000개로 균등 확장")
para(d,"편향이 큰 단순 복원추출 대신 모든 원본 세션을 거의 같은 횟수로 사용한다. 1,000을 43으로 나누면 몫 q=23, 나머지 r=11이다.")
code(d,"q, r = divmod(1000, 43)  # q=23, r=11\nchosen = 각 원본 session_id를 23회씩 배열한 989개 목록\nextra = RNG(seed=20260903)로 43개 중 11개를 비복원 추출\nchosen = RNG.permutation(chosen + extra)  # 최종 1,000개")
bullets(d,[
 "각 원본 세션은 최소 23회 사용되고, 선택된 11개 세션만 24회 사용된다.",
 "세그먼트 목표 비율은 강제하지 않는다. 결과 세그먼트는 원본 세션 구조를 복제한 결과로 자연스럽게 나온다.",
 "chosen 배열의 위치 i(1~1,000)가 합성 사용자·세션 번호가 된다.",
])

heading(d,"17.4.6 합성 ID와 기준시각 생성")
table(d,["객체","새 ID 규칙","예시적 형식","관계"],[
 ["사용자","SYN_U + i 4자리","SYN_U0001","합성 세션 1개에 사용자 1명"],
 ["세션","SYN_S + i 4자리","SYN_S0001","해당 세션의 SEARCH·EVENT에 공통"],
 ["검색","SYN_Q + i 4자리 + j 3자리","SYN_Q0001_001","세션 내부 j번째 검색"],
 ["필터","SYN_F + i 4자리 + j 3자리","SYN_F0001_001","j번째 검색과 1:1"],
 ["검색결과","SYN_R + i 4자리 + j 3자리 + k 4자리","SYN_R0001_001_0001","j번째 검색의 result_rank 순 k번째"],
 ["이벤트","SYN_E + i 4자리 + k 5자리","SYN_E0001_00001","세션 내부 event_at·event_id 순 k번째"],
],6.7)
para(d,"합성 세션 i의 기준시각은 2027-01-01 00:00:00 KST + floor(i/4)일 + (i mod 4)분이다. 이 방식은 합성 세션의 기준시각을 결정적으로 분산시킨다.")

heading(d,"17.4.7 SEARCH 시간과 값 복제")
numbered(d,[
 "선택된 원본 세션의 SEARCH를 search_time, search_id 순으로 정렬한다.",
 "원본 세션의 최초 search_time을 search_origin으로 둔다.",
 "각 검색의 새 시간은 synthetic_base_time + (original_search_time − search_origin)으로 계산한다. 이로써 세션 내부 검색 간 시간차와 순서를 보존한다.",
 "검색 내용(query_text, 목적지, 체크인·체크아웃, 결과 수, 정렬, 인원 등)은 원본 세션 값을 복제한다.",
 "search_id와 session_id는 합성 키로 치환하고 data_origin='synthetic_augmentation'으로 설정한다.",
])
code(d,"new_search_time = base_time_i + (old_search_time - min_search_time_of_source_session)")

heading(d,"17.4.8 숙박일 역전 보정")
para(d,"먼저 원본 SEARCH 전체에서 checkout_date−checkin_date가 1일 이상인 정상 숙박기간을 일수 목록 valid_durations로 만든다. 복제 대상 검색에서 checkout_date≤checkin_date이면 고정 RNG로 valid_durations 중 하나를 선택해 checkout_date=checkin_date+선택 일수로 바꾼다.")
code(d,"valid_durations = [(checkout - checkin).days for 모든 원본 검색 if days > 0]\nif copied_checkout <= copied_checkin:\n    copied_checkout = copied_checkin + RNG.choice(valid_durations) days")
para(d,"주의: config에는 ‘같은 원본 세션에서 정상 기간을 뽑는다’고 적혀 있지만 승인 생성 코드의 실제 구현은 원본 전체 정상 숙박기간 풀을 사용한다. 완전 재현 시에는 실제 코드와 동일하게 전체 풀을 사용해야 현재 승인 DB와 일치한다. 향후 생성기 개정 시 config와 코드를 일치시키는 별도 결정이 필요하다.")

heading(d,"17.4.9 SEARCH_FILTER와 SEARCH_RESULT 복제")
bullets(d,[
 "SEARCH_FILTER: old_search_id에 해당하는 단일 행을 찾아 필터 값을 그대로 복제한다. 새 search_filter_id와 new_search_id를 기록하고 old_filter_id→new_filter_id 매핑을 보관한다.",
 "SEARCH_RESULT: old_search_id의 결과를 result_rank 순으로 정렬한다. hotel_id, room_id, result_score, result_rank, price_rank는 유지하고 search_result_id와 search_id만 새 키로 바꾼다.",
 "각 복제 검색의 SEARCH_RESULT 행 수가 해당 SEARCH.total_result_count와 같아야 한다.",
])

heading(d,"17.4.10 EVENT 복제와 제외 순서")
numbered(d,[
 "원본 세션 EVENT를 event_at, event_id 순으로 정렬하고 최초 event_at을 event_origin으로 둔다.",
 "booking_start, booking_complete, booking_cancel 이벤트는 복제하지 않는다.",
 "hotel_click 또는 hotel_detail_view 이벤트의 hotel_id가 새 검색의 SEARCH_RESULT에 없으면 해당 이벤트를 복제하지 않는다.",
 "남은 이벤트의 event_id, session_id, user_id, search_id, search_filter_id를 합성 키로 치환한다. 검색·필터 키가 없는 이벤트는 원본처럼 NULL을 유지한다.",
 "새 event_at은 synthetic_base_time + (old_event_at − event_origin)으로 계산한다. session_end_time이 있으면 같은 event_origin 기준으로 이동한다.",
 "event_type, hotel_id, rating, review 관련 값, device 등 나머지 필드는 유지하고 data_origin='synthetic_augmentation'으로 설정한다.",
])
code(d,"if event_type in booking event types: skip\nif event_type in {hotel_click, hotel_detail_view} and (new_search_id, hotel_id) not in exposed_pairs: skip\nnew_event_at = base_time_i + (old_event_at - min_event_time_of_source_session)")

heading(d,"17.4.11 USER와 BOOKING 처리")
bullets(d,[
 "각 i에 대해 합성 USER 1행을 생성한다. signup_at은 합성 기준시각 30일 전, age_group은 NULL, data_origin은 synthetic_augmentation이다.",
 "생성 코드는 합성 사용자명·이메일 placeholder를 만들지만 실제 개인정보가 아니다. 그래도 보고·분석 산출물에서는 user_name과 email을 내보내지 않는다.",
 "BOOKING은 테이블 스키마만 만들고 어떤 행도 삽입하지 않는다. 원본 BOOKING 36행도 복사하지 않는다.",
])

heading(d,"17.4.12 적재 순서, 메타데이터와 인덱스")
para(d,"참조 무결성을 고려해 HOTEL·ROOM을 먼저 적재하고, 이후 USER, SEARCH, SEARCH_FILTER, SEARCH_RESULT, EVENT를 적재한다. BOOKING은 빈 상태로 둔다. _generation_metadata에는 sample_set_type, scenario_id, random_seed, config_version, generation_version, n_users, n_sessions를 JSON 문자열로 저장한다.")
table(d,["인덱스","컬럼","목적"],[
 ["idx_syn_search_session_time","SEARCH(session_id, search_time, search_id)","세션 내 안정 검색 정렬"],
 ["idx_syn_result_search_hotel","SEARCH_RESULT(search_id, hotel_id)","노출 결과·미노출 클릭 검사"],
 ["idx_syn_event_search_hotel_type","EVENT(search_id, hotel_id, event_type)","검색별 행동 조회"],
 ["idx_syn_event_session_type","EVENT(session_id, event_type)","세션별 행동 조회"],
],7)
para(d,"모든 행과 메타데이터·인덱스 생성이 끝난 뒤 한 번 commit하고 연결을 닫는다. 중간 실패 시 불완전한 DB를 승인 산출물로 사용하지 않는다.")

heading(d,"17.4.13 동일 생성 코드를 이용한 실행 명령")
para(d,"저장소 루트에서 아래 형식으로 실행한다. 출력 폴더는 반드시 새 빈 폴더로 지정한다. 승인 실행 묶음 폴더를 대상으로 다시 실행하거나 기존 SQLite를 덮어쓰면 안 된다.")
code(d,'python "09_단계별 분석/2단계_1000건_증감_분석/02_관측형합성1000명_실행묶음_260903_1606_01/호텔검색_관측형합성1000명_생성코드_260903_1606_01.py" --db "travel_data_filtered_complete_2026-09-03_v02_비식별.sqlite" --config "09_단계별 분석/2단계_1000건_증감_분석/02_관측형합성1000명_실행묶음_260903_1606_01/호텔검색_관측형합성1000명_생성설정_260903_1606_01.json" --output-dir "<새로운_빈_출력폴더>"')
para(d,"코드는 설정의 stage2_bundle_run_id에서 260903_1606을 읽어 호텔검색_관측형합성1000명_데이터_260903_1606_01.sqlite라는 이름을 만든다. 같은 파일이 있으면 FileExistsError로 중단하는 것이 정상 동작이다.")

heading(d,"17.4.14 생성 직후 필수 QA와 승인 기준")
table(d,["검사","합격 기준"],[
 ["SQLite","PRAGMA integrity_check='ok'"],
 ["규모","USER 1,000; 세션 1,000; SEARCH 6,900; SEARCH_FILTER 6,900; SEARCH_RESULT 198,128; EVENT 238,851; BOOKING 0"],
 ["기준정보","HOTEL 1,000; ROOM 3,000"],
 ["검색·필터","SEARCH와 SEARCH_FILTER가 search_id 기준 1:1"],
 ["결과 수","SUM(SEARCH.total_result_count)=SEARCH_RESULT 행 수=198,128"],
 ["중복 결과","(search_id, hotel_id) 중복 0건"],
 ["호텔 연결","SEARCH_RESULT.hotel_id=ROOM.hotel_id 불일치 0건"],
 ["숙박일","checkout_date≤checkin_date 0건"],
 ["0건 행동","0건 검색의 impression/click/detail 0건"],
 ["클릭 노출","검색 결과에 없는 hotel_click 0건"],
 ["행 출처","SEARCH의 data_origin이 모두 synthetic_augmentation"],
 ["핵심 패턴","결과 없음률 3,434/6,900, 0건 후 후속검색 3,271/3,434, 즉시 회복 558/3,271, 세션 최종 회복 488/651, hotel_click 1,052/6,900"],
],6.8)
para(d,"위 값은 seed=20260903과 승인 코드·설정·원본 해시가 모두 같을 때의 재현 기준이다. 값이 다르면 seed를 바꾸거나 결과를 덮어쓰지 말고 원본 해시, 세션 정렬, 키 매핑, 이벤트 제외, 숙박일 보정, 패키지 버전을 순서대로 점검한다.")

heading(d,"17.4.15 구현용 의사코드")
code(d,"open source DB read-only\ncreate empty destination schema\ncopy HOTEL and ROOM exactly\nload original SEARCH, FILTER, RESULT, EVENT\ncreate valid stay-duration pool\nchoose 1,000 source sessions using 23 balanced copies + 11 seeded extras\nshuffle chosen sessions with seed 20260903\nfor i, source_session in chosen_sessions:\n    create synthetic USER and SESSION IDs\n    set deterministic base time\n    for each ordered SEARCH j:\n        map old search/filter IDs to new IDs\n        shift search time; repair reversed stay date if needed\n        copy one FILTER and all ranked RESULTS\n    build exposed (new_search_id, hotel_id) set\n    for each ordered EVENT k:\n        skip booking events\n        skip click/detail if hotel was not exposed\n        map keys and shift event/session-end times\ninsert USER, SEARCH, FILTER, RESULT, EVENT\nleave BOOKING empty\nwrite generation metadata and four indexes\ncommit, close, run all QA, calculate SHA-256")

d.save(docx)
d2=Document(docx);assert len(d2.inline_shapes)==6 and len(d2.tables)>=20 and any(p.text=="17.4 작업자가 그대로 실행할 수 있는 증강 재현 절차" for p in d2.paragraphs)
with zipfile.ZipFile(docx) as z:assert z.testzip() is None and len([n for n in z.namelist() if n.startswith("word/media/")])==6

text=log.read_text(encoding="utf-8")+f"\n## 재현 가능한 증강 절차 상세 보완\n\n- 기존 최종 보고서 한 파일에 17.4 절을 추가했다.\n- 포함 범위: 준비물·해시, 원본 점검, 스키마, 43→1,000 균등 표집 공식, ID 규칙, 시간 변환 공식, 테이블별 복제, 예외 처리, 적재·인덱스, 실행 명령, QA 기준, 의사코드.\n- 합성 DB 재생성: 미실행\n- 갱신 보고서 SHA-256: `{sh(docx)}`\n"
log.write_text(text,encoding="utf-8")
m=json.loads(manifest.read_text(encoding="utf-8"));stage=repo/"09_단계별 분석"/"2단계_1000건_증감_분석"
m["report_docx"]["path"]=str(docx.resolve().relative_to(repo.resolve())).replace("\\","/");m["report_docx"]["size_bytes"]=docx.stat().st_size;m["report_docx"]["sha256"]=sh(docx)
m["packaging_execution_log"]["size_bytes"]=log.stat().st_size;m["packaging_execution_log"]["sha256"]=sh(log)
m["report_content_addendum"]["human_reproducible_augmentation_procedure"]=True;m["report_content_addendum"]["detailed_procedure_section"]="17.4";m["report_content_addendum"]["synthetic_regeneration_performed_during_documentation"]=False;m["report_content_addendum"]["single_final_report_path"]=m["report_docx"]["path"]
manifest.write_text(json.dumps(m,ensure_ascii=False,indent=2),encoding="utf-8")
print(json.dumps({"report":str(docx),"report_sha256":sh(docx),"paragraphs":len(d2.paragraphs),"tables":len(d2.tables),"manifest_sha256":sh(manifest),"synthetic_regeneration":False},ensure_ascii=True,indent=2))
