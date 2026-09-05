from pathlib import Path
import hashlib,json,sys,zipfile
from docx import Document

docx=Path(sys.argv[1]);manifest=Path(sys.argv[2]);log=Path(sys.argv[3]);repo=Path(sys.argv[4])
def sh(p):return hashlib.sha256(Path(p).read_bytes()).hexdigest()
d=Document(docx)
mapping={"17.4 작업자가 그대로 실행할 수 있는 증강 재현 절차":"19. 작업자가 그대로 실행할 수 있는 증강 재현 절차"}
for i in range(1,16):mapping[f"17.4.{i} "]=f"19.{i} "
for p in d.paragraphs:
    if p.text in mapping:p.text=mapping[p.text]
    else:
        for old,new in mapping.items():
            if old.endswith(" ") and p.text.startswith(old):p.text=new+p.text[len(old):];break
d.save(docx)
d=Document(docx)
assert len(d.paragraphs)==181 and len(d.tables)==18 and len(d.inline_shapes)==6
assert any(p.text=="19. 작업자가 그대로 실행할 수 있는 증강 재현 절차" for p in d.paragraphs)
assert all(any(p.text.startswith(f"19.{i} ") for p in d.paragraphs) for i in range(1,16))
with zipfile.ZipFile(docx) as z:
    assert z.testzip() is None
    assert len([n for n in z.namelist() if n.startswith("word/media/")])==6
text=log.read_text(encoding="utf-8")
text += f"\n## 사람이 재현 가능한 증강 방법 최종 보완\n\n- 기존 최종 보고서 한 파일의 19절에 실행 가능한 상세 증강 절차를 추가했다.\n- 문단 181개, 표 18개, 차트 6개, DOCX ZIP/OOXML 재열기 PASS.\n- 합성 데이터 재생성 및 생성기 실행: 미실행\n- 최종 보고서 SHA-256: `{sh(docx)}`\n"
log.write_text(text,encoding="utf-8")
m=json.loads(manifest.read_text(encoding="utf-8"))
m["report_docx"]["path"]=str(docx.resolve().relative_to(repo.resolve())).replace("\\","/")
m["report_docx"]["size_bytes"]=docx.stat().st_size;m["report_docx"]["sha256"]=sh(docx)
m["packaging_execution_log"]["size_bytes"]=log.stat().st_size;m["packaging_execution_log"]["sha256"]=sh(log)
a=m["report_content_addendum"];a["human_reproducible_augmentation_procedure"]=True;a["detailed_procedure_section"]="19";a["synthetic_regeneration_performed_during_documentation"]=False;a["single_final_report_path"]=m["report_docx"]["path"]
manifest.write_text(json.dumps(m,ensure_ascii=False,indent=2),encoding="utf-8")
print(json.dumps({"report_sha256":sh(docx),"log_sha256":sh(log),"manifest_sha256":sh(manifest),"paragraphs":181,"tables":18,"images":6,"generator_executed":False},ensure_ascii=True,indent=2))
