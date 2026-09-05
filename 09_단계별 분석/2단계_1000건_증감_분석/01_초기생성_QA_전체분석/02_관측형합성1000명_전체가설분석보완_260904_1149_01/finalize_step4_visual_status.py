from pathlib import Path
import hashlib, json, sys, zipfile
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

root=Path(sys.argv[1]);docx=Path(sys.argv[2]);log=Path(sys.argv[3]);manifest=Path(sys.argv[4])
def sh(p):return hashlib.sha256(Path(p).read_bytes()).hexdigest()

d=Document(docx)
p=d.add_heading("16. 문서 시각 QA 상태",level=1)
for r in p.runs:
    r.font.name="Malgun Gothic";r._element.rPr.rFonts.set(qn("w:eastAsia"),"맑은 고딕")
p=d.add_paragraph("document_visual_qa=STRUCTURAL_ONLY / MANUAL_VISUAL_REVIEW_REQUIRED. Microsoft Word 자동 렌더러가 HRESULT 0x800A13E9로 문서를 열지 못해 전 페이지 시각 검증을 완료하지 못했다. DOCX 재열기, ZIP·OOXML 관계, 표·문단 및 내장 이미지 6개 검사는 통과했다. 이 제한은 분석 파이프라인의 확장 가능 여부와 분리한다.")
for r in p.runs:
    r.font.name="Malgun Gothic";r._element.rPr.rFonts.set(qn("w:eastAsia"),"맑은 고딕");r.font.size=Pt(9)
d.save(docx)

# Structural validation after the final document mutation.
d=Document(docx)
assert len(d.inline_shapes)==6 and len(d.tables)>=8 and len(d.paragraphs)>40
with zipfile.ZipFile(docx) as z:
    bad=z.testzip();assert bad is None
    names=z.namelist();media=[x for x in names if x.startswith("word/media/")]
    assert len(media)==6 and "word/document.xml" in names and "word/_rels/document.xml.rels" in names

text=log.read_text(encoding="utf-8")
text=text.replace("Word: 재열기, 표/문단, 내장 이미지 6개 구조 검사 PASS; Word PDF 렌더링 후 외부 시각 검수 예정","Word: document_visual_qa=STRUCTURAL_ONLY / MANUAL_VISUAL_REVIEW_REQUIRED; 자동 렌더러 HRESULT 0x800A13E9. DOCX 재열기·ZIP/OOXML·표/문단·내장 이미지 6개 검사 PASS")
text += f"\n## 최종 문서 구조 검증\n\n- DOCX SHA-256: `{sh(docx)}`\n- ZIP test: PASS\n- document.xml 및 관계 파일: PASS\n- 내장 이미지: 6개\n- document_visual_qa: `STRUCTURAL_ONLY`\n- 수동 확인: `MANUAL_VISUAL_REVIEW_REQUIRED`\n"
log.write_text(text,encoding="utf-8")

m=json.loads(manifest.read_text(encoding="utf-8"))
m["document_visual_qa"]="STRUCTURAL_ONLY"
m["manual_visual_review_required"]=True
m["document_render_attempt"]={"renderer":"Microsoft Word COM","status":"FAILED","error":"HRESULT 0x800A13E9","attempts":2,"structural_validation":"PASS","embedded_image_count":6}
m["report_docx"]["size_bytes"]=docx.stat().st_size;m["report_docx"]["sha256"]=sh(docx)
m["packaging_execution_log"]["size_bytes"]=log.stat().st_size;m["packaging_execution_log"]["sha256"]=sh(log)
m["final_status"]="PASS"
manifest.write_text(json.dumps(m,ensure_ascii=False,indent=2),encoding="utf-8")
print(json.dumps({"docx_sha256":sh(docx),"log_sha256":sh(log),"manifest_sha256":sh(manifest),"document_visual_qa":"STRUCTURAL_ONLY","manual_visual_review_required":True},ensure_ascii=True,indent=2))
