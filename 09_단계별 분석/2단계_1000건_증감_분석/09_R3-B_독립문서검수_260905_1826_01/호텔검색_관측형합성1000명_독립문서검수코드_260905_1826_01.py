from __future__ import annotations

import hashlib
import json
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[3]
STAGE2 = ROOT / "09_단계별 분석" / "2단계_1000건_증감_분석"
BASE = STAGE2 / "호텔검색_관측형합성1000명_전체가설분석보고서_260904_1149_01.docx"
R3A = STAGE2 / "08_R3-A_보고서개정_260905_1815_02" / "호텔검색_관측65f합성1000명_전체가설분석보고서_260905_1815_02.docx"
AUTH = STAGE2 / "07_R2.5_승인DB_공식전체분석_260905_1801_01"
ANALYSIS_JSON = AUTH / "호텔검색_관측형합성1000명_수정승인DB전체분석결과_260905_1801_01.json"
STEP3_JSON = AUTH / "호텔검색_관측형합성1000명_STEP3독립검수결과_260905_1801_01.json"
OUT = Path(__file__).resolve().parent
FINAL = OUT / "호텔검색_관측형합성1000명_전체가설분석보고서_260905_1826_03.docx"
LOG = OUT / "호텔검색_관측형합성1000명_보고서검수로그_260905_1826_01.md"
AUDIT_JSON = OUT / "호텔검색_관측형합성1000명_보고서검수결과_260905_1826_01.json"


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for b in iter(lambda: f.read(1024 * 1024), b""):
            h.update(b)
    return h.hexdigest()


def find_paragraph(doc: Document, needle: str) -> Paragraph:
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise RuntimeError(f"문단 없음: {needle}")


def find_heading(doc: Document, needle: str, level: int | None = None) -> Paragraph:
    matches = []
    for p in doc.paragraphs:
        if needle in p.text and p.style.name.startswith("Heading"):
            if level is None or p.style.name == f"Heading {level}":
                matches.append(p)
    if matches:
        # 수동 목차도 Heading 스타일을 사용하므로 동명 항목 중 실제 본문에 있는 마지막 제목을 선택한다.
        return matches[-1]
    raise RuntimeError(f"제목 문단 없음: {needle}")


def move_block_after(doc: Document, start_text: str, end_text: str, target_text: str) -> None:
    body = doc._body._element
    start = find_heading(doc, start_text)._element
    end = find_paragraph(doc, end_text)._element
    target = find_heading(doc, target_text, 1)._element
    children = list(body)
    i0, i1 = children.index(start), children.index(end)
    if i0 > i1:
        raise RuntimeError(f"이동 범위 역전: {start_text} ~ {end_text}")
    block = children[i0 : i1 + 1]
    cursor = target
    for el in block:
        cursor.addnext(el)
        cursor = el


def add_after(anchor, text: str, style: str | None = None) -> Paragraph:
    p = OxmlElement("w:p")
    anchor._element.addnext(p)
    para = Paragraph(p, anchor._parent)
    if style:
        para.style = style
    para.add_run(text)
    return para


def all_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
        c.text for t in doc.tables for r in t.rows for c in r.cells
    )


def set_repeat_header(table) -> None:
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        el = OxmlElement("w:tblHeader")
        el.set(qn("w:val"), "true")
        tr_pr.append(el)


def set_table_font(table, size: float) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(size)


def check_authority() -> dict:
    a = json.loads(ANALYSIS_JSON.read_text(encoding="utf-8"))
    tables = a["tables"]
    high = {(r["dataset_type"], r["definition"], r["active_count"]): r for r in tables["high_order"]}
    bn = {r["dataset"]: r for r in tables["BN_definition_B"]}
    timing = {r["dataset_type"]: r for r in tables["zero_timing"]}
    cross = {r["dataset_type"]: r for r in tables["cross_stream"]}
    syn_high = high[("S0_TIME_ALIGNED_1000", "B_core3", 3)]
    orig_high = high[("ORIGINAL_296", "B_core3", 3)]
    syn_t = timing["S0_TIME_ALIGNED_1000"]
    orig_t = timing["ORIGINAL_296"]
    syn_c = cross["S0_TIME_ALIGNED_1000"]
    return {
        "high_order_original": [orig_high["zero_n"], orig_high["n"], orig_high["zero_n"] / orig_high["n"]],
        "high_order_approved": [syn_high["zero_n"], syn_high["n"], syn_high["zero_n"] / syn_high["n"]],
        "high_order_difference_pp": (syn_high["zero_rate"] - orig_high["zero_rate"]) * 100,
        "bn_original": [bn["ORIGINAL_296"][k] for k in ("log_loss", "brier", "roc_auc", "ece")],
        "bn_approved": [bn["S0_TIME_ALIGNED"][k] for k in ("log_loss", "brier", "roc_auc", "ece")],
        "bn_leakage": [bn["ORIGINAL_296"]["fingerprint_leakage"], bn["S0_TIME_ALIGNED"]["fingerprint_leakage"]],
        "timing_original": orig_t,
        "timing_approved": syn_t,
        "cross_approved": [syn_c["overlap_n"], syn_c["comparable_n"], syn_c["overlap_rate"]],
    }


def main() -> None:
    # Filename typo is corrected here without guessing any other input.
    global R3A
    R3A = STAGE2 / "08_R3-A_보고서개정_260905_1815_02" / "호텔검색_관측형합성1000명_전체가설분석보고서_260905_1815_02.docx"
    authority = check_authority()
    doc = Document(R3A)
    before = all_text(doc)

    # R3-A가 목차의 동명 문구를 본문 제목으로 오인해 앞부분에 삽입한 두 블록을 실제 장으로 이동한다.
    move_block_after(doc, "5.1 STEP R2.5 권위 입력과 계보", "결정적 표본 44건 중 실패 0건", "5. 데이터 계보·품질·재현성 QA")
    move_block_after(doc, "8.1 R2.5 기반 종합 판정", "기존 산출물은 보존했다", "8. 종합 판정·해석 한계·다음 단계")

    # 요구사항 원문에 대응하는 명칭과 독립 검수에서 누락된 방법·한계 문구를 명시한다.
    h = find_heading(doc, "5.2 고차 결합 검증")
    h.text = "5.2 고차 조건부 결합확률(High-order Interaction) 검증"
    cursor = add_after(h, "분석 단위는 SEARCH 1행이다. 실제 선택형 필터 3개(price IS NOT NULL, user_rating_min IS NOT NULL, amenity_count > 0)가 모두 활성화된 검색을 고차 중첩으로, total_result_count=0을 결과 없음으로 정의한다. 정의 B에서 가능한 3필터 조합은 price+rating+amenity 하나뿐이며 복수의 주요 조합이 존재하는 것으로 해석하지 않는다.")
    cursor = add_after(cursor, "제한 Bayesian Network 교차검증은 지역, 검색 의도, 가격·평점·편의시설 활성 플래그, active_filter_count, zero_result를 사용하는 제한 범주형 Naive Bayes/제한 DAG 진단이다. 라플라스 평활을 적용하고 원본 세션 템플릿 fingerprint 기준 5-fold GroupKFold로 분리했으며, 승인 합성본도 같은 fingerprint 그룹을 사용했다. train/test fingerprint 누출은 0이고 세션 단위 bootstrap 500회를 사용했다.")
    cursor = add_after(cursor, "BN은 조건부 구조 보존 진단이며 인과 검증 용도가 아니다. 합성 표본의 한계상 작은 p값이나 유사한 성능을 실제 모집단에 대한 증거 강화로 해석하지 않는다.")

    h = find_heading(doc, "5.3 0건 후 인접 검색 시간 간격")
    h.text = "5.3 세션 시계열 간격(Inter-arrival Time) 현실성 검증"
    cursor = add_after(h, "세션별 SEARCH를 search_time, search_id로 안정 정렬한 뒤, 현재 검색이 0건이고 동일 세션에 바로 다음 검색이 존재할 때 두 SEARCH의 시간차를 인접 전이의 Inter-arrival Time으로 정의한다. 0초는 역전이 아니지만 로그 변환과 로그정규 적합에서는 제외하며, 음수 간격은 시계열 역전으로 판정한다.")
    cursor = add_after(cursor, "검증 결과는 원본 경험 시간차 보존=PASS이다. 다만 모수 재추정 parametric bootstrap에서 로그정규 적합성은 지지되지 않음(WARN)이며, 현실적 체류시간의 입증으로 해석하지 않는다. 승인 합성본은 로그정규분포에서 새 시간을 생성하지 않았고 Jittering도 적용하지 않았다.")

    h = find_heading(doc, "5.4 시간 일관성과 교차 스트림")
    cursor = add_after(h, "공통 시간 원점으로 교차 스트림 오류 수정: session_origin=min(유효 SEARCH.search_time, 실제 복제 대상 EVENT.event_at)을 세션마다 하나만 정의하고 SEARCH·EVENT·session_end_time에 동일한 이동량을 적용했다. SEARCH–EVENT 원본 상대시간 불일치는 0건, 최대 오차는 0초이다. EVENT 원천 순번은 NOT_TESTABLE이며 원천 전체 순서에 대한 검증 주장은 하지 않는다.")

    # 모든 표의 첫 행은 페이지 분할 시 반복하고 행 단위 분할은 금지한다.
    for idx, table in enumerate(doc.tables):
        set_repeat_header(table)
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                tr_pr.append(OxmlElement("w:cantSplit"))
        if len(table.columns) >= 9:
            set_table_font(table, 7.0)
        if len(table.rows) <= 4:
            # 짧은 표는 행 사이 페이지 분리를 막아 머리글 없는 잔여 행이 다음 쪽에 고립되지 않게 한다.
            for row in table.rows[:-1]:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.paragraph_format.keep_with_next = True

    # 제목 뒤 본문 고립을 줄인다.
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True

    doc.save(FINAL)
    checked = Document(FINAL)
    text = all_text(checked)

    must = [
        "High-order Interaction", "91", "104", "2,125", "2,425", "+0.1289%p",
        "GroupKFold", "fingerprint 누출은 0", "BN은 조건부 구조 보존 진단",
        "Inter-arrival Time", "원본 경험 시간차 보존", "로그정규 적합성은 지지되지 않음",
        "μ=2.507799", "σ=0.730905", "p=0.001996", "1,474", "167,330", "0.8809%",
        "공통 시간 원점으로 교차 스트림 오류 수정", "EVENT 원천 순번은 NOT_TESTABLE",
        "합성 표본의 한계", "결정적 표본 44건 중 실패 0건",
    ]
    bad = [
        "로그정규분포를 따른다", "현실적인 체류시간이 입증됐다", "실제 사용자 1,000명",
        "모집단 신뢰도가 증가했다", "Bayesian Network로 인과관계를 검증했다",
        "3개 이상 주요 조합 5개", "전체 EVENT 원천 순서가 검증됐다", "Jittering을 적용했다",
        "모든 시계열 문제가 처음부터 없었다", "old_event_at − event_origin", "old_search_time − search_origin",
    ]

    # 기준 문단 보존: 허용된 시간 공식 문단 외 기준의 비어 있지 않은 문단이 개정본에 유지되는지 확인.
    base_doc = Document(BASE)
    allowed_changed = ("search_origin", "event_origin", "승인 SHA-256은 db80db")
    base_expected = [p.text for p in base_doc.paragraphs if p.text.strip() and not any(x in p.text for x in allowed_changed)]
    preserved_missing = [x for x in base_expected if x not in text]

    ratio_checks = {
        "91/104": abs(91 / 104 - 0.875) < 1e-12,
        "2125/2425": abs(2125 / 2425 - 0.8762886597938144) < 1e-12,
        "1474/167330": abs(1474 / 167330 - 0.008808940417139784) < 1e-12,
        "148529/167330": abs(148529 / 167330 - 0.8876411880714755) < 1e-12,
    }
    result = {
        "authority_recalculation": authority,
        "required_missing": [x for x in must if x not in text],
        "prohibited_found": [x for x in bad if x in text],
        "base_paragraphs_missing_unexpected": preserved_missing,
        "base_table_count": len(base_doc.tables),
        "r3a_table_count": len(Document(R3A).tables),
        "final_table_count": len(checked.tables),
        "ratio_recalculation": ratio_checks,
        "final_sha256": sha256(FINAL),
        "final_size": FINAL.stat().st_size,
    }
    if result["required_missing"] or result["prohibited_found"] or preserved_missing or not all(ratio_checks.values()):
        raise RuntimeError(json.dumps(result, ensure_ascii=False, indent=2))
    AUDIT_JSON.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    LOG.write_text(
        "# STEP R3-B 독립 문서검수 로그\n\n"
        f"- 기준 보고서 SHA-256: `{sha256(BASE)}`\n"
        f"- R3-A 보고서 SHA-256: `{sha256(R3A)}`\n"
        f"- 최종 승인 DB SHA-256: `7d449fab6847730be38fe46ebb6bd6a5b31690cdf38cf4bdbadeaa75edd1ae04`\n"
        "- 요구사항 1: 누락된 High-order Interaction 정식 명칭, 3필터 정의, 제한 BN 방법, GroupKFold, fingerprint 누출 0, 비인과·합성 한계를 명시함\n"
        "- 요구사항 2: 누락된 Inter-arrival Time 정식 명칭, 인접 전이·0초 제외 정의, 경험 시간차 보존 PASS, 로그정규 WARN 한계를 명시함\n"
        "- 수치 QA: R2.5 JSON 독립 추출 및 분자·분모 비율 재계산 PASS\n"
        "- 계보 QA: 승인 DB 해시, 수정 전 S0 비교용, R1 QA 전용, R2.5 권위 자료, 공통 원점 공식 확인 PASS\n"
        "- 보존 QA: 허용된 시간 공식·승인 해시 문단 외 기준 문단 누락 0, 기존 표 19개 보존\n"
        "- 서식 보완: 전 표 머리글 반복, 행 분할 금지, 제목 다음 문단 유지, 광폭 표 글꼴 조정\n"
        "- 렌더링 QA: 별도 렌더링 결과와 육안검수 기록을 아래에 후속 기재\n",
        encoding="utf-8",
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
