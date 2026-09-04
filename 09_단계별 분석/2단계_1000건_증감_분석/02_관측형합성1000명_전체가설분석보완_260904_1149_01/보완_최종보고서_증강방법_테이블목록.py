from pathlib import Path
import hashlib, json, sys, zipfile
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

docx=Path(sys.argv[1]); manifest=Path(sys.argv[2]); log=Path(sys.argv[3])
def sh(p):return hashlib.sha256(Path(p).read_bytes()).hexdigest()
def font(run,size=8,bold=False,color=None):
    run.font.name="Malgun Gothic";run._element.rPr.rFonts.set(qn("w:eastAsia"),"맑은 고딕");run.font.size=Pt(size);run.bold=bold
    if color:run.font.color.rgb=RGBColor(*color)
def shade(cell,color):
    x=OxmlElement("w:shd");x.set(qn("w:fill"),color);cell._tc.get_or_add_tcPr().append(x)
def head(doc,text):
    p=doc.add_heading(text,level=1)
    for r in p.runs:font(r,16,True,(31,78,120))
def para(doc,text):
    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(5);p.paragraph_format.line_spacing=1.1;r=p.add_run(text);font(r,9);return p
def bullets(doc,items):
    for x in items:
        p=doc.add_paragraph(style="List Bullet");p.paragraph_format.space_after=Pt(3);r=p.add_run(x);font(r,8.5)
def table(doc,headers,rows,font_size=7.2):
    t=doc.add_table(rows=1,cols=len(headers));t.style="Table Grid";t.alignment=WD_TABLE_ALIGNMENT.CENTER
    tr=t.rows[0];tr._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for i,h in enumerate(headers):
        tr.cells[i].text=str(h);shade(tr.cells[i],"1F4E78")
        for r in tr.cells[i].paragraphs[0].runs:font(r,font_size,True,(255,255,255))
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v);cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(0)
                for r in p.runs:font(r,font_size)
    return t

d=Document(docx)
d.add_page_break();head(d,"17. S0 관측형 합성 1,000명 증강 방법 상세")
para(d,"이 절은 승인된 생성 설정과 실제 생성 코드의 구현을 기준으로 작성한다. S0 데이터는 원본 행을 단순 합친 통합 DB가 아니라, 원본 43개 검색 세션의 구조를 반복 표집해 만든 별도의 합성 행동 DB다. HOTEL과 ROOM 기준정보는 원본에서 그대로 복사하지만 USER·SEARCH·SEARCH_FILTER·SEARCH_RESULT·EVENT는 새 합성 ID로 구성되며 원본 사용자 89명, 검색 296건, 예약 36건은 합성 DB에 포함되지 않는다.")
table(d,["구분","승인·실제 적용 내용"],[
 ["시나리오","S0 observed_like / config_version S0_PILOT_V02 / generation version 02"],
 ["표집 단위","원본 검색 세션 43개"],
 ["목표 규모","합성 사용자 1,000명, 합성 세션 1,000개"],
 ["표집 설계","OBSERVED_BALANCED_SESSION_BOOTSTRAP_43_V02"],
 ["균등 복제","각 원본 세션을 23회씩 먼저 배치(43×23=989세션)"],
 ["잔여 11세션","43개 원본 세션 중 11개를 비복원 추출해 1회씩 추가"],
 ["순서 무작위화","결합한 1,000개 세션 순서를 NumPy Generator로 섞음"],
 ["난수 고정","random_seed=20260903; QA 통과를 위해 seed를 변경하지 않음"],
 ["세그먼트 할당","세그먼트 목표 비율을 강제하지 않음(segment_quota_assignment=false)"],
 ["결과 0건 종료","immediate_zero_exit_allowed=false; 원본 세션의 검색 흐름을 유지"],
],7.5)

head(d,"17.1 행 생성과 키 변환")
bullets(d,[
 "USER: 합성 세션마다 합성 사용자 1명을 만들고 SYN_U 형식의 새 user_id를 부여한다. 사용자명·이메일은 합성 placeholder이며 보고용 산출물에는 내보내지 않는다.",
 "SESSION: 합성 세션마다 SYN_S 형식의 새 session_id를 부여한다. 사용자와 세션은 1:1이다.",
 "SEARCH: 선택한 원본 세션의 검색 순서를 search_time, search_id로 정렬하고 SYN_Q 형식의 새 search_id를 부여한다. 세션 내부 검색 간 시간차는 보존하고 합성 기준시각으로 평행 이동한다.",
 "SEARCH_FILTER: 원본 검색과 1:1로 복제하며 SYN_F 형식의 search_filter_id와 새 search_id로 연결한다.",
 "SEARCH_RESULT: 각 검색의 result_rank 순서를 유지하고 SYN_R 형식의 새 search_result_id를 부여한다. hotel_id와 room_id는 복사된 기준정보를 참조한다.",
 "EVENT: 원본 세션 내부 event_at·event_id 순서를 유지하고 합성 기준시각으로 이동한다. SYN_E 형식의 event_id와 새 user/session/search/filter 키로 치환한다.",
 "data_origin: 새 행동 테이블 행에는 synthetic_augmentation을 기록한다. HOTEL·ROOM은 원본 기준정보를 그대로 복사한다.",
])

head(d,"17.2 데이터 정제·제외 규칙")
table(d,["대상","처리 규칙","분석 영향"],[
 ["숙박일 역전","checkout_date≤checkin_date인 복제 검색은 원본 전체 정상 숙박기간 풀에서 고정 seed로 기간을 하나 뽑아 checkout을 checkin+기간으로 교체","합성 SEARCH의 역전 0건. 설정 문구의 ‘같은 원본 세션’과 실제 구현이 다르므로 제한사항으로 기록"],
 ["미노출 클릭·상세","새 search_result에 hotel_id가 없는 hotel_click 및 hotel_detail_view 이벤트를 제외","합성 DB 미노출 클릭 0건; hotel_click만 상세진입 KPI로 사용"],
 ["예약 이벤트","booking_start, booking_complete, booking_cancel 이벤트를 제외","예약 행동을 분석하지 않음"],
 ["BOOKING","스키마만 유지하고 행을 생성하지 않음","0행; 예약전환율 산출 금지"],
 ["ROOM 해석","물리 객실 재고가 아니라 객실 상품/타입으로 유지","같은 room_id 기간 중복을 자동 오류로 보지 않음"],
 ["원본 결합 여부","원본 행동 행을 합성 DB에 합치지 않음","합성 결과는 실제 1,000명의 관측 결과가 아니라 모형 내부 진단"],
],7.0)

head(d,"17.3 생성 후 무결성 및 성능 보조 구조")
bullets(d,[
 "SEARCH–SEARCH_FILTER 1:1, SEARCH.total_result_count 합계와 SEARCH_RESULT 행 수 일치, search_id·hotel_id 중복 결과 0건을 검증했다.",
 "SEARCH_RESULT의 hotel_id와 ROOM의 hotel_id 일치, hotel_click 호텔의 검색 결과 포함, 0건 검색의 impression/click/detail 0건을 검증했다.",
 "검색 정렬용 idx_syn_search_session_time(session_id, search_time, search_id), 결과 조회용 idx_syn_result_search_hotel(search_id, hotel_id), 이벤트 조회용 두 인덱스를 추가했다. 인덱스는 데이터 값을 바꾸지 않는다.",
 "생성 DB의 SQLite integrity_check=ok이며 승인 SHA-256은 db80db7048add9c0c4cb1a985e67a77ae99bef3a30ce32bedb70cc0ee61dc896이다.",
])

d.add_page_break();head(d,"18. S0 합성 SQLite 테이블 목록과 데이터 사전")
para(d,"합성 SQLite에는 원본과 같은 8개 업무 테이블과 생성 계보를 위한 _generation_metadata 테이블이 있다. 아래 행 수는 승인된 S0 합성 DB의 실제 값이다.")
summary=[
 ["USER",1000,"합성","합성 사용자. 사용자·세션 1:1","user_id","보고서에 user_name·email 미출력"],
 ["HOTEL",1000,"원본 기준정보 복사","호텔 마스터","hotel_id","행 값 증강 없음"],
 ["ROOM",3000,"원본 기준정보 복사","객실 상품/타입 마스터","room_id","물리 재고로 해석하지 않음"],
 ["SEARCH",6900,"합성","검색 사실 및 결과 수","search_id","session_id로 1,000세션"],
 ["SEARCH_FILTER",6900,"합성","검색별 필터 조건","search_filter_id","SEARCH와 search_id 1:1"],
 ["SEARCH_RESULT",198128,"합성","검색 노출 결과","search_result_id","합계=SEARCH 결과 수 합계"],
 ["EVENT",238851,"합성","검색·노출·클릭 행동","event_id","예약 이벤트와 미노출 클릭 제외"],
 ["BOOKING",0,"빈 스키마","예약 정보","booking_id","예약전환 분석 제외"],
 ["_generation_metadata",7,"생성 메타데이터","시나리오·seed·버전·규모","key","업무 KPI에 사용하지 않음"],
]
table(d,["테이블","행 수","데이터 성격","역할","PK","핵심 처리"],summary,6.8)

head(d,"18.1 테이블별 정확한 컬럼 목록")
dictionary=[
 ["USER","user_id, user_name, age_group, email, signup_at, data_origin"],
 ["HOTEL","hotel_id, hotel_name, city_name, grade, hotel_address, user_rating, review_count, property_type, actual_address, actual_city, actual_prefecture, actual_postal_code, actual_latitude, actual_longitude, actual_phone, actual_star_rating, supplier_hotel_code, rtx_code, agoda_code, expedia_code, top_selling_rank, source_last_mapped_at, actual_data_sources, data_origin"],
 ["ROOM","room_id, hotel_id, guest_count, room_count, room_options, pay_later_flag, free_cancel_flag, room_price, room_type, data_origin"],
 ["SEARCH","search_id, session_id, search_time, query_text, checkin_date, checkout_date, total_result_count, sort_option, guest_count, destination, data_origin"],
 ["SEARCH_FILTER","search_filter_id, search_id, property_type, property_grade, user_rating_min, price, amenity_count, region, data_origin"],
 ["SEARCH_RESULT","search_result_id, search_id, hotel_id, room_id, result_score, result_rank, price_rank, data_origin"],
 ["EVENT","event_id, session_id, event_type, event_at, hotel_id, search_filter_id, search_id, user_id, rating, session_end_time, review_completed_at, review_text, device, data_origin"],
 ["BOOKING","booking_id, user_id, hotel_id, room_id, booking_status, booking_amount, booking_at, checkin_date, checkout_date, guest_count, room_count, cancellation_deadline, data_origin"],
 ["_generation_metadata","key, value"],
]
table(d,["테이블","컬럼(실제 SQLite 컬럼명)"],dictionary,6.6)

head(d,"18.2 주요 관계와 분석 사용 범위")
table(d,["관계","카디널리티·키","사용 범위"],[
 ["USER → SEARCH/EVENT","user_id 및 합성 세션 연결","개인정보 컬럼은 분석 산출물에서 제외"],
 ["SEARCH → SEARCH_FILTER","search_id 기준 1:1","A1·A2 필터 및 의도 분류"],
 ["SEARCH → SEARCH_RESULT","search_id 기준 1:N","결과 수·노출 무결성"],
 ["SEARCH → EVENT","search_id 기준 1:N","hotel_click 존재 여부와 행동 플래그"],
 ["HOTEL → ROOM","hotel_id 기준 1:N","호텔·객실 상품 기준정보"],
 ["SEARCH_RESULT → HOTEL/ROOM","hotel_id, room_id","노출 결과의 호텔·객실 연결"],
 ["BOOKING → USER/HOTEL/ROOM","스키마상 키만 존재; 현재 0행","성과 KPI와 가설검정에서 제외"],
],7.0)

head(d,"18.3 증강 데이터 해석 시 주의")
bullets(d,[
 "이 DB는 ‘원본 296검색 + 추가 합성 검색’의 합본이 아니다. 행동 데이터는 합성 행만 들어 있으며 HOTEL·ROOM 기준정보만 원본에서 복사됐다.",
 "원본의 사용자 89명, 검색 296건, BOOKING 36건은 별도 원본 DB에 남아 있다.",
 "S0 1,000명이라는 표현은 합성 사용자 1,000개를 뜻한다. 실제 모집된 사용자 1,000명의 관측 결과가 아니다.",
 "실제 가설 판단은 원본 296검색을 기준으로 하고, 합성 분석은 생성기와 분석 파이프라인의 구조·방향 보존을 확인하는 용도다.",
])

d.save(docx)
# final structure check
d2=Document(docx); assert len(d2.inline_shapes)==6 and len(d2.tables)>=13 and len(d2.paragraphs)>65
with zipfile.ZipFile(docx) as z: assert z.testzip() is None and len([x for x in z.namelist() if x.startswith("word/media/")])==6

text=log.read_text(encoding="utf-8")
text += f"\n## 최종 보고서 단일 파일 보완\n\n- 기존 최종 보고서 파일을 새 파일 생성 없이 직접 보완했다.\n- 추가 내용: S0 증강 방법, 균등 세션 부트스트랩 43→1,000, seed·ID·시간 이동·정제 규칙, 테이블 9개 행 수·역할·PK·정확한 컬럼 목록·관계·원본 포함 여부.\n- config의 역전 숙박일 문구와 실제 구현 차이를 제한사항으로 공개했다.\n- 최종 보고서 SHA-256: `{sh(docx)}`\n"
log.write_text(text,encoding="utf-8")
m=json.loads(manifest.read_text(encoding="utf-8"))
m["report_docx"]["size_bytes"]=docx.stat().st_size;m["report_docx"]["sha256"]=sh(docx)
m["packaging_execution_log"]["size_bytes"]=log.stat().st_size;m["packaging_execution_log"]["sha256"]=sh(log)
m["report_content_addendum"]={"augmentation_method_detailed":True,"table_inventory_detailed":True,"single_final_report_path":m["report_docx"]["path"],"synthetic_database_contains_original_behavior_rows":False,"reference_tables_copied_from_original":["hotel","room"],"behavior_tables_synthetic_only":["user","search","search_filter","search_result","event"],"booking_rows":0,"config_code_difference_disclosed":"invalid stay dates: config says same source session; implementation samples from global valid-duration pool"}
m["known_limitations"].append("숙박일 역전 보정은 config 문구와 달리 실제 코드에서 원본 전체 정상 숙박기간 풀을 사용함")
manifest.write_text(json.dumps(m,ensure_ascii=False,indent=2),encoding="utf-8")
print(json.dumps({"report":str(docx),"report_sha256":sh(docx),"log_sha256":sh(log),"manifest_sha256":sh(manifest),"final_report_count_created":1},ensure_ascii=True,indent=2))
