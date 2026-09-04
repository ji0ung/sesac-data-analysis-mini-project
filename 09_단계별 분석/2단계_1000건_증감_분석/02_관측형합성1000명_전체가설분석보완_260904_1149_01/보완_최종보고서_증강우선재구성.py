from pathlib import Path
import copy,hashlib,json,sys,zipfile
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

docx=Path(sys.argv[1]);manifest=Path(sys.argv[2]);log=Path(sys.argv[3]);repo=Path(sys.argv[4])
def sh(p):return hashlib.sha256(Path(p).read_bytes()).hexdigest()
def ptext(el):return "".join(x.text or "" for x in el.xpath(".//w:t")) if el.tag==qn("w:p") else ""
def style_id(el):
    if el.tag!=qn("w:p"):return None
    pPr=el.find(qn("w:pPr"));ps=pPr.find(qn("w:pStyle")) if pPr is not None else None
    return ps.get(qn("w:val")) if ps is not None else None
def set_style(el,style):
    pPr=el.find(qn("w:pPr"))
    if pPr is None:pPr=OxmlElement("w:pPr");el.insert(0,pPr)
    ps=pPr.find(qn("w:pStyle"))
    if ps is None:ps=OxmlElement("w:pStyle");pPr.insert(0,ps)
    ps.set(qn("w:val"),style)
def set_text(el,text):
    # Preserve paragraph properties and replace runs with a single formatted run.
    for child in list(el):
        if child.tag!=qn("w:pPr"):el.remove(child)
    r=OxmlElement("w:r");rPr=OxmlElement("w:rPr");fonts=OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"),"Malgun Gothic");fonts.set(qn("w:hAnsi"),"Malgun Gothic");fonts.set(qn("w:eastAsia"),"맑은 고딕");rPr.append(fonts);r.append(rPr);t=OxmlElement("w:t");t.text=text;r.append(t);el.append(r)
def make_toc_line(text,indent=False):
    p=OxmlElement("w:p");pPr=OxmlElement("w:pPr");spacing=OxmlElement("w:spacing");spacing.set(qn("w:after"),"55");pPr.append(spacing)
    if indent:
        ind=OxmlElement("w:ind");ind.set(qn("w:left"),"360");pPr.append(ind)
    p.append(pPr);r=OxmlElement("w:r");rPr=OxmlElement("w:rPr");fonts=OxmlElement("w:rFonts");fonts.set(qn("w:ascii"),"Malgun Gothic");fonts.set(qn("w:hAnsi"),"Malgun Gothic");fonts.set(qn("w:eastAsia"),"맑은 고딕");rPr.append(fonts);sz=OxmlElement("w:sz");sz.set(qn("w:val"),"19");rPr.append(sz);r.append(rPr);t=OxmlElement("w:t");t.text=text;r.append(t);p.append(r);return p
def pagebreak():
    p=OxmlElement("w:p");r=OxmlElement("w:r");b=OxmlElement("w:br");b.set(qn("w:type"),"page");r.append(b);p.append(r);return p
def insert_after_paragraph(paragraph,text,shade="E2F0D9"):
    p=OxmlElement("w:p");pPr=OxmlElement("w:pPr");shd=OxmlElement("w:shd");shd.set(qn("w:fill"),shade);pPr.append(shd);spacing=OxmlElement("w:spacing");spacing.set(qn("w:before"),"70");spacing.set(qn("w:after"),"120");pPr.append(spacing);p.append(pPr)
    r=OxmlElement("w:r");rPr=OxmlElement("w:rPr");fonts=OxmlElement("w:rFonts");fonts.set(qn("w:ascii"),"Malgun Gothic");fonts.set(qn("w:hAnsi"),"Malgun Gothic");fonts.set(qn("w:eastAsia"),"맑은 고딕");rPr.append(fonts);color=OxmlElement("w:color");color.set(qn("w:val"),"375623");rPr.append(color);sz=OxmlElement("w:sz");sz.set(qn("w:val"),"17");rPr.append(sz);r.append(rPr);t=OxmlElement("w:t");t.text=text;r.append(t);p.append(r);paragraph._p.addnext(p);return p

d=Document(docx);body=d._element.body;children=list(body);sect=children[-1] if children[-1].tag==qn("w:sectPr") else None
# Build ranges for current Heading 1 blocks.
h1=[]
for i,el in enumerate(children):
    if style_id(el)=="1":h1.append((i,ptext(el)))
names=[x[1] for x in h1]
required=["목차","1. 한 페이지 핵심 요약","2. 분석 목적과 범위","3. 입력 파일·데이터 계보·재현성","4. 지표 사전과 분석 단위","5. 데이터 품질과 분석 QA","6. A1 결과","7. A2 결과","8. B1 결과","9. B2 결과","10. B3 결과","11. 세션 결과 4개 세그먼트","12. H3 결과","13. 원본과 S0 합성 1,000명 비교","14. 통계적·데이터적 한계","15. 10,000명 확장 진입 판정과 다음 단계","16. 문서 시각 QA 상태","17. S0 관측형 합성 1,000명 증강 방법 상세","18. S0 합성 SQLite 테이블 목록과 데이터 사전","19. 작업자가 그대로 실행할 수 있는 증강 재현 절차"]
assert all(x in names for x in required)
blocks={}
for j,(start,name) in enumerate(h1):
    end=h1[j+1][0] if j+1<len(h1) else (len(children)-1 if sect is not None else len(children))
    blocks[name]=children[start:end]
toc_start=next(i for i,x in h1 if x=="목차")
prefix=children[:toc_start]

# Rename primary chapters and convert merged chapters to subsections.
rename={
"2. 분석 목적과 범위":("1. 증강 목적과 데이터 범위","1"),
"17. S0 관측형 합성 1,000명 증강 방법 상세":("2. S0 1,000명 증강 방법 한눈에 보기","1"),
"19. 작업자가 그대로 실행할 수 있는 증강 재현 절차":("3. 작업자가 그대로 실행할 수 있는 단계별 증강 절차","1"),
"18. S0 합성 SQLite 테이블 목록과 데이터 사전":("4. 합성 SQLite 테이블 구조와 데이터 사전","1"),
"3. 입력 파일·데이터 계보·재현성":("5. 데이터 계보·품질·재현성 QA","1"),
"4. 지표 사전과 분석 단위":("5.1 지표 사전과 분석 단위","2"),
"5. 데이터 품질과 분석 QA":("5.2 계산 및 독립 검수 결과","2"),
"16. 문서 시각 QA 상태":("5.3 문서 QA 상태","2"),
"1. 한 페이지 핵심 요약":("6. S0 1,000명 분석 핵심 요약","1"),
"13. 원본과 S0 합성 1,000명 비교":("6.1 원본과 S0 합성 1,000명 비교","2"),
"6. A1 결과":("7. A1·A2: 필터·지역·검색의도 결과","1"),
"7. A2 결과":("7.1 A2 지역×검색의도 결과","2"),
"8. B1 결과":("8. B1·B2: 후속검색과 탐색지속성 결과","1"),
"9. B2 결과":("8.1 B2 세션 검색 횟수 결과","2"),
"10. B3 결과":("9. B3: 서로 다른 세 가지 회복 지표","1"),
"11. 세션 결과 4개 세그먼트":("10. 세션 결과 4개 세그먼트","1"),
"12. H3 결과":("11. H3: 재검색 전이 유형 결과","1"),
"14. 통계적·데이터적 한계":("12. 해석 한계와 다음 단계","1"),
"15. 10,000명 확장 진입 판정과 다음 단계":("12.1 10,000명 확장 진입 판정","2"),
}
for old,(new,style) in rename.items():
    set_text(blocks[old][0],new);set_style(blocks[old][0],style)
# Renumber detailed method/data-dictionary subsections for intuitive local hierarchy.
for old,prefix_old,prefix_new in [("17. S0 관측형 합성 1,000명 증강 방법 상세","17.","2."),("19. 작업자가 그대로 실행할 수 있는 증강 재현 절차","19.","3."),("18. S0 합성 SQLite 테이블 목록과 데이터 사전","18.","4.")]:
    for el in blocks[old][1:]:
        txt=ptext(el)
        if txt.startswith(prefix_old):set_text(el,prefix_new+txt[len(prefix_old):]);set_style(el,"2")

# The chapter order now leads with generation, then structure/QA, then results.
order=[
"2. 분석 목적과 범위",
"17. S0 관측형 합성 1,000명 증강 방법 상세",
"19. 작업자가 그대로 실행할 수 있는 증강 재현 절차",
"18. S0 합성 SQLite 테이블 목록과 데이터 사전",
"3. 입력 파일·데이터 계보·재현성","4. 지표 사전과 분석 단위","5. 데이터 품질과 분석 QA","16. 문서 시각 QA 상태",
"1. 한 페이지 핵심 요약","13. 원본과 S0 합성 1,000명 비교",
"6. A1 결과","7. A2 결과","8. B1 결과","9. B2 결과","10. B3 결과","11. 세션 결과 4개 세그먼트","12. H3 결과",
"14. 통계적·데이터적 한계","15. 10,000명 확장 진입 판정과 다음 단계"]

# Rebuild the body while preserving cover prefix and section properties.
for el in list(body):body.remove(el)
for el in prefix:body.append(el)
toc=blocks["목차"][0];body.append(toc)
toc_lines=[
"1. 증강 목적과 데이터 범위","2. S0 1,000명 증강 방법 한눈에 보기","3. 작업자가 그대로 실행할 수 있는 단계별 증강 절차","4. 합성 SQLite 테이블 구조와 데이터 사전","5. 데이터 계보·품질·재현성 QA","6. S0 1,000명 분석 핵심 요약","7. A1·A2: 필터·지역·검색의도 결과","8. B1·B2: 후속검색과 탐색지속성 결과","9. B3: 서로 다른 세 가지 회복 지표","10. 세션 결과 4개 세그먼트","11. H3: 재검색 전이 유형 결과","12. 해석 한계와 다음 단계"]
for line in toc_lines:body.append(make_toc_line(line))
body.append(pagebreak())
for name in order:
    for el in blocks[name]:body.append(el)
if sect is not None:body.append(sect)

# Replace obsolete reader-route wording with the new order.
for p in d.paragraphs:
    if p.text.startswith("추천 읽기 순서 |"):
        set_text(p._p,"추천 읽기 순서 | 먼저 1~5절에서 증강 설계·재현 방법·테이블·QA를 확인한 뒤, 6~11절에서 1,000명 분석 결과를 읽는다. 마지막 12절에서 해석 한계와 다음 단계를 확인한다.")

# Add a plain-language explanation immediately after every figure caption.
graph_notes={
"그림 1. A1 필터별 결과 없음률":"그래프 해설 | 파란 막대는 실제 관측 원본, 주황 막대는 S0 합성 모형이다. 편의시설 3개 이상·최소평점 설정·가격 설정 집단에서 결과 없음률이 비교집단보다 높다. 두 데이터의 막대 높이가 비슷한 것은 방향 재현을 뜻하며, 합성 결과가 새로운 실제 효과를 입증하는 것은 아니다.",
"그림 2. A2 지역×검색의도 결과 없음률":"그래프 해설 | 색이 진할수록 결과 없음률이 높고 각 셀의 n이 표본 크기다. 먼저 n을 확인한 뒤 비율을 읽어야 한다. N/A는 해당 조합이 관측되지 않았다는 뜻이며 0%나 불가능을 뜻하지 않는다. 원본의 작은 n 셀은 기술통계로만 본다.",
"그림 3. B1 후속검색률":"그래프 해설 | 0건 검색 뒤에는 원본 140/147, 합성 3,271/3,434에서 바로 다음 검색이 있었다. 비0건 검색보다 후속검색률이 높지만 이는 관련성이다. 동일 세션 반복 검색의 상관과 사용자의 탐색 의지가 함께 영향을 줄 수 있다.",
"그림 4. B3 회복 지표와 서로 다른 분모":"그래프 해설 | 세 막대 묶음은 서로 다른 질문이다. 즉시 회복은 0건→다음 검색 전이, 세션 최종 회복은 0건 경험 세션, 첫 검색 0건 회복은 별도 진단 세션이 분모다. 따라서 17.1%, 75.0%, 66.7%를 한 순위로 직접 비교하면 안 된다.",
"그림 5. 상호배타 세션 결과 세그먼트":"그래프 해설 | 모든 세션은 직접 성공·결과 노출 미선택·재검색 회복·지속 실패 중 하나에만 속한다. 원본 43세션과 합성 1,000세션의 구성비가 유사해 세션 구조가 보존됐음을 보여준다. 합성 세션은 실제 모집 사용자가 아니다.",
"그림 6. H3 재검색 유형별 성공률과 hotel_click률":"그래프 해설 | 왼쪽은 다음 검색이 비0건이 된 비율, 오른쪽은 다음 검색에 hotel_click이 있었던 비율이다. 유형별 막대 위 분자/분모를 함께 봐야 한다. 원본 지역 변경의 회복률이 높지만 작은 유형 표본과 탐색적 분류이므로 제품 효과로 단정하지 않는다."
}
for p in list(d.paragraphs):
    if p.text in graph_notes:insert_after_paragraph(p,graph_notes[p.text],shade="E2F0D9")
d.save(docx)

# Structural/order verification.
d2=Document(docx);top=[p.text for p in d2.paragraphs if p.style.name=="Heading 1"]
expected=["목차",*toc_lines]
assert top==expected,(top,expected)
assert len(d2.tables)==18 and len(d2.inline_shapes)==6 and sum(p.text.startswith("표 해설 |") for p in d2.paragraphs)==18 and sum(p.text.startswith("그래프 해설 |") for p in d2.paragraphs)==6
all_text="\n".join(p.text for p in d2.paragraphs)
assert all(x in all_text for x in ["각 원본 세션을 23회씩","random_seed=20260903","SEARCH_RESULT 198,128","3.15 구현용 의사코드"])
with zipfile.ZipFile(docx) as z:assert z.testzip() is None and len([x for x in z.namelist() if x.startswith("word/media/")])==6

text=log.read_text(encoding="utf-8")+f"\n## 증강 방법 우선 문서 재구성\n\n- 목차를 12개 핵심 장으로 축소했다.\n- 문서 순서를 증강 목적·방법·재현 절차·테이블·QA → 1,000명 분석 결과 → 한계·다음 단계로 변경했다.\n- 하위 절은 본문 탐색에 유지하되 목차에서는 생략해 복잡도를 낮췄다.\n- 표 18개 해설과 그래프 6개 해설을 각각 해당 표·그림 바로 아래에 배치했다.\n- 기존 최종 보고서 한 파일만 수정했으며 합성 데이터·생성기는 실행하지 않았다.\n- 최종 보고서 SHA-256: `{sh(docx)}`\n"
log.write_text(text,encoding="utf-8")
m=json.loads(manifest.read_text(encoding="utf-8"));m["report_docx"]["path"]=str(docx.resolve().relative_to(repo.resolve())).replace("\\","/");m["report_docx"]["size_bytes"]=docx.stat().st_size;m["report_docx"]["sha256"]=sh(docx);m["packaging_execution_log"]["size_bytes"]=log.stat().st_size;m["packaging_execution_log"]["sha256"]=sh(log)
a=m["report_content_addendum"];a["top_level_toc_count"]=12;a["content_order"]="augmentation_method_first_then_S0_1000_analysis_results";a["table_explanation_count"]=18;a["chart_explanation_count"]=6;a["single_final_report_path"]=m["report_docx"]["path"]
manifest.write_text(json.dumps(m,ensure_ascii=False,indent=2),encoding="utf-8")
print(json.dumps({"report_sha256":sh(docx),"manifest_sha256":sh(manifest),"top_level_toc_count":12,"tables":18,"table_explanations":18,"images":6,"chart_explanations":6,"new_report_created":False,"generator_executed":False},ensure_ascii=True,indent=2))
