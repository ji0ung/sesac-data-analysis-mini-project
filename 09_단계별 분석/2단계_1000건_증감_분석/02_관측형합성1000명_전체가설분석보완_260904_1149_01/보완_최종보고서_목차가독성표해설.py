from pathlib import Path
import hashlib,json,sys,zipfile
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor

docx=Path(sys.argv[1]);manifest=Path(sys.argv[2]);log=Path(sys.argv[3]);repo=Path(sys.argv[4])
def sh(p):return hashlib.sha256(Path(p).read_bytes()).hexdigest()
def fmt(run,size=9,bold=False,color=None):
    run.font.name="Malgun Gothic";run._element.rPr.rFonts.set(qn("w:eastAsia"),"맑은 고딕");run.font.size=Pt(size);run.bold=bold
    if color:run.font.color.rgb=RGBColor(*color)
def insert_after_paragraph(paragraph,text,style=None,shade=None):
    p=OxmlElement("w:p");paragraph._p.addnext(p)
    if style:
        pPr=OxmlElement("w:pPr");pStyle=OxmlElement("w:pStyle");pStyle.set(qn("w:val"),style);pPr.append(pStyle);p.append(pPr)
    if shade:
        pPr=p.find(qn("w:pPr"))
        if pPr is None:pPr=OxmlElement("w:pPr");p.insert(0,pPr)
        s=OxmlElement("w:shd");s.set(qn("w:fill"),shade);pPr.append(s)
        sp=OxmlElement("w:spacing");sp.set(qn("w:before"),"80");sp.set(qn("w:after"),"80");pPr.append(sp)
    r=OxmlElement("w:r");rPr=OxmlElement("w:rPr");fonts=OxmlElement("w:rFonts");fonts.set(qn("w:ascii"),"Malgun Gothic");fonts.set(qn("w:hAnsi"),"Malgun Gothic");fonts.set(qn("w:eastAsia"),"맑은 고딕");rPr.append(fonts);sz=OxmlElement("w:sz");sz.set(qn("w:val"),"18");rPr.append(sz);r.append(rPr);t=OxmlElement("w:t");t.text=text;r.append(t);p.append(r)
    return p
def insert_after_table(table,text):
    p=OxmlElement("w:p");pPr=OxmlElement("w:pPr");shd=OxmlElement("w:shd");shd.set(qn("w:fill"),"EAF2F8");pPr.append(shd);spacing=OxmlElement("w:spacing");spacing.set(qn("w:before"),"70");spacing.set(qn("w:after"),"120");pPr.append(spacing);p.append(pPr)
    r=OxmlElement("w:r");rPr=OxmlElement("w:rPr");fonts=OxmlElement("w:rFonts");fonts.set(qn("w:ascii"),"Malgun Gothic");fonts.set(qn("w:hAnsi"),"Malgun Gothic");fonts.set(qn("w:eastAsia"),"맑은 고딕");rPr.append(fonts);bold=OxmlElement("w:b");rPr.append(bold);color=OxmlElement("w:color");color.set(qn("w:val"),"1F4E78");rPr.append(color);sz=OxmlElement("w:sz");sz.set(qn("w:val"),"17");rPr.append(sz);r.append(rPr);t=OxmlElement("w:t");t.text="표 해설 | "+text;r.append(t);p.append(r);table._tbl.addnext(p)

d=Document(docx)
if any("목차와 표 해설 가독성 보완 완료" in p.text for p in d.paragraphs):raise RuntimeError("Readability amendment already applied")

# Correct heading hierarchy for method/data-dictionary subsections.
for p in d.paragraphs:
    if any(p.text.startswith(x) for x in ["17.1 ","17.2 ","17.3 ","18.1 ","18.2 ","18.3 "]):p.style=d.styles["Heading 2"]

# Replace the old static TOC block between the TOC heading and chapter 1.
paras=d.paragraphs;toc=next(p for p in paras if p.text=="목차");first=next(p for p in paras if p.text=="1. 한 페이지 핵심 요약")
node=toc._p.getnext()
while node is not None and node is not first._p:
    nxt=node.getnext();node.getparent().remove(node);node=nxt
toc_lines=[
"1. 한 페이지 핵심 요약","2. 분석 목적과 범위","3. 입력 파일·데이터 계보·재현성","4. 지표 사전과 분석 단위","5. 데이터 품질과 분석 QA",
"6. A1 결과 — 필터 조건과 결과 없음","7. A2 결과 — 지역×검색의도","8. B1 결과 — 후속검색 관련성","9. B2 결과 — 세션 검색 횟수","10. B3 결과 — 세 가지 회복 지표",
"11. 세션 결과 4개 세그먼트","12. H3 결과 — 재검색 전이 유형","13. 원본과 S0 합성 1,000명 비교","14. 통계적·데이터적 한계","15. 10,000명 확장 진입 판정과 다음 단계","16. 문서 시각 QA 상태",
"17. S0 관측형 합성 1,000명 증강 방법 상세","    17.1 행 생성과 키 변환","    17.2 데이터 정제·제외 규칙","    17.3 생성 후 무결성 및 성능 보조 구조",
"18. S0 합성 SQLite 테이블 목록과 데이터 사전","    18.1 테이블별 정확한 컬럼 목록","    18.2 주요 관계와 분석 사용 범위","    18.3 증강 데이터 해석 시 주의",
"19. 작업자가 그대로 실행할 수 있는 증강 재현 절차","    19.1 준비물과 고정값","    19.2 실행 전 원본 점검","    19.3 새 SQLite 스키마 생성","    19.4 HOTEL·ROOM 기준정보 복사","    19.5 원본 43세션을 1,000세션으로 균등 확장","    19.6 합성 ID와 기준시각 생성","    19.7 SEARCH 시간과 값 복제","    19.8 숙박일 역전 보정","    19.9 SEARCH_FILTER·SEARCH_RESULT 복제","    19.10 EVENT 복제와 제외 순서","    19.11 USER·BOOKING 처리","    19.12 적재 순서·메타데이터·인덱스","    19.13 실행 명령","    19.14 생성 직후 필수 QA","    19.15 구현용 의사코드"
]
anchor=toc._p
for line in toc_lines:
    p=OxmlElement("w:p");pPr=OxmlElement("w:pPr");ind=OxmlElement("w:ind");ind.set(qn("w:left"),"360" if line.startswith("    ") else "0");pPr.append(ind);spacing=OxmlElement("w:spacing");spacing.set(qn("w:after"),"35");pPr.append(spacing);p.append(pPr)
    r=OxmlElement("w:r");rPr=OxmlElement("w:rPr");fonts=OxmlElement("w:rFonts");fonts.set(qn("w:ascii"),"Malgun Gothic");fonts.set(qn("w:hAnsi"),"Malgun Gothic");fonts.set(qn("w:eastAsia"),"맑은 고딕");rPr.append(fonts);sz=OxmlElement("w:sz");sz.set(qn("w:val"),"17" if line.startswith("    ") else "19");rPr.append(sz);r.append(rPr);t=OxmlElement("w:t");t.set(qn("xml:space"),"preserve");t.text=line.strip();r.append(t);p.append(r);anchor.addnext(p);anchor=p

# Add a reader-oriented route immediately after chapter 1 heading.
chapter1=next(p for p in d.paragraphs if p.text=="1. 한 페이지 핵심 요약")
insert_after_paragraph(chapter1,"추천 읽기 순서 | 의사결정자는 1→6~15절을, 데이터 작업자는 17→18→19절을 먼저 읽으면 된다. 파란색은 실제 관측 원본, 주황색은 S0 합성 모형 결과다.",shade="FFF2CC")

explanations=[
"같은 지표를 실제 관측 원본과 S0 합성 모형으로 나란히 제시한다. 두 값이 가깝다는 것은 생성 구조가 핵심 비율을 재현했다는 뜻이지, 합성 데이터가 새로운 실증 근거라는 뜻은 아니다.",
"STEP1은 입력·해시, STEP2는 계산, STEP3는 독립 재검산을 담당한다. 세 단계가 모두 PASS여야 이 보고서의 수치를 사용할 수 있다.",
"제한 또는 설정 집단의 결과 없음 오즈를 비교집단과 비교한다. 원본에서 OR이 모두 1보다 크지만, 입력 상태 컬럼 부재 때문에 A1은 부분 채택으로 제한한다.",
"0건 검색은 비0건 검색보다 바로 다음 검색으로 이어진 비율이 높다. 이는 재검색과의 관련성이지 0건이 재검색을 일으켰다는 인과 증거는 아니다.",
"0건 경험 세션의 검색 횟수가 더 많지만, 오래 탐색한 세션일수록 0건을 경험할 기회도 늘어난다. 따라서 U검정과 효과크기는 탐색지속성 진단으로 읽는다.",
"즉시 회복, 세션 최종 회복, 첫 검색 0건 회복은 분모가 서로 다르다. 특히 21/28과 4/6을 같은 지표로 혼동하지 않아야 한다.",
"네 세그먼트는 모든 세션을 정확히 한 번만 분류한다. 원본 43세션과 합성 1,000세션 모두 비율 구조가 유사하지만 합성값은 모형 내부 결과다.",
"유형별 n을 먼저 보고 성공률과 click률을 해석한다. 원본의 혼합 변경처럼 n이 매우 작은 셀은 순위나 효과를 일반화하지 않는다.",
"두 승인 상태가 true라는 것은 다음 확장 프롬프트를 준비할 수 있다는 뜻이다. 이번 보고서가 실제 10,000명 효과를 입증했다는 의미는 아니다.",
"핵심은 ‘43세션을 모두 23회 사용하고 11세션만 한 번 더 사용’하는 균등 부트스트랩이다. seed를 고정해 같은 선택과 순서를 재현한다.",
"원본 품질 문제는 원본에서 고치지 않고 합성 복제 단계에서만 규칙적으로 처리한다. BOOKING과 예약 이벤트는 S0 분석 범위에서 제외한다.",
"행 수 열은 실제 합성 DB 규모다. HOTEL·ROOM만 원본 기준정보이고 USER·SEARCH·FILTER·RESULT·EVENT는 합성 행이며 BOOKING은 빈 스키마다.",
"컬럼명은 실제 SQLite 이름 그대로다. 구현할 때 임의 별칭을 쓰더라도 최종 적재 시 이 이름과 데이터형을 유지해야 기존 분석 코드가 작동한다.",
"SEARCH–FILTER는 1:1, SEARCH–RESULT와 SEARCH–EVENT는 1:N이다. 조인 시 SEARCH_RESULT와 EVENT를 동시에 붙이면 행이 곱해질 수 있으므로 각각 집계한 뒤 결합한다.",
"재현 전에 파일 해시와 seed를 먼저 고정한다. 하나라도 다르면 같은 알고리즘이어도 승인 DB와 바이트·행 결과가 달라질 수 있다.",
"합성 ID는 원본 ID와 충돌하지 않도록 객체별 접두어와 세션·순번을 결합한다. old→new 매핑을 먼저 만든 뒤 자식 테이블의 FK를 치환한다.",
"네 인덱스는 세션 정렬과 노출·행동 조회를 빠르게 할 뿐 데이터 값이나 분석 정의를 바꾸지 않는다.",
"이 표의 모든 조건을 통과해야 생성물을 승인할 수 있다. 불일치가 나면 seed를 바꾸지 말고 해시→정렬→키 매핑→제외 규칙 순으로 원인을 찾는다."
]
assert len(d.tables)==len(explanations)==18
for t,text in zip(d.tables,explanations):insert_after_table(t,text)

# Add a completion marker used only to prevent accidental duplicate application.
last=d.paragraphs[-1];insert_after_paragraph(last,"목차와 표 해설 가독성 보완 완료",shade="FFFFFF")
d.save(docx)

# Structural QA and consistency checks.
d2=Document(docx);heads=[p.text for p in d2.paragraphs if p.style.name.startswith("Heading")]
assert all(x in heads for x in ["1. 한 페이지 핵심 요약","17. S0 관측형 합성 1,000명 증강 방법 상세","18. S0 합성 SQLite 테이블 목록과 데이터 사전","19. 작업자가 그대로 실행할 수 있는 증강 재현 절차"])
assert sum(p.text.startswith("표 해설 |") for p in d2.paragraphs)==18
assert len(d2.tables)==18 and len(d2.inline_shapes)==6
with zipfile.ZipFile(docx) as z:assert z.testzip() is None and len([x for x in z.namelist() if x.startswith("word/media/")])==6

text=log.read_text(encoding="utf-8")+f"\n## 목차·직관성·가독성 보완\n\n- 실제 1~19절과 17·18·19절 하위 항목을 반영한 정적 목차로 교체했다.\n- 핵심 요약에 독자별 추천 읽기 순서와 색상 의미를 추가했다.\n- 표 18개 각각 바로 아래에 핵심 관찰·해석·주의사항을 추가했다.\n- 17.1~18.3을 Heading 2로 변경해 제목 계층을 정리했다.\n- 새 최종 보고서를 만들지 않고 기존 단일 DOCX를 수정했다.\n- 합성 데이터 및 생성기 실행: 미실행\n- 최종 보고서 SHA-256: `{sh(docx)}`\n"
log.write_text(text,encoding="utf-8")
m=json.loads(manifest.read_text(encoding="utf-8"));m["report_docx"]["path"]=str(docx.resolve().relative_to(repo.resolve())).replace("\\","/");m["report_docx"]["size_bytes"]=docx.stat().st_size;m["report_docx"]["sha256"]=sh(docx);m["packaging_execution_log"]["size_bytes"]=log.stat().st_size;m["packaging_execution_log"]["sha256"]=sh(log)
a=m["report_content_addendum"];a["table_of_contents_matches_document"]=True;a["table_explanation_count"]=18;a["reader_route_added"]=True;a["heading_hierarchy_normalized"]=True;a["single_final_report_path"]=m["report_docx"]["path"]
manifest.write_text(json.dumps(m,ensure_ascii=False,indent=2),encoding="utf-8")
print(json.dumps({"report_sha256":sh(docx),"manifest_sha256":sh(manifest),"headings":len(heads),"table_explanations":18,"tables":18,"images":6,"new_report_created":False},ensure_ascii=True,indent=2))
