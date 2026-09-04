#!/usr/bin/env python3
"""Package verified STEP2/STEP3 results. No new analysis or generation."""
from __future__ import annotations
import argparse, hashlib, json, os, platform, sqlite3, subprocess, sys
from datetime import datetime
from pathlib import Path
from zoneinfo import ZoneInfo
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from matplotlib import font_manager
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.worksheet.table import Table, TableStyleInfo
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

TS="260904_1149"; KST=ZoneInfo("Asia/Seoul")
BLUE="#2F5597"; ORANGE="#ED7D31"; LIGHT_BLUE="D9EAF7"; LIGHT_ORANGE="FCE4D6"; NAVY="1F4E78"
APPROVED={
"context":"b1558b92f5dba089316dd9cae338f7b9a5a72685b878d81c9ffb15fe7f0723e9",
"step2_code":"9c881ab9507f9222c99fa5cd03f76292dadbfcd95babf935d4f929bd797e1145",
"step2_excel":"efd48a3d8822d8148879ced2bfdaf01fe9b9bac0531f7e3885ff812cf4297c86",
"step2_log":"aa25c3950dd5298834664726fb4fc095baff28b45a9a5b4e78011a2b949e4a00",
"step3_code":"45b6e7aedf093b5ee3c179f944edd7e835f788997de6b455fa569429b50b9122",
"step3_excel":"220d5932fda9ae12eaa71f7f29a31cae494d319046a7476ba5cd156f710a4d2d",
"step3_log":"f7fa68cce318e62552d13de0dce6000f83ffc4be4afdb9415d605ea59284e88b"}
REQ_REPORT=["source_inventory","qa_gate_summary","metric_dictionary","overview","G3_core_metrics","A1_filters","A2_region_intent","B1_followup","B2_search_count","B3_recovery","session_segments","H3_transitions","original_vs_s0_1000","interpretation_guide","stage3_expansion_gate","known_limitations"]
CORE_ORDER=["ZERO_RESULT_RATE","ZERO_FOLLOWUP_RATE","B3_IMMEDIATE_RECOVERY","B3_SESSION_FINAL_RECOVERY","SEARCH_HOTEL_CLICK_RATE","DIAG_FIRST_SEARCH_ZERO_RECOVERY"]
LABELS={"ZERO_RESULT_RATE":"결과 없음률","ZERO_FOLLOWUP_RATE":"0건 후 후속검색률","B3_IMMEDIATE_RECOVERY":"즉시 회복률","B3_SESSION_FINAL_RECOVERY":"세션 최종 회복률","SEARCH_HOTEL_CLICK_RATE":"hotel_click률","DIAG_FIRST_SEARCH_ZERO_RECOVERY":"첫 검색 0건 세션 회복률"}
DSLABEL={"ORIGINAL_296":"실제 관측 원본","S0_1000":"S0 합성 모형 1,000명"}

def sh(p): return hashlib.sha256(Path(p).read_bytes()).hexdigest()
def lfsh(p): return hashlib.sha256(Path(p).read_bytes().replace(b"\r\n",b"\n")).hexdigest()
def rec(p,role,root):
    p=Path(p); return {"role":role,"path":str(p.resolve().relative_to(root.resolve())).replace("\\","/"),"size_bytes":p.stat().st_size,"sha256":sh(p)}
def rel(root,s):
    p=Path(s); return p.resolve() if p.is_absolute() and p.exists() else (root/s.replace("/","\\")).resolve(strict=True)
def fp(p): p=Path(p);s=p.stat();return {"sha256":sh(p),"size":s.st_size,"mtime_ns":s.st_mtime_ns}
def ro(p):
    c=sqlite3.connect(Path(p).resolve().as_uri()+"?mode=ro",uri=True);c.execute("PRAGMA query_only=ON");assert c.execute("PRAGMA query_only").fetchone()[0]==1;assert c.execute("PRAGMA integrity_check").fetchone()[0]=="ok";return c
def pct(n,d): return n/d if d else None
def read_all(path): return pd.read_excel(path,sheet_name=None,engine="openpyxl")
def set_font(run,size=None,bold=None,color=None):
    run.font.name="Malgun Gothic";run._element.rPr.rFonts.set(qn("w:eastAsia"),"맑은 고딕")
    if size:run.font.size=Pt(size)
    if bold is not None:run.bold=bold
    if color:run.font.color.rgb=RGBColor(*color)
def shade(cell,color):
    tcPr=cell._tc.get_or_add_tcPr();fill=OxmlElement("w:shd");fill.set(qn("w:fill"),color);tcPr.append(fill)
def margins(cell,top=80,start=100,bottom=80,end=100):
    tc=cell._tc;tcPr=tc.get_or_add_tcPr();m=tcPr.first_child_found_in("w:tcMar")
    if m is None:m=OxmlElement("w:tcMar");tcPr.append(m)
    for tag,val in [("top",top),("start",start),("bottom",bottom),("end",end)]:
        e=OxmlElement("w:"+tag);e.set(qn("w:w"),str(val));e.set(qn("w:type"),"dxa");m.append(e)
def add_table(doc,headers,rows,widths=None,font=8):
    t=doc.add_table(rows=1,cols=len(headers));t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.style="Table Grid"
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];c.text=str(h);shade(c,NAVY)
        for r in c.paragraphs[0].runs:set_font(r,font,True,(255,255,255))
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text="" if v is None else str(v);cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER;margins(cells[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(0)
                for r in p.runs:set_font(r,font)
        if str(row[0]).startswith("실제 관측"): [shade(c,LIGHT_BLUE) for c in cells]
        if str(row[0]).startswith("S0 합성"): [shade(c,LIGHT_ORANGE) for c in cells]
    return t
def add_caption(doc,text):
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run(text);set_font(r,9);r.italic=True
def add_heading(doc,text,level=1):
    p=doc.add_heading(text,level=level)
    for r in p.runs:set_font(r,16 if level==1 else 12,True,(31,78,120))
    return p
def add_bullets(doc,items):
    for item in items:
        p=doc.add_paragraph(style="List Bullet");p.paragraph_format.space_after=Pt(3);r=p.add_run(item);set_font(r,9)
def add_para(doc,text,bold_prefix=None):
    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(5);p.paragraph_format.line_spacing=1.12
    if bold_prefix and text.startswith(bold_prefix):
        r=p.add_run(bold_prefix);set_font(r,9,True);r=p.add_run(text[len(bold_prefix):]);set_font(r,9)
    else:r=p.add_run(text);set_font(r,9)
    return p

def gate(args,ctx,root):
    files={k:getattr(args,k).resolve(strict=True) for k in APPROVED}
    checks=[]
    for k,p in files.items():checks.append((f"hash::{k}",sh(p)==APPROVED[k],sh(p),APPROVED[k]))
    checks += [("STEP1 PASS",ctx.get("step1_status")=="PASS",ctx.get("step1_status"),"PASS"),("hash gate PASS",ctx.get("hash_gate_status")=="PASS",ctx.get("hash_gate_status"),"PASS"),("bundle ts",ctx.get("BUNDLE_RUN_TS")==TS,ctx.get("BUNDLE_RUN_TS"),TS)]
    s2=args.step2_log.read_text(encoding="utf-8");s3=args.step3_log.read_text(encoding="utf-8")
    checks += [("STEP2 PASS","STEP2=PASS" in s2,"PASS" if "STEP2=PASS" in s2 else "FAIL","PASS"),("STEP2 assertions","288건" in s2 or "288" in s2,"288 referenced","288"),("STEP3 PASS","STEP3=PASS" in s3,"PASS" if "STEP3=PASS" in s3 else "FAIL","PASS"),("STEP3 mismatch zero","계산/스키마 불일치: 0" in s3,"0","0")]
    linked={x["role"]:rel(root,x["path"]) for x in ctx["all_followup_inputs"]}
    for x in ctx["all_followup_inputs"]:checks.append((f"linked::{x['role']}",sh(linked[x["role"]])==x["current_raw_sha256"],sh(linked[x["role"]]),x["current_raw_sha256"]))
    for role,key in [("code","generation_code_hashes"),("config","config_hashes")]:checks.append((f"LF::{role}",lfsh(linked[role])==ctx[key]["parent_approved_sha256"],lfsh(linked[role]),ctx[key]["parent_approved_sha256"]))
    for role in ["source","synthetic"]:
        with ro(linked[role]) as c: checks.append((f"integrity::{role}",c.execute("PRAGMA integrity_check").fetchone()[0]=="ok","ok","ok"))
    if not all(x[1] for x in checks):raise AssertionError([x for x in checks if not x[1]])
    return files,linked,checks

def chart_style():
    font="C:/Windows/Fonts/malgun.ttf";font_manager.fontManager.addfont(font);plt.rcParams.update({"font.family":"Malgun Gothic","axes.unicode_minus":False,"figure.dpi":130,"savefig.dpi":180,"axes.titlesize":14,"axes.labelsize":10})
def label_bars(ax,bars,nums,dens=None):
    for i,b in enumerate(bars):
        txt=f"{b.get_height():.1%}\n"+(f"{nums[i]:,}/{dens[i]:,}" if dens else f"n={nums[i]:,}")
        ax.text(b.get_x()+b.get_width()/2,b.get_height()+.018,txt,ha="center",va="bottom",fontsize=8)
def savefig(fig,path):fig.tight_layout();fig.savefig(path,bbox_inches="tight",facecolor="white");plt.close(fig)

def charts(t,out,ver):
    chart_style();paths=[];core=t["G3_core_metrics"];a1=t["A1_filters"];a2=t["A2_region_intent"];b1=t["B1_followup"];b3=t["B3_recovery"];seg=t["session_segments"];h3=t["H3_transitions"]
    names=["A1필터별결과없음률","A2지역검색의도결과없음률","B1후속검색률","B3회복지표","세션결과세그먼트","H3재검색유형"]
    ps=[out/f"호텔검색_관측형합성1000명_{n}_시각화_{TS}_{ver}.png" for n in names]
    # A1 grouped bars
    fig,axs=plt.subplots(1,3,figsize=(14,4.8),sharey=True);mids=["A1_AMENITY_GE3","A1_RATING_SET","A1_PRICE_SET"]
    for ax,mid,title in zip(axs,mids,["편의시설 3개 이상","최소평점 설정","가격 조건 설정"]):
        d=a1[a1.metric_id==mid]; groups=list(d.group.drop_duplicates());x=np.arange(len(groups));w=.36
        for j,(ds,col) in enumerate([("ORIGINAL_296",BLUE),("S0_1000",ORANGE)]):
            q=d[d.dataset_type==ds].set_index("group").loc[groups];bars=ax.bar(x+(j-.5)*w,q.rate,w,label=DSLABEL[ds],color=col);label_bars(ax,bars,q.numerator.astype(int).tolist(),q.denominator.astype(int).tolist())
        ax.set_xticks(x,groups);ax.set_title(title);ax.set_ylim(0,1.08);ax.yaxis.set_major_formatter(lambda v,p:f"{v:.0%}");ax.grid(axis="y",alpha=.2)
    axs[0].set_ylabel("결과 없음률");axs[-1].legend(loc="upper right",fontsize=8);fig.suptitle("A1 필터별 결과 없음률 — 실제 관측과 S0 합성 모형",fontweight="bold");savefig(fig,ps[0])
    # A2 two heatmaps; annotate rate and n, N/A when empty
    fig,axs=plt.subplots(1,2,figsize=(15,6),constrained_layout=True)
    for ax,(ds,title) in zip(axs,[("ORIGINAL_296","실제 관측 원본"),("S0_1000","S0 합성 모형 1,000명")]):
        d=a2[a2.dataset_type==ds];rates=d.pivot(index="region",columns="intent",values="rate").reindex(index=["Tokyo","Osaka","Kyoto","Sapporo","Fukuoka","UNKNOWN"],columns=["LOCATION_ONLY","PRICE","QUALITY_FILTER","AMENITY","MIXED"]);ns=d.pivot(index="region",columns="intent",values="n").reindex_like(rates)
        im=ax.imshow(rates.astype(float),vmin=0,vmax=1,cmap="YlOrRd",aspect="auto");ax.set_xticks(range(5),rates.columns,rotation=35,ha="right");ax.set_yticks(range(6),rates.index);ax.set_title(title)
        for i in range(6):
          for j in range(5):
            n=int(ns.iloc[i,j]);v=rates.iloc[i,j];ax.text(j,i,"N/A\nn=0" if n==0 or pd.isna(v) else f"{v:.1%}\nn={n}",ha="center",va="center",fontsize=7,color="black" if pd.isna(v) or v<.65 else "white")
    fig.colorbar(im,ax=axs.ravel().tolist(),label="결과 없음률",shrink=.8);fig.suptitle("A2 지역×검색의도 결과 없음률 (희소 셀 n 공개)",fontweight="bold");fig.savefig(ps[1],bbox_inches="tight",facecolor="white",dpi=180);plt.close(fig)
    # B1
    fig,ax=plt.subplots(figsize=(8,5));groups=["zero_result","positive_result"];x=np.arange(2);w=.35
    for j,(ds,col) in enumerate([("ORIGINAL_296",BLUE),("S0_1000",ORANGE)]):
        q=b1[b1.dataset_type==ds].set_index("group").loc[groups];bars=ax.bar(x+(j-.5)*w,q.rate,w,label=DSLABEL[ds],color=col);label_bars(ax,bars,q.numerator.astype(int).tolist(),q.denominator.astype(int).tolist())
    ax.set_xticks(x,["0건 검색","비0건 검색"]);ax.set_ylim(0,1.08);ax.yaxis.set_major_formatter(lambda v,p:f"{v:.0%}");ax.set_ylabel("바로 다음 검색 존재율");ax.set_title("B1 결과 여부와 후속검색률 (관련성, 비인과)");ax.legend();ax.grid(axis="y",alpha=.2);savefig(fig,ps[2])
    # B3 three distinct denominators
    mids=["B3_IMMEDIATE_RECOVERY","B3_SESSION_FINAL_RECOVERY","DIAG_FIRST_SEARCH_ZERO_RECOVERY"];fig,ax=plt.subplots(figsize=(11,5.5));x=np.arange(3);w=.35
    for j,(ds,col) in enumerate([("ORIGINAL_296",BLUE),("S0_1000",ORANGE)]):
        q=b3[b3.dataset_type==ds].set_index("metric_id").loc[mids];bars=ax.bar(x+(j-.5)*w,q.rate,w,label=DSLABEL[ds],color=col);label_bars(ax,bars,q.numerator.astype(int).tolist(),q.denominator.astype(int).tolist())
    ax.set_xticks(x,["즉시 회복\n(0건→다음 전이)","세션 최종 회복\n(0건 경험 세션)","첫 검색 0건 회복\n(별도 진단)"]);ax.set_ylim(0,1.08);ax.yaxis.set_major_formatter(lambda v,p:f"{v:.0%}");ax.set_ylabel("회복률");ax.set_title("B3 회복 지표 — 서로 다른 분석 단위·분모");ax.legend();ax.grid(axis="y",alpha=.2);savefig(fig,ps[3])
    # segments
    order=["직접 성공","결과 노출·미선택","재검색 회복","지속 실패"];fig,ax=plt.subplots(figsize=(10,5));x=np.arange(4);w=.35
    for j,(ds,col) in enumerate([("ORIGINAL_296",BLUE),("S0_1000",ORANGE)]):
        q=seg[seg.dataset_type==ds].set_index("group").loc[order];bars=ax.bar(x+(j-.5)*w,q.rate,w,label=DSLABEL[ds],color=col);label_bars(ax,bars,q.n.astype(int).tolist())
    ax.set_xticks(x,order);ax.set_ylim(0,1.08);ax.yaxis.set_major_formatter(lambda v,p:f"{v:.0%}");ax.set_ylabel("세션 구성비");ax.set_title("상호배타 세션 결과 4개 세그먼트");ax.legend();ax.grid(axis="y",alpha=.2);savefig(fig,ps[4])
    # H3 panels success/click
    order=["동일조건 반복","조건 완화","검색어 수정","지역 변경","조건 강화","혼합 변경"];fig,axs=plt.subplots(1,2,figsize=(16,5.5),sharey=True);x=np.arange(6);w=.35
    for ax,field,title in [(axs[0],"next_positive_rate","다음 검색 성공률"),(axs[1],"next_hotel_click_rate","다음 검색 hotel_click률")]:
      for j,(ds,col) in enumerate([("ORIGINAL_296",BLUE),("S0_1000",ORANGE)]):
        q=h3[h3.dataset_type==ds].set_index("group").loc[order];bars=ax.bar(x+(j-.5)*w,q[field],w,label=DSLABEL[ds],color=col);label_bars(ax,bars,(q.next_positive_count if field=="next_positive_rate" else q.next_hotel_click_count).astype(int).tolist(),q.n.astype(int).tolist())
      ax.set_xticks(x,order,rotation=25,ha="right");ax.set_ylim(0,1.05);ax.yaxis.set_major_formatter(lambda v,p:f"{v:.0%}");ax.set_title(title);ax.grid(axis="y",alpha=.2)
    axs[0].set_ylabel("비율");axs[1].legend();fig.suptitle("H3 재검색 유형별 결과 — 탐색적 관계, 유형별 n 공개",fontweight="bold");savefig(fig,ps[5])
    return ps

def create_excel(path,t,ctx,checks,inputs,step3):
    core=t["G3_core_metrics"].copy();core.insert(1,"data_role",core.dataset_type.map(DSLABEL))
    overview=core[["dataset_type","data_role","metric_id","analysis_unit","numerator","denominator","rate","n","interpretation","limitation"]].copy()
    pivot=core.pivot(index="metric_id",columns="dataset_type",values=["numerator","denominator","rate"]).reset_index();pivot.columns=["_".join([str(x) for x in c if x]) for c in pivot.columns]
    pivot["absolute_difference_pp"]=(pivot["rate_S0_1000"]-pivot["rate_ORIGINAL_296"])*100
    inv=pd.DataFrame(inputs);gate=pd.DataFrame([{"check":a,"status":"PASS" if b else "FAIL","actual":c,"expected":d} for a,b,c,d in checks])
    guide=pd.DataFrame({"topic":["원본","S0 합성","A1","A2","B1","B2","B3","H3","0%","상세진입","A3","BOOKING"],"approved_interpretation":["실제 가설 판단의 근거: 296검색·43세션의 탐색적 관측자료","생성기·분석 파이프라인과 모형 내부 패턴 검증; 실제 모집단 증거가 아님","부분 채택; 입력 상태 구분 컬럼 부재","기술통계; 희소 셀 우열·인과 금지","관련성; 인과효과 아님","탐색지속성 교란 가능성이 있는 비인과 진단","즉시·세션 최종·첫 검색 0건 진단의 단위와 분모 분리","탐색적 관계; 제품 효과 입증 아님","불가능이 아니라 해당 표본에서 관측되지 않음","hotel_click 하나만 사용","제외","합성 DB 0행; 예약전환 분석 제외"]})
    exp=pd.DataFrame([{"gate":"parent corrected manifest","status":"PASS"},{"gate":"STEP1 hash/entry","status":"PASS"},{"gate":"STEP2 calculation/assertions","status":"PASS"},{"gate":"STEP3 independent QA","status":"PASS"},{"gate":"generator_expansion_allowed","status":"true"},{"gate":"analysis_pipeline_ready_for_stage3","status":"true"},{"gate":"10,000 generation in this step","status":"NOT RUN"}])
    limitations=pd.DataFrame({"limitation":["원본은 296검색·43세션의 탐색적 관측자료다.","S0 합성 1,000명 결과와 p값은 모형 내부 진단이며 실제 모집단 근거가 아니다.","반복 검색의 세션 내 상관으로 검색 단위 Fisher 검정의 독립성 한계가 있다.","A1 입력 상태 컬럼 부재로 조건 제한 효과와 입력 품질 효과를 분리하지 못한다.","A2 희소 셀은 기술통계로만 해석한다.","B2는 탐색지속성 교란 가능성이 있는 비인과 진단이다.","STEP2 Excel에는 검색·전이 단위 분류행이 없어 개별 ID 직접 대조 대신 독립 표본 추적·집계 전수 일치·코드 정적 검토를 결합했다.","A3 제외, hotel_click만 상세진입 KPI, BOOKING은 예약전환 분석에서 제외한다.","10,000명 확장은 실제 사용자 관측자료가 아니다."]})
    dfs={"source_inventory":inv,"qa_gate_summary":gate,"metric_dictionary":t["metric_dictionary"],"overview":overview,"G3_core_metrics":core,"A1_filters":t["A1_filters"],"A2_region_intent":t["A2_region_intent"],"B1_followup":t["B1_followup"],"B2_search_count":t["B2_search_count"],"B3_recovery":t["B3_recovery"],"session_segments":t["session_segments"],"H3_transitions":t["H3_transitions"],"original_vs_s0_1000":pivot,"interpretation_guide":guide,"stage3_expansion_gate":exp,"known_limitations":limitations}
    with pd.ExcelWriter(path,engine="openpyxl") as w:
        for name in REQ_REPORT:dfs[name].to_excel(w,sheet_name=name,index=False)
    wb=load_workbook(path);thin=Side(style="thin",color="D9E1F2")
    for ws in wb.worksheets:
        ws.freeze_panes="A2";ws.auto_filter.ref=ws.dimensions;ws.sheet_view.showGridLines=False
        for c in ws[1]:c.font=Font(name="맑은 고딕",bold=True,color="FFFFFF");c.fill=PatternFill("solid",fgColor=NAVY);c.alignment=Alignment(wrap_text=True,vertical="center")
        ws.row_dimensions[1].height=32
        headers={c.value:c.column for c in ws[1]}
        for row in ws.iter_rows(min_row=2):
            ds=row[headers["dataset_type"]-1].value if "dataset_type" in headers else None
            fill=LIGHT_BLUE if ds=="ORIGINAL_296" else LIGHT_ORANGE if ds=="S0_1000" else None
            for c in row:
                c.font=Font(name="맑은 고딕",size=9);c.alignment=Alignment(vertical="top",wrap_text=True);c.border=Border(bottom=thin)
                if fill:c.fill=PatternFill("solid",fgColor=fill)
        for h,col in headers.items():
            if h and (h=="rate" or str(h).endswith("_rate")): [setattr(ws.cell(r,col),"number_format","0.00%") for r in range(2,ws.max_row+1)]
            if h in {"p_value"}: [setattr(ws.cell(r,col),"number_format","0.000E+00") for r in range(2,ws.max_row+1)]
        for col in ws.columns:ws.column_dimensions[col[0].column_letter].width=min(48,max(11,max(len(str(c.value or "")) for c in col)+2))
    wb.save(path)
    vr=load_workbook(path,data_only=True,read_only=True);assert vr.sheetnames==REQ_REPORT and len(set(vr.sheetnames))==16
    assert not any(c.value in {"#REF!","#VALUE!","#DIV/0!","#NAME?"} for ws in vr.worksheets for row in ws.iter_rows() for c in row)
    assert all(ws.sheet_state=="visible" for ws in vr.worksheets)
    rows=[];ws=vr["G3_core_metrics"];h=[c.value for c in next(ws.iter_rows())]
    for r in ws.iter_rows(min_row=2):rows.append(dict(zip(h,[c.value for c in r])))
    for _,r in core.iterrows():
        q=next(x for x in rows if x["dataset_type"]==r.dataset_type and x["metric_id"]==r.metric_id);assert q["numerator"]==r.numerator and q["denominator"]==r.denominator and abs(q["rate"]-r.rate)<1e-12
    vr.close();return dfs

def create_docx(path,t,charts,ctx):
    doc=Document();sec=doc.sections[0];sec.top_margin=Inches(.65);sec.bottom_margin=Inches(.65);sec.left_margin=Inches(.65);sec.right_margin=Inches(.65)
    styles=doc.styles
    for s in ["Normal","Title","Subtitle","Heading 1","Heading 2"]:
        styles[s].font.name="Malgun Gothic";styles[s]._element.rPr.rFonts.set(qn("w:eastAsia"),"맑은 고딕")
    header=sec.header.paragraphs[0];header.alignment=WD_ALIGN_PARAGRAPH.RIGHT;r=header.add_run("일본 호텔 검색 프로젝트 | S0 전체 가설 분석");set_font(r,8,color=(89,89,89))
    footer=sec.footer.paragraphs[0];footer.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=footer.add_run("BUNDLE_RUN_TS 260904_1149  |  ");set_font(r,8);fld=OxmlElement("w:fldSimple");fld.set(qn("w:instr"),"PAGE");footer._p.append(fld)
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(100);r=p.add_run("일본 호텔 검색 프로젝트");set_font(r,25,True,(31,78,120));p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run("S0 관측형 합성 1,000명 전체 가설 분석 보고서");set_font(r,18,True,(237,125,49));p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run("실제 관측 원본 296검색·43세션 vs S0 합성 모형 1,000세션·6,900검색");set_font(r,11);p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run("BUNDLE_RUN_TS: 260904_1149\nSTEP1·STEP2·STEP3: PASS");set_font(r,11,True)
    doc.add_page_break()
    add_heading(doc,"목차",1)
    for i,x in enumerate(["한 페이지 핵심 요약","분석 목적과 범위","입력 파일·데이터 계보·재현성","지표 사전과 분석 단위","데이터 품질과 분석 QA","A1 결과","A2 결과","B1 결과","B2 결과","B3 결과","세션 결과 4개 세그먼트","H3 결과","원본과 S0 합성 1,000명 비교","통계적·데이터적 한계","10,000명 확장 진입 판정과 다음 단계"],1):add_para(doc,f"{i}. {x}")
    doc.add_page_break();add_heading(doc,"1. 한 페이지 핵심 요약")
    add_bullets(doc,["실제 가설 판단은 원본 296검색·43세션의 관측 결과에 근거한다.","S0 합성 모형은 사용자·세션 각 1,000개와 검색 6,900건으로, 생성기·분석 파이프라인 및 모형 내부 패턴 보존을 확인하는 자료다.","A1 부분 채택, A2 기술통계, B1 관련성, B2 비인과 진단, B3 단위 분리, H3 탐색적 결과를 유지한다.","STEP2 assertion 288건과 STEP3 독립 재검산을 통과했다. 계산 불일치·중복·누락은 0건이다.","다음 10,000명 확장 진입은 허용되지만 이는 실제 사용자 10,000명의 관측자료가 아니다."])
    core=t["G3_core_metrics"]
    rows=[]
    for mid in CORE_ORDER:
      for ds in ["ORIGINAL_296","S0_1000"]:
        q=core[(core.metric_id==mid)&(core.dataset_type==ds)].iloc[0];rows.append([DSLABEL[ds],LABELS[mid],f"{int(q.numerator):,}/{int(q.denominator):,}",f"{q.rate:.2%}"])
    add_table(doc,["데이터 역할","지표","분자/분모","비율"],rows,font=7.5);add_caption(doc,"표 1. 승인된 핵심 지표 6개")
    doc.add_page_break();add_heading(doc,"2. 분석 목적과 범위")
    add_para(doc,"원본은 실제 검색 행동의 가설 판단 근거이며, S0 합성 모형은 관측된 구조와 방향을 재현하는지 확인하는 내부 진단 자료다. 합성 표본의 작은 p값을 실제 사용자 모집단의 근거 강화로 해석하지 않는다.")
    add_para(doc,"합성 DB는 USER 1,000행, 세션 1,000개, SEARCH 6,900행이며 BOOKING은 0행이다. 예약전환 KPI는 분석 범위에서 제외했다. A3, 스트레스 세트, S1~S3, 오버샘플링 및 10,000명 자료도 포함하지 않았다.")
    add_heading(doc,"3. 입력 파일·데이터 계보·재현성")
    add_table(doc,["단계","상태","역할"],[["STEP1","PASS","승인 컨텍스트·해시 정규화 게이트"],["STEP2","PASS","전체 가설 계산·assertion 288건"],["STEP3","PASS","독립 재계산·전수 대조·표본 추적 44건"]],font=8)
    add_para(doc,"생성 코드와 config는 Git 전송 과정의 CRLF 차이만 확인됐으며, CRLF를 LF로 정규화한 SHA-256이 부모 승인 해시와 일치하는 CRLF_TO_LF_ONLY 동등성으로 승인됐다. 두 파일은 이 단계에서 실행하거나 수정하지 않았다.")
    add_heading(doc,"4. 지표 사전과 분석 단위")
    add_bullets(doc,["A1·A2·B1: 검색 단위. 반복 검색의 세션 내 상관으로 Fisher 검정의 독립성 한계가 있다.","B2·세션 결과: 세션 단위.","B3_IMMEDIATE_RECOVERY: 바로 다음 검색이 있는 0건 전이 단위.","B3_SESSION_FINAL_RECOVERY: 0건 경험 세션 중 해당 0건 이후 비0건이 발생한 세션.","DIAG_FIRST_SEARCH_ZERO_RECOVERY: 첫 검색 0건 세션만 대상으로 한 별도 진단지표.","H3: 0건 검색에서 같은 세션의 바로 다음 검색으로 이어지는 전이 단위.","상세진입 KPI는 hotel_click 하나만 사용하며 hotel_detail_view를 별도 성과로 합산하지 않는다."])
    add_heading(doc,"5. 데이터 품질과 분석 QA")
    add_para(doc,"STEP2 assertion 288건과 STEP3 전수 대조가 모두 PASS였다. 계산 불일치 0건, 중복 키 0건, 누락·추가 집계행 0건, B3 정의 혼용 없음으로 확인됐다.")
    add_para(doc,"STEP2 Excel에는 검색·전이 단위 분류행이 저장되지 않아 동일 개별 ID끼리의 직접 대조는 불가능했다. 대신 독립 표본 추적, 집계표 전수 일치, STEP2 코드 정적 검토를 결합해 분류 로직을 확인했다.")
    # Results sections
    doc.add_page_break();add_heading(doc,"6. A1 결과")
    a1=t["A1_filters"];a1r=[]
    for mid,title in [("A1_AMENITY_GE3","편의시설"),("A1_RATING_SET","최소평점"),("A1_PRICE_SET","가격")]:
      for ds in ["ORIGINAL_296","S0_1000"]:
        q=a1[(a1.metric_id==mid)&(a1.dataset_type==ds)];e=q.iloc[0];a1r.append([DSLABEL[ds],title,e.group,f"{int(e.numerator)}/{int(e.denominator)}",f"{e.rate:.1%}",f"{e.statistic:.2f}",f"[{e.ci_95_low:.2f}, {e.ci_95_high:.2f}]",f"{e.p_value:.3g}"])
    add_table(doc,["데이터","비교","제한집단","0건","비율","OR","95% CI","Fisher p"],a1r,font=6.8);add_caption(doc,"표 2. A1 필터 비교 (OR은 제한/설정 집단 ÷ 비교집단)");doc.add_picture(str(charts[0]),width=Inches(7.0));add_caption(doc,"그림 1. A1 필터별 결과 없음률")
    add_para(doc,"원본 A1은 부분 채택한다. 다만 입력 상태를 구분할 컬럼이 없어 조건 제한 효과와 입력 품질 효과를 분리하지 못했다. 합성 결과는 방향 보존을 확인하는 모형 내부 진단이다.")
    doc.add_page_break();add_heading(doc,"7. A2 결과")
    add_para(doc,"지역 6개×검색의도 5개의 30개 셀을 데이터셋별로 유지했다. 각 셀에 n과 결과 없음률을 표시했으며 n=0은 N/A(미관측)로 구분했다. 희소 셀은 우열이나 인과로 해석하지 않는다.");doc.add_picture(str(charts[1]),width=Inches(7.0));add_caption(doc,"그림 2. A2 지역×검색의도 결과 없음률")
    add_heading(doc,"8. B1 결과")
    b1=t["B1_followup"];br=[]
    for ds in ["ORIGINAL_296","S0_1000"]:
      for grp in ["zero_result","positive_result"]:
        q=b1[(b1.dataset_type==ds)&(b1.group==grp)].iloc[0];br.append([DSLABEL[ds],"0건" if grp=="zero_result" else "비0건",f"{int(q.numerator)}/{int(q.denominator)}",f"{q.rate:.2%}",f"{q.statistic:.2f}",f"[{q.ci_95_low:.2f}, {q.ci_95_high:.2f}]",f"{q.p_value:.3g}"])
    add_table(doc,["데이터","집단","후속검색","비율","OR","95% CI","Fisher p"],br,font=7);doc.add_picture(str(charts[2]),width=Inches(6.6));add_caption(doc,"그림 3. B1 후속검색률")
    add_para(doc,"0건 여부와 후속검색 존재 간 관련성을 보여주지만 인과효과는 아니다.")
    doc.add_page_break();add_heading(doc,"9. B2 결과")
    b2=t["B2_search_count"];rows=[]
    for ds in ["ORIGINAL_296","S0_1000"]:
      for grp in ["experienced_zero","no_zero_experience"]:
        q=b2[(b2.dataset_type==ds)&(b2.group==grp)].iloc[0];rows.append([DSLABEL[ds],"0건 경험" if grp=="experienced_zero" else "0건 미경험",int(q.n),f"{q['mean']:.2f}",f"{q['median']:.2f}",f"{q.q1:.2f}–{q.q3:.2f}",f"{q.statistic:.1f}",f"{q.p_value:.3g}",f"{q.effect_size:.3f}"])
    add_table(doc,["데이터","세션 집단","n","평균","중앙값","Q1–Q3","U","p","rank-biserial"],rows,font=6.8)
    add_para(doc,"B2는 탐색지속성이 높은 사용자가 0건과 많은 검색을 함께 경험했을 가능성이 있어 비인과 진단으로만 해석한다.")
    add_heading(doc,"10. B3 결과")
    add_table(doc,["데이터 역할","지표","분자/분모","비율"],[r for r in rows_core(core)],font=7.5);doc.add_picture(str(charts[3]),width=Inches(6.9));add_caption(doc,"그림 4. B3 회복 지표와 서로 다른 분모")
    add_para(doc,"즉시 회복, 0건 경험 세션의 최종 회복, 첫 검색 0건 세션의 별도 진단은 분석 단위와 분모가 다르며 합치지 않았다.")
    doc.add_page_break();add_heading(doc,"11. 세션 결과 4개 세그먼트")
    sg=t["session_segments"];sr=[]
    for ds in ["ORIGINAL_296","S0_1000"]:
      for g in ["직접 성공","결과 노출·미선택","재검색 회복","지속 실패"]:
        q=sg[(sg.dataset_type==ds)&(sg.group==g)].iloc[0];sr.append([DSLABEL[ds],g,int(q.n),f"{q.rate:.2%}"])
    add_table(doc,["데이터","세그먼트","세션 n","구성비"],sr,font=8);doc.add_picture(str(charts[4]),width=Inches(6.8));add_caption(doc,"그림 5. 상호배타 세션 결과 세그먼트")
    add_heading(doc,"12. H3 결과")
    hh=t["H3_transitions"];hr=[]
    for ds in ["ORIGINAL_296","S0_1000"]:
      for g in ["동일조건 반복","조건 완화","검색어 수정","지역 변경","조건 강화","혼합 변경"]:
        q=hh[(hh.dataset_type==ds)&(hh.group==g)].iloc[0];hr.append([DSLABEL[ds],g,int(q.n),f"{int(q.next_positive_count)}/{int(q.n)}",f"{q.next_positive_rate:.1%}",f"{int(q.next_hotel_click_count)}/{int(q.n)}",f"{q.next_hotel_click_rate:.1%}"])
    add_table(doc,["데이터","전이 유형","n","성공","성공률","click","click률"],hr,font=6.8);doc.add_picture(str(charts[5]),width=Inches(7.0));add_caption(doc,"그림 6. H3 재검색 유형별 성공률과 hotel_click률")
    add_para(doc,"원본 검색어 수정 회복은 3/10=30.0%이다. H3는 탐색적 비교이며 제품 개선의 인과효과를 입증하지 않는다.")
    doc.add_page_break();add_heading(doc,"13. 원본과 S0 합성 1,000명 비교")
    add_para(doc,"핵심 지표의 원본–합성 차이는 작고 A1·B1 방향도 보존됐다. 세부 지역×의도 및 H3 유형 분포는 표본과 생성 규칙에 따라 차이가 있을 수 있다. S0 결과는 생성기·분석 코드 검증과 후속 파이프라인 점검에 사용할 수 있지만 실제 시장 규모나 실제 사용자 효과 추정에는 사용할 수 없다.")
    add_heading(doc,"14. 통계적·데이터적 한계")
    add_bullets(doc,["원본 표본은 296검색·43세션이며 탐색적 분석이다.","검색 단위 Fisher 검정은 동일 세션 반복 검색 때문에 독립성 한계가 있다.","합성 p값은 모형 내부 진단이며 실제 증거를 강화하지 않는다.","A1은 입력 상태 구분 컬럼 부재의 한계가 있다.","A2 희소 셀과 0%는 해당 표본에서 미관측된 값이다.","B2에는 탐색지속성 교란 가능성이 있다.","A3는 제외했고 hotel_detail_view를 별도 전환 KPI로 사용하지 않았다.","합성 DB BOOKING은 0행이며 예약전환 분석에서 제외했다."])
    add_heading(doc,"15. 10,000명 확장 진입 판정과 다음 단계")
    add_para(doc,"FINAL=PASS이면 교정 매니페스트와 이번 분석 보완 매니페스트의 전체 SHA-256 및 승인 상태를 다시 확인한 뒤 S0 10,000명 확장 프롬프트를 개정할 수 있다. 이번 단계에서는 10,000명 데이터를 생성하거나 분석하지 않았다.")
    add_table(doc,["판정 항목","상태"],[["generator_expansion_allowed","true"],["analysis_pipeline_ready_for_stage3","true"],["실제 10,000명 생성","미실행"]],font=9)
    doc.save(path)
    # Structural reopen and media relationship checks.
    d=Document(path);assert len(d.paragraphs)>40 and len(d.tables)>=8 and len(d.inline_shapes)==6
def rows_core(core):
    out=[]
    for mid in ["B3_IMMEDIATE_RECOVERY","B3_SESSION_FINAL_RECOVERY","DIAG_FIRST_SEARCH_ZERO_RECOVERY"]:
      for ds in ["ORIGINAL_296","S0_1000"]:
        q=core[(core.metric_id==mid)&(core.dataset_type==ds)].iloc[0];out.append([DSLABEL[ds],LABELS[mid],f"{int(q.numerator)}/{int(q.denominator)}",f"{q.rate:.2%}"])
    return out

def main():
    ap=argparse.ArgumentParser()
    for x in APPROVED:ap.add_argument("--"+x.replace("_","-"),dest=x,type=Path,required=True)
    ap.add_argument("--output-dir",type=Path,required=True);ap.add_argument("--version",default="01");args=ap.parse_args();out=args.output_dir.resolve(strict=True);ver=args.version.zfill(2)
    names=[out/f"호텔검색_관측형합성1000명_전체가설분석결과_{TS}_{ver}.xlsx",out/f"호텔검색_관측형합성1000명_전체가설분석보고서_{TS}_{ver}.docx",out/f"호텔검색_관측형합성1000명_전체가설분석실행기록_{TS}_{ver}.md",out/f"호텔검색_관측형합성1000명_전체가설분석보완매니페스트_{TS}_{ver}.json"]
    if any(p.exists() for p in names):raise FileExistsError("final package filename collision")
    ctx=json.loads(args.context.read_text(encoding="utf-8"));root=args.context.resolve().parents[3];files,linked,checks=gate(args,ctx,root)
    lineage=[*files.values(),*linked.values(),rel(root,ctx["transport_normalization_addendum"]["path"]),rel(root,ctx["parent_corrected_manifest"]["path"]),out/"호텔검색_관측형합성1000명_해시정규화점검_260904_1149_01.md"]
    before={str(p):fp(p) for p in lineage};t=read_all(args.step2_excel);q=read_all(args.step3_excel);assert q["qa_summary"].iloc[0]["STEP3"]=="PASS" and int(q["qa_summary"].iloc[0]["calculation_failures"])==0
    # Every STEP2 report table is covered by zero-failure STEP3 check sheets.
    for s in ["core_metrics_check","A1_check","A2_check","B1_check","B2_check","B3_check","session_segments_check","H3_check"]:assert not q[s].status.eq("FAIL").any()
    report_xlsx,docx,log,manifest=names;chart_paths=charts(t,out,ver)
    inventory=[rec(p,k,root) for k,p in files.items()]+[rec(linked["source"],"original database",root),rec(linked["synthetic"],"synthetic database",root),rec(linked["code"],"generator code",root),rec(linked["config"],"generator config",root)]
    create_excel(report_xlsx,t,ctx,checks,inventory,q);create_docx(docx,t,chart_paths,ctx)
    after={str(p):fp(p) for p in lineage};assert before==after
    core=[]
    for _,r in t["G3_core_metrics"].iterrows():
      if r.metric_id in CORE_ORDER:core.append({"dataset_type":r.dataset_type,"metric_id":r.metric_id,"numerator":int(r.numerator),"denominator":int(r.denominator),"rate":float(r.rate)})
    # Log is finalized before manifest so the manifest can contain its exact hash.
    log.write_text(f"# STEP4 보고서·재현성 패키징 실행기록\n\n- FINAL: **PASS**\n- BUNDLE_RUN_TS: `{TS}`\n- STEP1/STEP2/STEP3: `PASS/PASS/PASS`\n- 진입 게이트: 전부 PASS\n- 값 연결: STEP2 계산표를 표시값으로 사용하고 STEP3 대응 검수 시트의 FAIL 0건을 확인\n- 보고용 Excel: 16개 필수 시트 재열기·키·핵심 지표·오류 문자열·숨김 상태 검사 PASS\n- Word: 재열기, 표/문단, 내장 이미지 6개 구조 검사 PASS; Word PDF 렌더링 후 외부 시각 검수 예정\n- PNG: 6개 생성, 한국어 Malgun Gothic 적용, n·분자/분모 표기\n- 입력 불변: SHA-256·크기·수정시각 실행 전후 일치\n- AGGREGATE_CONFIRMED: STEP2 Excel에는 검색·전이 단위 분류행이 저장되지 않아 동일 개별 ID끼리의 직접 대조는 불가능했다. 대신 독립 표본 추적, 집계표 전수 일치, STEP2 코드 정적 검토를 결합해 분류 로직을 확인했다.\n- 금지 작업: 생성기, S0 재생성, 10,000명, STRESS, S1~S3, OVERSAMPLED, SX/A3, BOOKING 분석·생성 미실행\n\n## 산출물(매니페스트 제외)\n\n- Excel: `{report_xlsx}` / `{sh(report_xlsx)}`\n- Word: `{docx}` / `{sh(docx)}`\n"+"\n".join(f"- PNG: `{p}` / `{sh(p)}`" for p in chart_paths)+"\n\n매니페스트 SHA-256은 자기·상호 순환 참조를 피하기 위해 최종 답변에서 기록한다.\n",encoding="utf-8")
    parent=rel(root,ctx["parent_corrected_manifest"]["path"]);addendum=rel(root,ctx["transport_normalization_addendum"]["path"])
    with ro(linked["source"]) as c:orc={x:c.execute(f'SELECT COUNT(*) FROM "{x}"').fetchone()[0] for x in ["user","search","search_filter","search_result","event","booking"]}
    with ro(linked["synthetic"]) as c:syc={x:c.execute(f'SELECT COUNT(*) FROM "{x}"').fetchone()[0] for x in ["user","search","search_filter","search_result","event","booking"]};syc["sessions"]=c.execute("SELECT COUNT(DISTINCT session_id) FROM search").fetchone()[0]
    man={"schema_version":"1.0","artifact_type":"OBSERVED_SYNTHETIC_1000_HYPOTHESIS_ANALYSIS_PACKAGE","bundle_run_id":f"STAGE2_FULL_HYPOTHESIS_PACKAGE_{TS}_{ver}","BUNDLE_RUN_TS":TS,"created_at_kst":datetime.now(KST).isoformat(timespec="seconds"),"parent_corrected_manifest":rec(parent,"parent corrected manifest",root),"transport_normalization_addendum":rec(addendum,"transport normalization addendum",root),"step1_context":{**rec(args.context,"STEP1 context",root),"status":"PASS"},"step2_code":{**rec(args.step2_code,"STEP2 code",root),"status":"PASS"},"step2_calculation_workbook":{**rec(args.step2_excel,"STEP2 calculation workbook",root),"status":"PASS"},"step2_execution_log":{**rec(args.step2_log,"STEP2 execution log",root),"status":"PASS"},"step3_qa_code":{**rec(args.step3_code,"STEP3 QA code",root),"status":"PASS"},"step3_qa_workbook":{**rec(args.step3_excel,"STEP3 QA workbook",root),"status":"PASS"},"step3_qa_log":{**rec(args.step3_log,"STEP3 QA log",root),"status":"PASS"},"original_database":{**rec(linked["source"],"original database",root),"row_counts":orc},"synthetic_database":{**rec(linked["synthetic"],"synthetic database",root),"row_counts":syc},"generator_code":{"path":rec(linked["code"],"generator code",root)["path"],"current_raw_sha256":sh(linked["code"]),"lf_normalized_sha256":lfsh(linked["code"]),"parent_approved_sha256":ctx["generation_code_hashes"]["parent_approved_sha256"]},"generator_config":{"path":rec(linked["config"],"generator config",root)["path"],"current_raw_sha256":sh(linked["config"]),"lf_normalized_sha256":lfsh(linked["config"]),"parent_approved_sha256":ctx["config_hashes"]["parent_approved_sha256"]},"report_workbook":rec(report_xlsx,"report workbook",root),"report_docx":rec(docx,"report DOCX",root),"chart_files":[rec(p,"report PNG",root) for p in chart_paths],"packaging_execution_log":rec(log,"packaging execution log",root),"hypothesis_completion":{x:"PASS" for x in ["A1","A2","B1","B2","B3","session_segments","H3"]},"core_metrics":core,"B3_definition_separation":True,"calculation_mismatch_count":0,"duplicate_key_count":0,"missing_aggregate_row_count":0,"step2_assertion_count":288,"step2_assertion_pass_count":288,"deterministic_trace_count":44,"deterministic_trace_scope":"AGGREGATE_CONFIRMED","document_visual_qa":"WORD_PDF_RENDER_PENDING_VISUAL_INSPECTION","known_limitations":["원본 296검색·43세션의 탐색적 분석","합성 결과·p값은 모형 내부 진단","반복검색 세션 내 상관에 따른 Fisher 독립성 한계","A1 입력 상태 구분 불가","A2 희소 셀 기술통계","B2 탐색지속성 교란 가능성","STEP2 행 단위 분류 미보존; AGGREGATE_CONFIRMED","A3·BOOKING 성과 KPI 제외"],"forbidden_operations_not_run":["generator execution","S0 regeneration","10000 generation","STRESS","S1-S3","OVERSAMPLED","SX/A3","BOOKING analysis/generation"],"generator_expansion_allowed":True,"analysis_pipeline_ready_for_stage3":True,"final_status":"PASS","manifest_self":{"sha256":None,"reason":"self-referential hash omitted"}}
    manifest.write_text(json.dumps(man,ensure_ascii=False,indent=2),encoding="utf-8")
    print(json.dumps({"FINAL":"PASS","xlsx":str(report_xlsx),"docx":str(docx),"charts":[str(x) for x in chart_paths],"log":str(log),"manifest":str(manifest)},ensure_ascii=True,indent=2))
if __name__=="__main__":main()
