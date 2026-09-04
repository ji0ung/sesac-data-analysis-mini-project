from pathlib import Path
import hashlib,json,sys,zipfile
from docx import Document
docx=Path(sys.argv[1]);manifest=Path(sys.argv[2]);log=Path(sys.argv[3]);repo=Path(sys.argv[4])
def sh(p):return hashlib.sha256(Path(p).read_bytes()).hexdigest()
d=Document(docx);top=[p.text for p in d.paragraphs if p.style.name=="Heading 1"]
expected=["목차","1. 증강 목적과 데이터 범위","2. S0 1,000명 증강 방법 한눈에 보기","3. 작업자가 그대로 실행할 수 있는 단계별 증강 절차","4. 합성 SQLite 테이블 구조와 데이터 사전","5. 데이터 계보·품질·재현성 QA","6. S0 1,000명 분석 핵심 요약","7. A1·A2: 필터·지역·검색의도 결과","8. B1·B2: 후속검색과 탐색지속성 결과","9. B3: 서로 다른 세 가지 회복 지표","10. 세션 결과 4개 세그먼트","11. H3: 재검색 전이 유형 결과","12. 해석 한계와 다음 단계"]
assert top==expected
assert len(d.tables)==18 and len(d.inline_shapes)==6
assert sum(p.text.startswith("표 해설 |") for p in d.paragraphs)==18
assert sum(p.text.startswith("그래프 해설 |") for p in d.paragraphs)==6
all_text="\n".join([p.text for p in d.paragraphs]+[c.text for t in d.tables for row in t.rows for c in row.cells])
assert all(x in all_text for x in ["각 원본 세션을 23회씩","random_seed=20260903","SEARCH_RESULT","198,128","3.15 구현용 의사코드"])
with zipfile.ZipFile(docx) as z:assert z.testzip() is None and len([x for x in z.namelist() if x.startswith("word/media/")])==6
text=log.read_text(encoding="utf-8")+f"\n## 증강 우선 구성 최종 확인\n\n- 목차: 핵심 12개 장\n- 순서: 증강 목적·방법·재현 절차·테이블·QA(1~5절) → 1,000명 분석 결과(6~11절) → 한계·다음 단계(12절)\n- 표 해설: 18개 / 그래프 해설: 6개\n- DOCX 재열기·ZIP·OOXML·이미지 관계 검사: PASS\n- 기존 단일 최종 보고서만 수정; 생성기 및 합성 데이터 재생성 미실행\n- 최종 보고서 SHA-256: `{sh(docx)}`\n"
log.write_text(text,encoding="utf-8")
m=json.loads(manifest.read_text(encoding="utf-8"));m["report_docx"]["path"]=str(docx.resolve().relative_to(repo.resolve())).replace("\\","/");m["report_docx"]["size_bytes"]=docx.stat().st_size;m["report_docx"]["sha256"]=sh(docx);m["packaging_execution_log"]["size_bytes"]=log.stat().st_size;m["packaging_execution_log"]["sha256"]=sh(log)
a=m["report_content_addendum"];a["top_level_toc_count"]=12;a["content_order"]="augmentation_method_first_then_S0_1000_analysis_results";a["table_explanation_count"]=18;a["chart_explanation_count"]=6;a["single_final_report_path"]=m["report_docx"]["path"]
manifest.write_text(json.dumps(m,ensure_ascii=False,indent=2),encoding="utf-8")
print(json.dumps({"report_sha256":sh(docx),"log_sha256":sh(log),"manifest_sha256":sh(manifest),"top_level_chapters":12,"table_explanations":18,"chart_explanations":6,"generator_executed":False},ensure_ascii=True,indent=2))
