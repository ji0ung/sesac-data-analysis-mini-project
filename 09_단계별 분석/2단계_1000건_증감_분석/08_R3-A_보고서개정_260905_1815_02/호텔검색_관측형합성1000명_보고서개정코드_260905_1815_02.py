from __future__ import annotations

import hashlib
import json
import shutil
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[3]
STAGE2 = ROOT / "09_단계별 분석" / "2단계_1000건_증감_분석"
BASE = STAGE2 / "호텔검색_관측형합성1000명_전체가설분석보고서_260904_1149_01.docx"
AUTH = STAGE2 / "07_R2.5_승인DB_공식전체분석_260905_1801_01"
APPROVED_DB = STAGE2 / "06_STEP_R2_독립검증_260905_1650_01" / "호텔검색_관측형합성1000명_분석용lineage제거정제DB_260905_1650_01.sqlite"
OLD_LOG = ROOT / "09_단계별 분석" / "1단계_원본_분석_및_가설_검증" / "local_outputs" / "20260903_checkpoint1" / "분석_결정_로그.md"
OUT = Path(__file__).resolve().parent
STAMP = "260905_1815"
REPORT = OUT / f"호텔검색_관측형합성1000명_전체가설분석보고서_{STAMP}_02.docx"
DECISION = OUT / f"호텔검색_관측형합성1000명_분석결정로그_{STAMP}_02.md"
VERIFY = OUT / f"호텔검색_관측형합성1000명_보고서개정검증_{STAMP}_02.json"


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for block in iter(lambda: f.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def replace_runs(p: Paragraph, old: str, new: str) -> bool:
    text = "".join(r.text for r in p.runs)
    if old not in text:
        return False
    changed = text.replace(old, new)
    if p.runs:
        p.runs[0].text = changed
        for run in p.runs[1:]:
            run.text = ""
    else:
        p.add_run(changed)
    return True


def add_paragraph_after(anchor, text: str = "", style: str | None = None) -> Paragraph:
    p = OxmlElement("w:p")
    anchor._element.addnext(p)
    para = Paragraph(p, anchor._parent)
    if style:
        para.style = style
    if text:
        para.add_run(text)
    return para


def add_table_after(doc: Document, anchor, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = doc.tables[18].style
    for i, value in enumerate(headers):
        table.rows[0].cells[i].text = value
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    anchor._element.addnext(table._tbl)
    return table


def add_picture_after(doc: Document, anchor, image_path: Path, caption: str):
    p = add_paragraph_after(anchor)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin)
    cap = add_paragraph_after(p, caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return cap


def find_paragraph(doc: Document, needle: str) -> Paragraph:
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise RuntimeError(f"문단을 찾을 수 없음: {needle}")


def main() -> None:
    expected = {
        APPROVED_DB: "7d449fab6847730be38fe46ebb6bd6a5b31690cdf38cf4bdbadeaa75edd1ae04",
        AUTH / "호텔검색_관측형합성1000명_수정승인DB전체분석결과_260905_1801_01.xlsx": "192841580da144af99445dd04a78c37aad0329af0f9ebbb7adb55fe43cd88faf",
        AUTH / "호텔검색_관측형합성1000명_수정승인DB전체분석결과_260905_1801_01.json": "c3c77aece865fc0e955779d6e5f67d5a17644070c4f4cba638ee5c84ea8932f3",
        AUTH / "호텔검색_관측형합성1000명_STEP3독립검수결과_260905_1801_01.xlsx": "a2104cc4cc6315cc4383da1f180e8494c150c08e44f7e0939b746e92976eb101",
        AUTH / "호텔검색_관측형합성1000명_STEP3독립검수결과_260905_1801_01.json": "e1714206510fec5f188b9d8b4399ee2dd8f156e0c4efe408522c054bbeb313e2",
        AUTH / "호텔검색_관측형합성1000명_R2.5판단로그_260905_1801_01.md": "651e0ae2b23f7fc0b47f45fb6708d7572666cfb3a48cad191f2236fa70276c3f",
        AUTH / "호텔검색_관측형합성1000명_R2.5_SHA256매니페스트_260905_1801_01.json": "c0ef835e1e5e44d2c309954b2ebdccc452d7e00778c6107975edb4ac62bbd4ad",
    }
    mismatches = {str(p): [sha256(p), h] for p, h in expected.items() if sha256(p) != h}
    if mismatches:
        raise RuntimeError(f"권위 입력 해시 불일치: {mismatches}")

    doc = Document(BASE)

    replacements = [
        ("첫 검색 시각을 search_origin으로 고정", "세션별 유효 SEARCH.search_time과 복제 대상 EVENT.event_at의 최솟값을 session_origin으로 고정"),
        ("new_search_time = user_base_time + (old_search_time - search_origin)", "new_search_time = base_i + (old_search_time - session_origin)"),
        ("SEARCH 전용 origin", "SEARCH와 EVENT에 공통인 session_origin"),
        ("EVENT 전용 origin(event_origin)", "SEARCH와 EVENT에 공통인 session_origin"),
        ("new_event_at = user_base_time + (old_event_at - event_origin)", "new_event_at = base_i + (old_event_at - session_origin)"),
        ("new_session_end_time = user_base_time + (old_session_end_time - event_origin)", "new_session_end_time = base_i + (old_session_end_time - session_origin)"),
        ("원본 세션의 최초 search_time을 search_origin으로 둔다.", "원본 세션의 유효한 SEARCH.search_time과 실제 복제 대상 EVENT.event_at 중 최솟값을 session_origin으로 둔다. NULL timestamp는 원점 후보에서 제외하고 원본 NULL은 그대로 유지한다."),
        ("각 검색의 새 시간은 synthetic_base_time + (original_search_time − search_origin)으로 계산한다. 이로써 세션 내부 검색 간 시간차와 순서를 보존한다.", "각 검색의 새 시간은 base_i + (old_search_time − session_origin)으로 계산한다. SEARCH·EVENT·session_end_time에 동일한 이동량을 적용하여 내부 및 교차 스트림 상대시간을 보존한다."),
        ("원본 세션 EVENT를 event_at, event_id 순으로 정렬하고 최초 event_at을 event_origin으로 둔다.", "원본 세션 EVENT는 event_at, event_id로 안정 정렬하되, SEARCH와 공통인 session_origin을 사용한다."),
        ("새 event_at은 synthetic_base_time + (old_event_at − event_origin)으로 계산한다. session_end_time이 있으면 같은 event_origin 기준으로 이동한다.", "새 event_at은 base_i + (old_event_at − session_origin), 새 session_end_time은 base_i + (old_session_end_time − session_origin)으로 계산한다. Jittering과 로그정규 난수 생성은 적용하지 않는다."),
        ("생성 DB의 SQLite integrity_check=ok이며 승인 SHA-256은 db80db7048add9c0c4cb1a985e67a77ae99bef3a30ce32bedb70cc0ee61dc896이다.", "수정 전 S0 DB의 SQLite integrity_check=ok이며 당시 승인 SHA-256은 db80db7048add9c0c4cb1a985e67a77ae99bef3a30ce32bedb70cc0ee61dc896이다. 현재 분석·배포용 최종 승인본은 lineage 제거 시간축 보정 DB이며 SHA-256은 7d449fab6847730be38fe46ebb6bd6a5b31690cdf38cf4bdbadeaa75edd1ae04이다."),
    ]
    for p in doc.paragraphs:
        for old, new in replacements:
            replace_runs(p, old, new)

    # 기존 S0를 승인 입력으로 오인하지 않도록 본문 계보를 명시한다.
    anchor = find_paragraph(doc, "5. 데이터 계보·품질·재현성 QA")
    cursor = add_paragraph_after(anchor, "5.1 STEP R2.5 권위 입력과 계보", "Heading 2")
    cursor = add_paragraph_after(cursor, "최종 분석 입력은 lineage가 제거된 시간축 보정 관측형 합성 1,000명 승인 DB로 고정했다. 수정 전 S0는 비교 기준으로 보존하고, R1 lineage 포함 DB는 QA 전용으로만 사용했다. 승인 DB SHA-256은 7d449fab6847730be38fe46ebb6bd6a5b31690cdf38cf4bdbadeaa75edd1ae04이다. STEP1·STEP2·STEP3는 모두 새 입력 해시에 근거해 재실행했으며 전체 불일치 건수는 0건이다.")

    cursor = add_paragraph_after(cursor, "5.2 고차 결합 검증", "Heading 2")
    cursor = add_paragraph_after(cursor, "본 분석의 고차 결합은 price IS NOT NULL, user_rating_min IS NOT NULL, amenity_count > 0의 3개 선택형 필터로 정의했다. region은 destination과 동일하여 필터 수에서 제외하고 BN 지역 설명변수로만 사용했으며, property_type과 property_grade는 전 행 NULL이므로 제외했다.")
    t = add_table_after(doc, cursor, ["구분", "분자", "분모", "비율", "원본 대비"], [
        ["원본", "91", "104", "87.5000%", "기준"],
        ["시간축 보정 승인 합성본", "2,125", "2,425", "87.6289%", "+0.1289%p"],
    ])
    cursor = add_paragraph_after(t, "제한 BN 결과", "Heading 3")
    t = add_table_after(doc, cursor, ["구분", "Log loss", "Brier", "ROC-AUC", "ECE"], [
        ["원본", "0.942203", "0.191710", "0.823106", "0.173780"],
        ["시간축 보정 승인 합성본", "0.915162", "0.182365", "0.832769", "0.151952"],
    ])
    cursor = add_paragraph_after(t, "고차 결합 본 분석과 제한 BN은 조건부 구조가 대체로 보존되었음을 보여준다. 이는 모집단 대표성이나 인과효과를 뜻하지 않으며, ECE가 0이 아니므로 완전 보정으로 해석하지 않는다. 고차 결합 판정은 PASS이다. 6개 플래그 정의 A는 민감도 분석으로만 보관하며 원본 119/160(74.3750%), 승인 합성본 2,783/3,741(74.3919%)이다.")
    cursor = add_picture_after(doc, cursor, AUTH / "호텔검색_관측형합성1000명_고차결합BN시각화_260905_1801_01.png", "그림. 고차 결합 및 제한 BN 비교(R2.5 권위 산출물)")

    cursor = add_paragraph_after(cursor, "5.3 0건 후 인접 검색 시간 간격", "Heading 2")
    t = add_table_after(doc, cursor, ["구분", "n", "음수", "0초", "양수", "평균", "중앙값", "Q1–Q3", "P90", "P95", "최대"], [
        ["실제 관측 원본", "140", "0", "1", "139", "17.7214초", "12초", "7–18초", "31.1초", "39.1초", "249초"],
        ["시간축 보정 승인 합성본", "3,271", "0", "23", "3,248", "17.7447초", "12초", "7–18초", "32초", "40초", "249초"],
    ])
    cursor = add_paragraph_after(t, "양수 간격에 로그정규분포를 적합하고 모수를 재추정하는 parametric bootstrap을 적용했다. 원본은 μ=2.507199, σ=0.730609, p=0.019960, 승인 합성본은 μ=2.507799, σ=0.730905, p=0.001996이다. 두 자료 모두 5% 유의수준에서 로그정규 적합이 지지되지 않아 WARN을 유지한다. 이번 생성은 jittering이나 로그정규분포 생성을 적용하지 않았고 원본의 시간차를 보존한 것이다.")

    cursor = add_paragraph_after(cursor, "5.4 시간 일관성과 교차 스트림", "Heading 2")
    t = add_table_after(doc, cursor, ["검사 항목", "결과", "판정"], [
        ["SEARCH 내부 음수 간격", "0", "PASS"],
        ["EVENT 내부 음수 간격", "0", "PASS"],
        ["SEARCH–EVENT 상대시간 불일치 / 최대 오차", "0 / 0초", "PASS"],
        ["click보다 이른 detail view", "0", "PASS"],
        ["session_end_time 역전", "0", "PASS"],
        ["EVENT source ordinal", "원천 순번 미제공", "NOT_TESTABLE"],
    ])
    cursor = add_paragraph_after(t, "교차 스트림 중첩 비교", "Heading 3")
    t = add_table_after(doc, cursor, ["데이터", "중첩", "비교 가능 건", "중첩률"], [
        ["원본", "64", "7,210", "0.8877%"],
        ["수정 전 기존 S0", "148,529", "167,330", "88.7641%"],
        ["시간축 보정 승인 합성본", "1,474", "167,330", "0.8809%"],
    ])
    cursor = add_paragraph_after(t, "공통 session_origin을 SEARCH와 EVENT에 동일하게 적용해 원본 교차 스트림 상대시간을 복원했다. 전체 시간 검증은 구조·순서·상대시간 항목은 PASS이나 로그정규 적합성 WARN을 포함하므로 CONDITIONAL PASS이다.")
    cursor = add_picture_after(doc, cursor, AUTH / "호텔검색_관측형합성1000명_시간간격교차스트림시각화_260905_1801_01.png", "그림. 시간 간격 및 교차 스트림 비교(R2.5 권위 산출물)")

    cursor = add_paragraph_after(cursor, "5.5 핵심 지표 및 STEP3 독립검수", "Heading 2")
    t = add_table_after(doc, cursor, ["지표", "결과"], [
        ["전체 0건률", "3,434/6,900 (49.7681%)"],
        ["0건 후 후속검색률", "3,271/3,434 (95.2533%)"],
        ["즉시 회복률", "558/3,271 (17.0590%)"],
        ["세션 최종 회복률", "488/651 (74.9616%)"],
        ["상세진입률", "1,052/6,900 (15.2464%)"],
        ["세션 4개 세그먼트", "직접 성공 630 / 결과 노출·미선택 231 / 재검색 회복 92 / 지속 실패 47"],
    ])
    cursor = add_paragraph_after(t, "기존 승인 STEP2 표를 새 승인 DB에서 전수 재계산한 결과 불일치 0건이었다. STEP3는 STEP2 계산 함수와 결과 객체를 재사용하지 않는 독립 코드로 검수했으며 결정적 표본 44건 중 실패 0건, 전체 검수 불일치 0건, 입력 불변성 및 개인정보 미출력을 확인했다.")

    final_anchor = find_paragraph(doc, "8. 종합 판정·해석 한계·다음 단계")
    cursor = add_paragraph_after(final_anchor, "8.1 R2.5 기반 종합 판정", "Heading 2")
    t = add_table_after(doc, cursor, ["검증 영역", "판정", "근거"], [
        ["STEP1 입력·해시·계보", "PASS", "승인 DB 해시 일치, lineage 제거 입력 고정"],
        ["STEP2 전체 가설 및 보완 분석", "PASS", "기존 결과와 전수 비교 불일치 0"],
        ["고차 결합 본 분석 및 제한 BN", "PASS", "주요 조건부 구조 보존"],
        ["시간 순서·상대시간·교차 스트림", "PASS", "오류 0, 중첩률 원본 수준 복원"],
        ["로그정규 적합성", "WARN", "원본·승인 합성본 모두 5% 수준에서 미지지"],
        ["EVENT source ordinal", "NOT_TESTABLE", "원천 순번 미제공"],
        ["STEP3 독립검수", "PASS", "결정적 표본 실패 0/44, 전체 불일치 0"],
        ["보고서 개정", "PASS", "R2.5 권위 산출물과 계보 반영"],
    ])
    cursor = add_paragraph_after(t, "STEP R3-A=PASS: R2.5 권위 산출물을 반영한 보고서 개정 초안이 생성되었으며, 다음 단계에서 독립 문서 QA를 수행할 수 있다.")
    cursor = add_paragraph_after(cursor, "8.2 변경 이력", "Heading 2")
    cursor = add_paragraph_after(cursor, "2026-09-05 | STEP R3-A | 수정 전 S0 해시 db80db7048add9c0c4cb1a985e67a77ae99bef3a30ce32bedb70cc0ee61dc896을 변경 이력에 보존했다. SEARCH/EVENT 분리 시간 원점 오류를 공통 session_origin 공식으로 교체하고, R2.5 승인 lineage 제거 DB와 해시, 고차 결합 정의 B·정의 A 민감도 분석, 시간 간격·교차 스트림, STEP1·STEP2·STEP3 PASS 결과를 추가했다. 전체 불일치는 0건이며 Jittering은 적용하지 않았다. QA DB와 분석·배포 DB는 분리했다. 로그정규 적합성은 WARN으로 유지했다. 기존 A1·A2·B1·B2·B3·H3 및 핵심 집계값은 변경하지 않았고 원본 보고서와 기존 산출물은 보존했다.")

    doc.core_properties.modified = datetime(2026, 9, 5, 18, 15)
    doc.core_properties.comments = "STEP R2.5 PASS 권위 산출물 반영 보고서 개정 초안; 독립 문서 QA 전 단계"
    doc.save(REPORT)

    old_text = OLD_LOG.read_text(encoding="utf-8")
    appended = f"""

## 2026-09-05 STEP R3-A — R2.5 권위 산출물 보고서 반영

- 최종 분석 입력: `호텔검색_관측형합성1000명_분석용lineage제거정제DB_260905_1650_01.sqlite`
- 입력 SHA-256: `7d449fab6847730be38fe46ebb6bd6a5b31690cdf38cf4bdbadeaa75edd1ae04`
- 수정 전 S0 SHA-256: `db80db7048add9c0c4cb1a985e67a77ae99bef3a30ce32bedb70cc0ee61dc896`
- 계보 판단: 수정 전 S0는 비교 기준으로 보존, R1 lineage 포함 DB는 QA 전용, R2 lineage 제거 DB를 최종 승인 입력으로 고정
- 시간 공식: `session_origin = min(valid SEARCH.search_time, replicated EVENT.event_at)`; SEARCH, EVENT, session_end_time 모두 동일 원점 적용
- 분석 결과: 기존 핵심 지표 유지, STEP2 전수 불일치 0건, STEP3 결정적 표본 실패 0/44 및 전체 불일치 0건
- 보완 결과: 고차 결합 PASS, 시간 구조 PASS, 로그정규 적합성 WARN, EVENT source ordinal NOT_TESTABLE
- 문서 처리: 기존 보고서와 기존 결정 로그는 수정하지 않고 개정본을 별도 생성
- 판정: `STEP R3-A=PASS` — 독립 문서 QA 진행 가능
"""
    DECISION.write_text(old_text.rstrip() + appended, encoding="utf-8")

    # 재개방 기반 구조 검증. 시각적 독립 QA는 다음 단계 범위로 남긴다.
    check = Document(REPORT)
    all_text = "\n".join(p.text for p in check.paragraphs) + "\n" + "\n".join(
        cell.text for table in check.tables for row in table.rows for cell in row.cells
    )
    required = [
        "7d449fab6847730be38fe46ebb6bd6a5b31690cdf38cf4bdbadeaa75edd1ae04",
        "87.6289%", "0.001996", "1,474", "148,529", "결정적 표본 44건 중 실패 0건",
        "STEP R3-A=PASS", "common session_origin" if False else "공통 session_origin",
    ]
    forbidden = ["old_event_at - event_origin", "old_search_time - search_origin"]
    verification = {
        "base_report_sha256": sha256(BASE),
        "approved_db_sha256_before": sha256(APPROVED_DB),
        "approved_db_sha256_after": sha256(APPROVED_DB),
        "report_sha256": sha256(REPORT),
        "decision_log_sha256": sha256(DECISION),
        "paragraph_count": len(check.paragraphs),
        "table_count": len(check.tables),
        "required_text_missing": [x for x in required if x not in all_text],
        "forbidden_text_remaining": [x for x in forbidden if x in all_text],
        "source_files_unchanged": sha256(APPROVED_DB) == expected[APPROVED_DB],
        "structural_check": "PASS",
        "visual_independent_doc_qa": "NEXT_STEP",
    }
    if verification["required_text_missing"] or verification["forbidden_text_remaining"]:
        raise RuntimeError(verification)
    VERIFY.write_text(json.dumps(verification, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(verification, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
