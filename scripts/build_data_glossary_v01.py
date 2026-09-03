from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path('/Users/macuser/sesac-data-analysis-mini-project')
OUT=ROOT/'06_guides_and_prompts'/'호텔검색_데이터분석_쉬운용어사전_표본설계가중치_20260902_v01.docx'
BLUE='2E74B5'; NAVY='0B2545'; PALE='E8EEF5'; LIGHT='F4F6F9'; GRAY='666666'

def font(run,size=11,bold=False,color='000000'):
    run.font.name='Calibri'; run.font.size=Pt(size); run.bold=bold; run.font.color.rgb=RGBColor.from_string(color)
    rpr=run._element.get_or_add_rPr(); fonts=rpr.get_or_add_rFonts(); fonts.set(qn('w:ascii'),'Calibri'); fonts.set(qn('w:hAnsi'),'Calibri'); fonts.set(qn('w:eastAsia'),'맑은 고딕')

def shade(cell,fill):
    shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); cell._tc.get_or_add_tcPr().append(shd)

def table_geom(t,widths):
    t.autofit=False; t.alignment=WD_TABLE_ALIGNMENT.LEFT; pr=t._tbl.tblPr
    for tag,val in [('tblW',9360),('tblInd',120)]:
        x=OxmlElement('w:'+tag); x.set(qn('w:w'),str(val)); x.set(qn('w:type'),'dxa'); pr.append(x)
    grid=t._tbl.tblGrid
    for x in list(grid): grid.remove(x)
    for w in widths:
        x=OxmlElement('w:gridCol'); x.set(qn('w:w'),str(w)); grid.append(x)
    for row in t.rows:
        for i,c in enumerate(row.cells):
            c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tcw=c._tc.get_or_add_tcPr().get_or_add_tcW(); tcw.set(qn('w:w'),str(widths[i])); tcw.set(qn('w:type'),'dxa')
            mar=OxmlElement('w:tcMar')
            for tag,val in [('top',90),('bottom',90),('start',120),('end',120)]:
                x=OxmlElement('w:'+tag); x.set(qn('w:w'),str(val)); x.set(qn('w:type'),'dxa'); mar.append(x)
            c._tc.get_or_add_tcPr().append(mar)

def add_table(doc,heads,rows,widths):
    t=doc.add_table(rows=1,cols=len(heads)); t.style='Table Grid'
    for i,h in enumerate(heads):
        shade(t.rows[0].cells[i],PALE); p=t.rows[0].cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run(h),9,True,NAVY)
    hdr=OxmlElement('w:tblHeader'); hdr.set(qn('w:val'),'true'); t.rows[0]._tr.get_or_add_trPr().append(hdr)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row):
            if ri%2: shade(cells[i],'FAFBFC')
            p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.08
            font(p.add_run(v),8.7, i==0, NAVY if i==0 else '222222')
    table_geom(t,widths); doc.add_paragraph()

def callout(doc,label,text):
    t=doc.add_table(rows=1,cols=1); c=t.cell(0,0); shade(c,LIGHT)
    p=c.paragraphs[0]; font(p.add_run(label+'  '),10.5,True,NAVY); font(p.add_run(text),10.5)
    table_geom(t,[9360]); doc.add_paragraph()

d=Document(); s=d.sections[0]; s.page_width=Inches(8.5); s.page_height=Inches(11); s.top_margin=s.bottom_margin=s.left_margin=s.right_margin=Inches(1); s.header_distance=s.footer_distance=Inches(.492)
normal=d.styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for name,size,before,after in [('Heading 1',16,18,10),('Heading 2',13,14,7)]:
    st=d.styles[name]; st.font.name='Calibri'; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(BLUE); st._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after)
font(s.header.paragraphs[0].add_run('TEAM 2 · DATA WORDBOOK'),8.5,True,GRAY)
s.footer.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(s.footer.paragraphs[0].add_run('호텔검색 데이터분석 쉬운 용어사전 · v01'),8.5,False,GRAY)
p=d.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4); font(p.add_run('호텔검색 데이터분석 쉬운 용어사전'),22,True,NAVY)
p=d.add_paragraph(); p.paragraph_format.space_after=Pt(14); font(p.add_run('표본설계·과표집·가중치를 발표 가능한 말로 풀어쓴 단어장'),13,False,BLUE)
callout(d,'30초 요약','관측형은 실제 비율을 닮게 만든 기본 세트, 과표집형은 드문 사례를 더 많이 넣은 별도 세트, 가중치는 과표집된 비율을 실제 비율로 되돌리는 계산값이다.')

d.add_heading('1. 꼭 알아야 할 핵심 용어',level=1)
rows=[
('모집단','우리가 결론을 적용하고 싶은 전체 대상','호텔검색 서비스의 전체 사용자','“전체 고객은…”이라고 말할 때의 전체'),
('표본','모집단에서 실제로 관찰하거나 뽑은 일부','검색 사용자 68명 또는 파일럿 1천 명','표본 결과를 전체처럼 말하지 않기'),
('표본설계','누구를 몇 명, 어떤 기준으로 뽑을지 정한 규칙','지역×의도×세그먼트별 생성 인원표','데이터 생성 전에 먼저 확정'),
('세그먼트','비슷한 행동을 보이는 사용자 묶음','직접 성공, 재검색 회복, 지속 실패','한 사용자에게 결과 세그먼트 1개'),
('관측비율','원본 데이터에서 실제로 관찰된 비율','실패군 12/68=17.6%','기준 분포로 사용'),
('관측형','관측비율을 최대한 그대로 유지한 생성 세트','1천 명 중 실패군 약 176명','발표의 기본 데이터 세트'),
('과표집','드문 집단을 분석하려고 실제 비율보다 많이 뽑는 것','실패군을 17.6% 대신 35%로 생성','관측형과 섞지 않고 별도 보관'),
('가중치','많이/적게 뽑힌 표본을 실제 비율로 되돌리는 값','17.6%÷35%=0.503','과표집 결과를 실제 비율로 해석할 때 적용'),
('층·층화','지역·의도처럼 표본을 나누는 기준과 그 묶음','Tokyo×PRICE, Osaka×AMENITY','각 층의 수가 너무 작지 않은지 확인'),
('추출확률','어떤 대상이 표본에 포함될 확률','실패군을 더 뽑으면 실패군 추출확률이 커짐','가중치 계산의 근거'),
('파일럿','본 작업 전 작은 규모로 시험하는 데이터','1천 명 먼저 생성','오류 확인 후 1만으로 확장'),
('허용오차','목표 비율과 생성 결과가 달라도 통과시키는 범위','핵심 ±3%p, 기타 ±5%p','교수님 승인 후 QA 기준으로 사용'),
]
add_table(d,['용어','쉬운 뜻','호텔검색 예시','기억할 점'],rows,[1400,2700,2860,2400])

d.add_heading('2. 관측형과 과표집형 비교',level=1)
add_table(d,['구분','관측형','과표집형'],[
('목적','실제 분포와 비슷한 결과·발표','희귀 실패경로를 자세히 분석'),
('실패군 비율','17.6% 유지','예: 35%로 확대'),
('1천 명 예시','실패군 약 176명','실패군 350명'),
('가중치','보통 1.0','실패군 예시 0.503'),
('보관','본 세트','별도 세트'),
('주의','기준 데이터로 사용','가중치 없이 실제 비율처럼 발표 금지'),
],[1500,3930,3930])
callout(d,'이번 파일럿','1천 명을 관측형으로 만들면 실패군은 약 176명이다. 지속 실패와 0건 즉시 이탈을 원본 비율대로 나누면 약 88명씩이므로 “각 유형 최소 50명”을 별도 과표집 없이 충족한다.')

d.add_heading('3. 표본설계·가중치 컬럼',level=1)
add_table(d,['컬럼','저장 내용','예시'],[
('sample_design_id','표본설계 버전','OBS_V01'),
('sample_set_type','관측형/과표집형 구분','OBSERVED / OVERSAMPLED'),
('sample_stratum','지역×의도×세그먼트 층','Tokyo×PRICE×재검색회복'),
('target_population_share','원본에서 관측된 목표 비율','0.176'),
('sample_share','생성 세트에서 차지하는 비율','0.350'),
('selection_probability','해당 층의 추출확률','설계표에서 계산'),
('sample_weight','목표비율÷생성비율','0.176÷0.350=0.503'),
('is_oversampled','과표집 여부','0 / 1'),
('weight_version','가중치 계산 버전','W_V01'),
],[2350,4000,3010])

d.add_heading('4. 의도 코드(intent_code)',level=1)
p=d.add_paragraph('의도 코드는 원본 DB의 기존 컬럼명이 아니라 query_text, destination, sort_option, SEARCH_FILTER 조건을 이용해 새로 만드는 분석용 분류값이다. A2는 이 코드와 지역별 표본 수를 먼저 정의한 뒤 진행한다.'); p.paragraph_format.space_after=Pt(8)
add_table(d,['코드','뜻','판정에 사용하는 값'],[
('LOCATION_ONLY','지역 중심','destination/region 중심, 강한 추가필터 없음'),
('PRICE','가격 중심','price 설정 또는 가격정렬'),
('AMENITY','편의시설 중심','amenity_count 및 편의시설 조건'),
('ACCOMMODATION_TYPE','숙소유형 중심','property_type'),
('HOTEL_NAME','특정 호텔 중심','query_text가 호텔명 사전과 일치/유사'),
('MIXED','둘 이상 혼합','의도 조건이 복수로 동시에 성립'),
('TYPO_INCOMPLETE','오타·중간입력','A3/SX 별도 실험군'),
('UNKNOWN','판별 불가','필요 값 부족 또는 규칙 불일치'),
],[2450,2400,4510])
callout(d,'C3 주의','QUALITY 코드는 기술적으로 만들 수 있지만 품질 우선형 C3는 현재 제외 상태다. 따라서 이번 핵심 증강 확률에는 사용하지 않고 필요하면 기술통계만 남긴다.')

d.add_heading('5. 기준 DB 확인 결과',level=1)
p=d.add_paragraph(); font(p.add_run('기준 파일: '),10.5,True); font(p.add_run('03_data_modeling/travel_data_filtered_complete_2026-09-01_v01_원본.sqlite'),10.5)
add_table(d,['테이블','행 수'],[('USER','89'),('SEARCH','296'),('SEARCH_FILTER','296'),('SEARCH_RESULT','8,555'),('EVENT','10,432'),('BOOKING','36')],[4680,4680])
callout(d,'버전 주의','현재 원본 SQLite의 SEARCH는 296건이다. 가설문서 v06의 608건과 다르므로 A2 증량계획을 확정하기 전에 어느 데이터가 최종 분석 기준인지 반드시 통일해야 한다.')

d.add_heading('6. 발표할 때 쓰는 쉬운 문장',level=1)
for text in [
    '“관측형은 실제 데이터 비율을 유지한 기본 데이터입니다.”',
    '“과표집형은 드문 실패 유형을 더 자세히 보기 위한 별도 데이터입니다.”',
    '“과표집 결과는 가중치를 적용하지 않으면 실제 비율처럼 해석할 수 없습니다.”',
    '“A2는 승인됐지만 지역과 검색 의도를 몇 건씩 만들지 먼저 정의해야 합니다.”',
    '“현재 SQLite는 검색 296건이고 가설문서는 608건이므로 기준 버전 확인이 선행됩니다.”']:
    p=d.add_paragraph(style='List Bullet'); font(p.add_run(text),10.5); p.paragraph_format.space_after=Pt(5)

d.core_properties.title='호텔검색 데이터분석 쉬운 용어사전'; d.core_properties.author='2팀'
OUT.parent.mkdir(parents=True,exist_ok=True); d.save(OUT); print(OUT)
