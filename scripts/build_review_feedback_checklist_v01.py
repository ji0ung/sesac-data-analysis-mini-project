from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path('/Users/macuser/sesac-data-analysis-mini-project')
OUT=ROOT/'04_analysis_design'/'32일차_제출검토의견_반영체크리스트_296건기준_20260902_v02_통합완성본.docx'
BLUE='2E74B5'; NAVY='0B2545'; PALE='E8EEF5'; LIGHT='F4F6F9'; GREEN='E2F0D9'; GOLD='FFF4E5'; RED='FCE8E6'; GRAY='666666'

def font(r,size=10.5,bold=False,color='000000'):
    r.font.name='Calibri'; r.font.size=Pt(size); r.bold=bold; r.font.color.rgb=RGBColor.from_string(color)
    f=r._element.get_or_add_rPr().get_or_add_rFonts(); f.set(qn('w:ascii'),'Calibri'); f.set(qn('w:hAnsi'),'Calibri'); f.set(qn('w:eastAsia'),'맑은 고딕')
def shade(c,fill):
    x=OxmlElement('w:shd'); x.set(qn('w:fill'),fill); c._tc.get_or_add_tcPr().append(x)
def geom(t,widths):
    t.autofit=False; t.alignment=WD_TABLE_ALIGNMENT.LEFT; pr=t._tbl.tblPr
    for tag,val in [('tblW',9360),('tblInd',120)]:
        x=OxmlElement('w:'+tag); x.set(qn('w:w'),str(val)); x.set(qn('w:type'),'dxa'); pr.append(x)
    g=t._tbl.tblGrid
    for x in list(g):g.remove(x)
    for w in widths:
        x=OxmlElement('w:gridCol'); x.set(qn('w:w'),str(w)); g.append(x)
    for row in t.rows:
        for i,c in enumerate(row.cells):
            c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tcw=c._tc.get_or_add_tcPr().get_or_add_tcW(); tcw.set(qn('w:w'),str(widths[i])); tcw.set(qn('w:type'),'dxa')
            mar=OxmlElement('w:tcMar')
            for tag,val in [('top',80),('bottom',80),('start',110),('end',110)]:
                x=OxmlElement('w:'+tag); x.set(qn('w:w'),str(val)); x.set(qn('w:type'),'dxa'); mar.append(x)
            c._tc.get_or_add_tcPr().append(mar)
def table(d,heads,rows,widths,size=7.9):
    t=d.add_table(rows=1,cols=len(heads)); t.style='Table Grid'
    for i,h in enumerate(heads):
        shade(t.rows[0].cells[i],PALE); p=t.rows[0].cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run(h),size,True,NAVY)
    rep=OxmlElement('w:tblHeader'); rep.set(qn('w:val'),'true'); t.rows[0]._tr.get_or_add_trPr().append(rep)
    for ri,row in enumerate(rows):
        cs=t.add_row().cells
        status=str(row[0])
        fill=GREEN if '완료' in status else GOLD if '부분' in status or '판단' in status else RED
        for i,v in enumerate(row):
            if i==0: shade(cs[i],fill)
            elif ri%2: shade(cs[i],'FAFBFC')
            p=cs[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.08; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if i==0 else WD_ALIGN_PARAGRAPH.LEFT
            font(p.add_run(str(v)),size,i==0,NAVY if i==0 else '222222')
    geom(t,widths); d.add_paragraph()
def callout(d,label,text,fill=LIGHT):
    t=d.add_table(rows=1,cols=1); c=t.cell(0,0); shade(c,fill); p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(label+'  '),10.3,True,NAVY); font(p.add_run(text),10.3); geom(t,[9360]); d.add_paragraph()

d=Document(); s=d.sections[0]; s.page_width=Inches(8.5); s.page_height=Inches(11); s.top_margin=s.bottom_margin=s.left_margin=s.right_margin=Inches(1); s.header_distance=s.footer_distance=Inches(.492)
n=d.styles['Normal']; n.font.name='Calibri'; n.font.size=Pt(11); n._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.25
for name,size,bef,aft in [('Heading 1',16,18,10),('Heading 2',13,14,7)]:
    st=d.styles[name]; st.font.name='Calibri'; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(BLUE); st._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); st.paragraph_format.space_before=Pt(bef); st.paragraph_format.space_after=Pt(aft)
font(s.header.paragraphs[0].add_run('TEAM 2 · REVIEW RESPONSE CHECKLIST'),8.5,True,GRAY); s.footer.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(s.footer.paragraphs[0].add_run('296건 기준 · 검토의견 반영 확인 v02 통합완성본'),8.5,False,GRAY)
p=d.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4); font(p.add_run('제출 검토의견 반영 체크리스트'),22,True,NAVY)
p=d.add_paragraph(); p.paragraph_format.space_after=Pt(12); font(p.add_run('교수 첨삭 × 가설문서 v08 통합완성본 × 데이터증강계획서 v06 통합완성본'),13,False,BLUE)
for a,b in [('검토 원문','BI시각화_32일차_2팀_제출본_검토의견_260902_1620_01.docx'),('수정 기준','travel_data_filtered_complete_2026-09-01_v01_원본.sqlite · SEARCH 296건'),('검토 대상','호텔검색 가설문서 v08 통합완성본 / 데이터증강계획서 v06 통합완성본'),('판정일','2026-09-02')]:
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(2); font(p.add_run(a+': '),10.3,True); font(p.add_run(b),10.3)
d.add_paragraph()
callout(d,'종합 판정','핵심 수정사항 14개 중 완료 11개, 부분 반영·추가 판단 필요 3개다. 미완료는 숨기지 않고 오타 정의, 희소 표본, 0건 즉시 이탈 표본 문제로 분리했다.',GREEN)

d.add_heading('1. 반드시 확인할 것 — 4건',level=1)
table(d,['상태','검토의견','반영 내용','근거 위치','남은 조치'],[
('☑ 완료','H1 판정이 부분 채택/채택으로 엇갈림','A1을 전 구간 “부분 채택”으로 통일','가설 v08 3.1·11·15절','없음'),
('△ 부분 반영','오타형 제외 후 H1 재계산','DB에 typo_flag 없음. 빈 query_text는 정상 destination 검색 가능, error 이벤트 5건도 모두 결과 있음','가설 v08 3.2절; 계획 v06 5·8절','오타 판정 규칙 또는 원본 입력상태 로그 승인 필요'),
('☑ 완료','H1 표에 분모·미설정 비교군 없음','편의시설·평점·가격별 설정군/비교군 분자·분모, OR, p-value 추가','가설 v08 3.1절','오타 기준 확정 후 민감도 재계산'),
('☑ 완료','호텔 클릭률과 상세조회율이 동일','231건씩이며 키 조합 완전 동일 확인; “호텔 상세진입” 한 지표로 통합','가설 v08 5·10절; 계획 v06 6절','발표 슬라이드에서도 이중 표기 금지'),
],[950,2100,3000,1760,1550],7.6)

d.add_heading('2. 발표 전에 손볼 것',level=1)
table(d,['상태','검토의견','반영 여부','근거 위치','발표 문장'],[
('☑ 완료','재검색과 클릭의 큰 차이를 인과처럼 말할 위험','H3를 복원하되 탐색적 채택·인과 보류를 함께 명시','가설 v08 9·14절','“연관이며 인과로 단정하지 않는다.”'),
('☑ 완료','재검색 상세표 모집단 불명확','모든 재검색 지표에 분자·분모와 분석단위를 명시','가설 v08 2·4·8·10절','“0건 검색 147건 중…”'),
('☑ 완료','소표본 지역변경 n을 함께 표시','지역·의도·재검색 유형 표에 모두 관측 n 표시','가설 v08 3.3·8·10절','“셀별 n이 작아 기술통계다.”'),
],[950,2400,2600,1800,1610],7.7)

d.add_heading('3. 문서 정합성',level=1)
table(d,['상태','검토의견','반영 내용','확인'],[
('☑ 완료','v06 파일명과 14절 버전 불일치','신규 가설문서를 v07로 작성하고 내부 버전도 v07로 통일','☑'),
('☑ 완료','0건 즉시 이탈형/0건 즉시 이탈 명칭 혼재','가설 v08·계획 v06 모두 “0건 즉시 이탈”로 통일','☑'),
('☑ 완료','608건과 296건 혼용 위험','296건을 현행 기준으로 명시하고 608건은 변경 사유 비교에만 사용','☑'),
],[1000,2900,4660,800],8.0)

d.add_heading('4. 교수 판단 8장 반영',level=1)
table(d,['상태','판단 항목','제출 의견','현재 반영','다음 결정'],[
('☑ 완료','최종 규모','1천 → 1만 점진 진행','계획 v06 2·10절에 반영','파일럿 QA 후 확장'),
('△ 판단 필요','세그먼트 비율','관측형 1종, 실패 유형 최소 50명','296건 재산정: 지속 실패 4.7%=약 47명, 즉시 이탈 0명','지속 실패 50명 보정·즉시이탈 별도 세트 승인'),
('☑ 완료','A2 범위','포함 승인, 지역·의도 표본 정의 선행','계획 v06 4절에 지역 6개·의도 5개 정의','희소셀 처리 승인'),
('△ 판단 필요','A3 범위','별도 SX 검토','명시적 오타 컬럼 없어 핵심 생성에서 제외','오타 정의 승인'),
('☑ 완료','C3 제외','제외 유지','가설·계획 모두 제외 유지','없음'),
('☑ 완료','HOTEL·ROOM','기존 풀 사용','기존 기준정보 참조 원칙 유지','없음'),
('☑ 완료','BOOKING','참고 시나리오','참고 지표·실제/합성 분리','없음'),
('☑ 완료','허용오차','핵심 ±3%p, 기타 ±5%p','교수 승인 대기 항목으로 유지','최종 승인 체크'),
],[950,1450,2350,3100,1510],7.5)

d.add_heading('5. 표본설계·가중치 지적',level=1)
table(d,['상태','지적','반영 컬럼','근거 위치'],[
('☑ 완료','과표집 비율을 실제 비율로 오해할 위험','sample_design_id, sample_set_type, sample_stratum','계획 v06 1절'),
('☑ 완료','가중치 컬럼 부재','target_population_share, sample_share, selection_probability, sample_weight, weight_version','계획 v06 1절'),
('☑ 완료','관측형과 과표집형 혼합 위험','물리적 파일·지표·발표 분리 원칙','계획 v06 2절'),
],[1000,2850,3600,1910],7.8)

d.add_heading('6. 296건 재기준화로 바뀐 결론',level=1)
table(d,['항목','이전 608건','현재 296건','처리'],[
('전체 0건률','310/608=51.0%','147/296=49.7%','현행은 296건'),
('B1 후속검색','96.1%','95.2%; OR 6.37, p<.001','채택'),
('B2 검색횟수','7.43 vs 7.88; p=.953, 기각','9.00 vs 2.93; p<.001','채택·주의로 변경'),
('실패군','17.6%','지속 실패 4.7%, 즉시 이탈 0%','이전 비율 사용 금지'),
],[1900,2300,2900,2260],7.9)
callout(d,'B2 주의','데이터가 바뀌어 통계 판정도 바뀌었다. 그러나 증강에서 “0건이 검색을 늘린다”를 직접 규칙으로 만들지 않고 탐색지속성을 먼저 생성한다. 인과 단정 금지는 그대로 유지한다.',GOLD)

d.add_heading('7. 제출 전 최종 체크',level=1)
items=[
'가설문서와 계획서의 기준 DB가 296건 원본으로 동일하다.',
'A1은 모든 절에서 부분 채택으로 표시됐다.',
'A3 미확정 사유와 필요한 데이터가 적혀 있다.',
'조건별 분자·분모·비교군·p-value가 있다.',
'클릭과 상세조회가 이중 KPI로 남아 있지 않다.',
'A2 지역·의도 정의와 관측 n이 있다.',
'관측형·과표집형과 가중치 컬럼이 분리됐다.',
'교수님 추가 판단 3건이 표시됐다.',
'608건 수치는 변경 이력 이외의 현행 생성값으로 사용하지 않는다.',
]
for x in items:
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(5); font(p.add_run('☑ '+x),10.4)
callout(d,'남은 3건','① 오타 판정 기준, ② 지속 실패 47명→최소 50명 보정 승인, ③ 관측 0명인 0건 즉시 이탈을 별도 스트레스 세트로 만들지 여부. 이 세 항목은 데이터로 자동 확정할 수 없어 교수 판단이 필요하다.',RED)

d.core_properties.title='32일차 제출 검토의견 반영 체크리스트'; d.core_properties.author='2팀'; OUT.parent.mkdir(parents=True,exist_ok=True); d.save(OUT); print(OUT)
