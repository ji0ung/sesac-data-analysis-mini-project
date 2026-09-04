from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "04_분석설계" / "팀프로젝트" / "2026" / "09" / "BI시각화_32-34일차_2팀실습_ASIS-TOBE_세그먼트AB재설계_20260904_v01_작업본.docx"

BLUE="2E74B5"; DARK="1F4D78"; NAVY="0B2545"; PALE="E8EEF5"; LIGHT="F2F4F7"
GRAY="666666"; RED="9B1C1C"; GREEN="1F5F4A"; GOLD="7A5A00"

def font(run,size=11,bold=False,color="000000",italic=False):
    run.font.name="Calibri"; rpr=run._element.get_or_add_rPr(); rf=rpr.get_or_add_rFonts()
    rf.set(qn("w:ascii"),"Calibri"); rf.set(qn("w:hAnsi"),"Calibri"); rf.set(qn("w:eastAsia"),"맑은 고딕")
    run.font.size=Pt(size); run.bold=bold; run.italic=italic; run.font.color.rgb=RGBColor.from_string(color)

def shade(cell,fill):
    pr=cell._tc.get_or_add_tcPr(); shd=pr.find(qn("w:shd"))
    if shd is None: shd=OxmlElement("w:shd"); pr.append(shd)
    shd.set(qn("w:fill"),fill)

def cell_margins(cell,top=90,bottom=90,start=120,end=120):
    pr=cell._tc.get_or_add_tcPr(); mar=pr.first_child_found_in("w:tcMar")
    if mar is None: mar=OxmlElement("w:tcMar"); pr.append(mar)
    for tag,val in [("top",top),("bottom",bottom),("start",start),("end",end)]:
        n=mar.find(qn("w:"+tag))
        if n is None: n=OxmlElement("w:"+tag); mar.append(n)
        n.set(qn("w:w"),str(val)); n.set(qn("w:type"),"dxa")

def geometry(table,widths):
    table.autofit=False; table.alignment=WD_TABLE_ALIGNMENT.LEFT; pr=table._tbl.tblPr
    for tag,val in [("tblW",sum(widths)),("tblInd",120)]:
        n=pr.find(qn("w:"+tag))
        if n is None: n=OxmlElement("w:"+tag); pr.append(n)
        n.set(qn("w:w"),str(val)); n.set(qn("w:type"),"dxa")
    grid=table._tbl.tblGrid
    for n in list(grid): grid.remove(n)
    for w in widths:
        n=OxmlElement("w:gridCol"); n.set(qn("w:w"),str(w)); grid.append(n)
    for row in table.rows:
        for i,c in enumerate(row.cells):
            tw=c._tc.get_or_add_tcPr().get_or_add_tcW(); tw.set(qn("w:w"),str(widths[i])); tw.set(qn("w:type"),"dxa")
            cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def repeat_header(row):
    p=row._tr.get_or_add_trPr(); n=OxmlElement("w:tblHeader"); n.set(qn("w:val"),"true"); p.append(n)

def table(doc,headers,rows,widths,size=8.4,center=(0,)):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Table Grid"
    for i,h in enumerate(headers):
        shade(t.rows[0].cells[i],PALE); p=t.rows[0].cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after=Pt(0); font(p.add_run(h),size,True,NAVY)
    repeat_header(t.rows[0])
    for ri,row in enumerate(rows):
        cs=t.add_row().cells
        for i,v in enumerate(row):
            if ri%2: shade(cs[i],"FAFBFC")
            p=cs[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if i in center else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.06; font(p.add_run(str(v)),size,False,"222222")
    geometry(t,widths); doc.add_paragraph().paragraph_format.space_after=Pt(1); return t

def heading(doc,text,level=1):
    p=doc.add_paragraph(text,style=f"Heading {level}"); p.paragraph_format.keep_with_next=True; return p

def para(doc,text,size=10.4,bold=False,color="222222",after=6,italic=False):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.15
    font(p.add_run(text),size,bold,color,italic); return p

def bullet(doc,text):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent=Inches(.5); p.paragraph_format.first_line_indent=Inches(-.25)
    p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.167; font(p.add_run(text),10.2); return p

def callout(doc,label,text,fill=LIGHT,color=NAVY):
    t=doc.add_table(rows=1,cols=1); c=t.cell(0,0); shade(c,fill); cell_margins(c,130,130,170,170)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.15
    font(p.add_run(label+"  "),10.5,True,color); font(p.add_run(text),10.5)
    geometry(t,[9360]); doc.add_paragraph().paragraph_format.space_after=Pt(2)

doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
sec.header_distance=sec.footer_distance=Inches(.492)
normal=doc.styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.10
for name,size,color,before,after in [("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",12,DARK,8,4)]:
    s=doc.styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(header.add_run("2팀 · 32~34일차 실습 재설계"),9,False,GRAY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(footer.add_run("AS-IS → TO-BE · 2026-09-04"),8.5,False,GRAY)

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(4); font(p.add_run("32~34일차 2팀 실습 재설계서"),23,True,NAVY)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(18); font(p.add_run("1,000명 관측형에서 10,000명 의도 세그먼트 A/B 시뮬레이션으로"),14,True,DARK)
table(doc,["항목","내용"],[
    ("문서 목적","32일차 가설·33일차 인사이트·34일차 개선안을 확장 설계에 맞게 연속 개정"),
    ("AS-IS","v03 원본 296건 관측 분석 + 관측형 합성 1,000명"),
    ("TO-BE","control 5,000명 + treatment 5,000명, 의도 세그먼트별 효과·불확실성 비교"),
    ("변경 원칙","원본 실습 문서는 보존하고 본 문서를 신규 변경본으로 사용"),
],[1800,7560],9.2)
callout(doc,"핵심 변경","32일차의 관측 가설은 기준선으로 남기고, 33일차 인사이트를 세그먼트 가설의 근거로 재배치한 뒤, 34일차 개선안을 실험군의 의도별 처치로 전환한다.",PALE)

heading(doc,"1. 변경 전·후 전체 구조",1)
table(doc,["일차","기존 AS-IS","변경 TO-BE","연결 산출물"],[
    ("32일차","원본 데이터에서 조건·재검색·회복 가설 판정","기준선 가설 + 의도 세그먼트별 A/B 가설 + 불확실성 설계","10,000명 실험 가설서"),
    ("33일차","관측 인사이트 카드 A·B·C·D·F·G·H","기존 카드를 기준선 근거로 보존 + 의도군별 실험 카드 추가","세그먼트 인사이트 카드"),
    ("34일차","전체 0건 사용자 대상 정적 선택지·관찰형 파일럿","대조군·실험군 분리 + 의도별 맞춤 선택지 + 시나리오 반복","우선순위·실행안·Power BI"),
],[1200,2770,3430,1960],8.2)

heading(doc,"2. 공통 기준선: 원본 v03 DB",1)
table(doc,["기준","원본 수치","비율","사용 용도"],[
    ("검색 연결 사용자","41명","-","10,000명 확률 추정 모수"),
    ("검색 세션","43개","-","세션 단위 QA"),
    ("0건 검색","147/296건","49.7%","검색 기준선"),
    ("0건 경험 사용자","27/41명","65.9%","실험 진입 확률"),
    ("0건 경험 세션","28/43개","65.1%","세션 기준선"),
],[2500,1900,1300,3660],8.7,center=(0,1,2))
table(doc,["첫 0건 필터 상태","0건 경험 사용자","27명 내 비중"],[
    ("가격+옵션 수 복합 설정","17명","63.0%"),
    ("가격·옵션 수 제한 없음","7명","25.9%"),
    ("옵션 수만 설정","2명","7.4%"),
    ("가격만 설정","1명","3.7%"),
],[4800,2500,2060],8.7,center=(0,1,2))
callout(doc,"분류 주의","합성 1,000명은 QA용이다. 필터 상태·추정 의도는 실험 전 분류, SG1~SG4는 처치 후 판정하는 검색·전환 성과 결과군으로 분리한다.","FFF3CD",RED)

heading(doc,"3. 32일차 가설 설계 AS-IS → TO-BE",1)
table(doc,["32일차 AS-IS","최신 판정·문제","TO-BE 가설"],[
    ("H1 제한 조건이 강할수록 0건률이 높다.","관측 근거는 유지되나 조건 중첩·지역·날짜 교란 가능","H1 예산·품질 유연형에게 조건 완화를 제안하면 control보다 결과 회복률이 높다."),
    ("H2-1 0건 후 다음 검색이 많다.","최신 140/147=95.2%; 이탈 로그 부족","H2 맞춤 제안을 받은 treatment의 조건 변경률이 control보다 높다."),
    ("H2 회복: 1회 성공 22.5%, 최종 71.4%","최신 정의로 즉시 24/140=17.1%, 최종 21/28=75.0%; grain 다름","H3 빠른 해결형의 treatment에서 즉시 회복·최종 회복이 각각 높다."),
    ("H3 재검색 방법별 결과·클릭 차이","지역 10/24=41.7%, 검색어 3/10=30.0%, 완화 11/41=26.8%; 탐색적","H4 위치 유연형에게 인접지역을 제안하면 control보다 회복률이 높다."),
    ("동일 조건·강화에서 회복 0%","동일 0/53, 강화 0/10; 전이 표본이며 인과 단정 불가","H5 조건 고수형의 treatment에서 동일 조건 반복률이 낮다."),
    ("검색어 수정 후 상세진입 차이","상세진입 3/10=30.0%; 표본 매우 작음","H6 표현 수정형에게 연관 검색어를 제공하면 control보다 상세진입률이 높다."),
],[2600,2820,3940],7.8)
heading(doc,"3.1 최종 가설 체계",2)
table(doc,["ID","가설","주 KPI","주요 세그먼트"],[
    ("H0","control과 treatment의 성과 차이가 없다.","결과 회복률","ALL"),
    ("H1","의도 맞춤 제안은 전체 결과 회복률을 높인다.","회복률 차이","ALL"),
    ("H2","개선 효과의 크기는 의도 세그먼트별로 다르다.","세그먼트×처치 상호작용","6개 의도군"),
    ("H3","결과 회복은 호텔 상세진입 증가로 이어진다.","회복 후 상세진입률","ALL"),
    ("H4","반복 방지 안내는 동일 조건 반복률을 낮춘다.","동일 조건 반복률","condition_keeper"),
    ("H5","지역 확대는 위치 유연형의 회복에 특히 효과적이다.","회복률","location_flexible"),
    ("H6","검색어 제안은 표현 수정형의 상세진입을 높인다.","상세진입률","query_reframer"),
],[800,4520,2240,1800],8.1,center=(0,3))
para(doc,"모든 세부 가설의 귀무가설은 '같은 의도군에서 control과 treatment의 차이가 없다'로 정의한다.",9.4,False,RED)

heading(doc,"4. 33일차 인사이트 카드 AS-IS → TO-BE",1)
table(doc,["카드","AS-IS 관측 인사이트","TO-BE 역할","연결 의도군·가설"],[
    ("A","제한 조건에 0건이 집중","기준선 근거로 보존; 가격·옵션 수 완화 treatment 정의","budget_flexible·option_count_flexible / H1·H2"),
    ("B","0건 후 후속 검색 95.2%, 이탈 로그 공백","효과 카드가 아닌 측정 게이트로 재분류","ALL / 로깅 선행조건"),
    ("C","즉시 17.1%·최종 75.0%; 분모 다름","두 KPI 분리 유지; 단계형 처치 근거","rapid_resolver / H3"),
    ("D","지역·검색어·완화 방법별 차이","세그먼트별 treatment 매핑 근거","location_flexible·query_reframer / H5·H6"),
    ("F","동일 조건·강화 회복 0%","반복 방지 treatment 근거; 0%를 절대 규칙으로 사용 금지","condition_keeper / H4"),
    ("G","직접 3/43, 재검색 후 28/43, 미상호작 12/43","직접·재검색 후 성공 구조 QA; 이탈률로 사용 금지","ALL / 보조 KPI"),
    ("H","고유 1위 클릭 29/149=19.5%","회복 후 결과 품질 통제 변수","ALL / H3 보정"),
],[800,2900,2900,2760],7.7,center=(0,))
callout(doc,"33일차 개정 원칙","기존 카드는 삭제하지 않는다. '실제 관측 기준선' 배지를 붙이고, 그 카드에서 도출된 세그먼트 가설·treatment·불확실성을 추가한다.",PALE)

heading(doc,"5. 34일차 개선안 AS-IS → TO-BE",1)
table(doc,["개선안","AS-IS","TO-BE","측정"],[
    ("A-1 조건 완화","0건 사용자에 가격 확대·옵션 수 감소 제안","treatment에서 추정 의도와 현재 필터 상태에 맞는 제안을 우선 노출","노출→선택→변경→회복"),
    ("D-1 지역 변경","전체 0건 사용자에게 인접지역 정적 제시","location_flexible에 우선 제시; 다른 의도군에서의 효과와 부정 효과도 비교","세그먼트별 회복·상세진입"),
    ("C-1 단계형 회복","첫 실패 후 아직 안 쓴 방법을 1개씩 제안","rapid_resolver에 회복 가능성 상위 제안; 반복 노출 제한","즉시 회복·최종 회복 분리"),
    ("F-1 반복 방지","A-1 반응 후 2차 개발","condition_keeper에 핵심 treatment로 포함; 반복 경고 + 조건 유지/완화 결과 비교","동일 조건 반복률↓"),
    ("D-2 검색어 제안","소규모 2주 측정","query_reframer에 연관 검색어 10개·자동완성; n=10 근거는 넓은 불확실성 적용","수정률·회복·상세진입"),
    ("H-1 순위 검증","가격·옵션 수·호텔 특성 통제 후 순위별 클릭 재분석","세그먼트·sample_set_type을 통제에 추가; 순위 효과와 treatment 효과 분리","고유 결과 단위 CTR"),
],[1500,2400,3560,1900],7.55)

heading(doc,"6. 34일차 우선순위 변경",1)
table(doc,["순위","TO-BE 실행안","근거 카드","우선 대상","판정"],[
    ("1","A-1 의도별 조건 완화","A","budget_flexible·option_count_flexible","높음/낮음 · 실행"),
    ("2","D-1 의도별 지역 확대","D","location_flexible","높음/낮음 · 실행"),
    ("3","F-1 동일 조건 반복 방지","F","condition_keeper","높음/중간 · 실험"),
    ("4","C-1 단계형 회복","C","rapid_resolver","높음/높음 · 단계"),
    ("5","D-2 연관 검색어","D","query_reframer","탐색/낮음 · 소규모"),
    ("6","H-1 순위 통제 분석","H","ALL","분석 실행"),
    ("7","B-1 이탈 로그 보강","B·G","ALL","선행조건"),
],[800,3050,1000,2150,2360],8.1,center=(0,2,3,4))
para(doc,"기존 34일차의 1차 관찰형 운영안은 삭제하지 않고 '라이브 서비스 전 초기 파일럿'으로 남긴다. 10,000명 시뮬레이션은 이와 별도의 사전 실험 설계이다.",9.4,False,RED)

heading(doc,"7. 의도 세그먼트·treatment 매핑",1)
table(doc,["intent_segment","의도","주 treatment","관측 행동"],[
    ("condition_keeper","조건을 쉽게 포기하지 않음","반복 경고·대안 비교","same / strengthen / relax"),
    ("location_flexible","위치를 넓혀서라도 찾음","인접지역 확대","region_change"),
    ("budget_flexible","예산을 조정할 수 있음","가격 범위 확대","price_relax"),
    ("option_count_flexible","요구 옵션 수를 줄일 수 있음","요구 옵션 수 줄이기","option_count_relax"),
    ("query_reframer","검색 표현을 수정","연관어·자동완성","query_change"),
    ("rapid_resolver","조건을 빠르게 바꿔 결과를 찾음","회복 가능성 순 단계 제안","mixed_change"),
],[2000,2500,2500,2360],8.25)

heading(doc,"7.1 실험 후 검색·전환 성과 결과군",2)
table(doc,["코드","결과군","판정 기준","역할"],[
    ("SG1","직접 성공형","첫 검색 결과≥1, 세션 내 상세진입","0건 실험 전 기준선"),
    ("SG2","결과 노출·미선택형","첫 검색 결과≥1, 상세진입 없음","정렬·추천 개선 후보"),
    ("SG3","재검색 회복형","첫 검색 0건 후 후속 검색 결과 발생","실험군에서 증가 목표"),
    ("SG4","지속 실패형","첫 검색 0건, 세션 종료까지 미회복","실험군에서 감소 목표"),
],[900,2300,3900,2260],8.4,center=(0,1))
para(doc,"정리: 추정 의도는 '누구에게 무엇을 보여줄지'를 정하고, SG1~SG4는 '보여준 후 어떤 성과가 나왔는지'를 판정한다. 두 분류를 다른 컬럼으로 모두 사용한다.",9.4,False,RED)

heading(doc,"8. 효과 가정·불확실성 개정",1)
table(doc,["요소","AS-IS","TO-BE"],[
    ("효과값","관측 비율 또는 방향성 중심","보수 +2~3%p, 기준 +5~8%p, 낙관 +10~15%p 시작값"),
    ("표본 표시","분자/분모·n 주석","세그먼트 n + 95% 불확실성 구간 + treatment 우세 확률"),
    ("희소 표본","검색어 수정 n=10을 한계로 표시","넓은 효과 분포·고불확실 배지·순위 확정 금지"),
    ("반복 실행","단일 관측값","시나리오별 최소 1,000회 Monte Carlo"),
    ("의도 비중","직접 필드 없음","확률적 부여 + 비중 ±5%p 민감도"),
],[1800,3200,4360],8.4)
callout(doc,"수치 해석","시뮬레이션의 +5~8%p는 가정이지 성과가 아니다. 원본 검색 사용자 41명·43세션·296검색의 관측값이 기준선이며, 합성 1,000명은 재현 QA용이다. 실험군의 결과는 가정·seed·run에 따라 달라진다.","FFF3CD",RED)

heading(doc,"9. 데이터·KPI 계약 변경",1)
table(doc,["구분","TO-BE 필수 컬럼·정의"],[
    ("표본","sample_set_type, sample_stratum, random_seed"),
    ("의도","intent_segment, intent_assignment_prob, intent_version"),
    ("처치","treatment_policy, suggestion_type, suggestion_exposed, suggestion_selected"),
    ("행동","observed_behavior, condition_changed, same_condition_repeat"),
    ("결과","result_recovered, hotel_detail_entered"),
    ("시나리오","scenario_type, simulation_run_id, uplift_assumption"),
],[1800,7560],8.8)
bullet(doc,"결과 회복률 = 결과≥1 후속 검색 / 0건 후 후속 검색")
bullet(doc,"조건 변경률 = 조건을 넓힌 후속 검색 / 제안 노출 0건 검색")
bullet(doc,"세션 최종 회복률 = 마지막 검색 결과≥1 세션 / 0건 경험 세션")
bullet(doc,"상세진입률 = 회복 후 hotel_detail_view 존재 전이 / 회복 전이")
para(doc,"v03에 BOOKING 36건·booking_complete 36건·연결 검색 30건이 있으므로 예약을 최종 보조 KPI로 확장한다. 다만 BOOKING.data_origin은 '전체 가상 예약 시뮬레이션'으로 표시되어 있으므로 출처 정정·승인, 중복 의심 6건, ROOM–HOTEL 불일치 2건을 먼저 해결한다. 연결은 booking_complete.search_id 기반 브리지를 사용한다.",9.4,False,RED)

heading(doc,"10. Power BI 표현 변경",1)
table(doc,["페이지","AS-IS","TO-BE 추가"],[
    ("Overview","관측 0건·회복·상세진입","A/B 차이, 95% 구간, treatment 우세 확률"),
    ("세그먼트","검색 조건·재검색 방법","의도군 비중, control/treatment 균형, 관측 행동"),
    ("A/B 성과","개선안 관찰 KPI","전체·세그먼트별 uplift, 퍼널, 효과 순위"),
    ("불확실성","소표본 주석","보수·기준·낙관, 95% 구간, 민감도"),
    ("QA","데이터 품질 주석","DB 용량, 배정 균형, PK/FK, 가정·seed·run"),
],[1700,3150,4510],8.5)

heading(doc,"11. 문서별 신규 버전 적용표",1)
table(doc,["대상 문서","보존할 내용","신규 버전에 추가·교체할 내용"],[
    ("32일차 2팀 실습","객체·grain·JOIN·분모 검증, 관측 가설 판정","최신 v03 수치, 1,000명 비중, H0~H6, 층화 5,000/5,000, 불확실성"),
    ("33일차 2팀 실습","카드 A·B·C·D·F·G·H의 관측 근거·한계","기준선 배지, 의도군, treatment, KPI, 불확실성, 반증 조건"),
    ("34일차 2팀 실습","개선안 목록·우선순위·선행조건·중단 기준","세그먼트 맞춤 treatment, A/B 판정, scenario_type, Power BI 비교"),
],[2400,2700,4260],8.25)

heading(doc,"12. 실행 전 게이트",1)
table(doc,["Gate","확인 항목","통과 기준"],[
    ("G0","1,000명 DB 경량화","VACUUM 포함 20MB 이하 + QA 재PASS"),
    ("G1","의도 부여 규칙·비중 승인","1,000명 전수 코드·확률·seed 추적"),
    ("G2","control/treatment 사전 균형","의도·지역·조건 허용차 이내"),
    ("G3","10,000명 DB 용량·무결성","150MB 이하, PK/FK·시간순서 위반 0건"),
    ("G4","시뮬레이션 가정 표시","가정·불확실성·run_id·seed 누락 0건"),
],[900,4660,3800],8.6,center=(0,))
callout(doc,"최종 정리","32일차는 '왜 검증하는가', 33일차는 '어떤 근거로 세그먼트·처치를 정했는가', 34일차는 '무엇을 실행·비교할 것인가'를 담는 연속 구조로 개정한다.",PALE)

heading(doc,"부록. 참고한 현행 문서",1)
for s in [
    "호텔검색_32일차_2팀_실행체크리스트_20260902_v03_현행본_데이터증강반영.docx",
    "BI시각화_32일차_2팀_기초실습_문제-1·2_20260902_v02_제출본.docx",
    "BI시각화_33일차_2팀프로젝트실습_해답_20260903_v03_제출본.docx",
    "BI시각화_34일차_2팀프로젝트실습_해답_20260904_v01_제출본.docx",
    "호텔검색_1000명_10000명_세그먼트AB시뮬레이션_증강계획서_20260904_v01_작업본.docx",
]: bullet(doc,s)

OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT); print(OUT)
