from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / '04_분석설계' / '이전버전' / '데이터증강계획서' / '호텔검색_데이터증강계획서_세그먼트가설중심_20260902_v04_이전본.docx'
SOURCE = '호텔검색_가설문서_결과없음_재검색_세그먼트_20260902_v06.docx'

BLUE = '2E74B5'; DARK = '1F4D78'; NAVY = '0B2545'; PALE = 'E8EEF5'
LIGHT = 'F2F4F7'; CALLOUT = 'F4F6F9'; GOLD = '7A5A00'; RED = '9B1C1C'; GREEN = '1F5F4A'; GRAY = '666666'

def set_font(run, size=11, bold=False, color='000000', name='Calibri'):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(size); run.bold = bold; run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def margins(cell, top=80, bottom=80, start=120, end=120):
    tcPr = cell._tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for tag, value in [('top',top),('bottom',bottom),('start',start),('end',end)]:
        node = tcMar.find(qn('w:'+tag))
        if node is None: node = OxmlElement('w:'+tag); tcMar.append(node)
        node.set(qn('w:w'), str(value)); node.set(qn('w:type'), 'dxa')

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr(); el = OxmlElement('w:tblHeader'); el.set(qn('w:val'), 'true'); trPr.append(el)

def set_table_geometry(table, widths):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    for tag, val in [('tblW', sum(widths)), ('tblInd', 120)]:
        el = tblPr.find(qn('w:'+tag))
        if el is None: el = OxmlElement('w:'+tag); tblPr.append(el)
        el.set(qn('w:w'), str(val)); el.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for c in list(grid): grid.remove(c)
    for w in widths:
        col = OxmlElement('w:gridCol'); col.set(qn('w:w'), str(w)); grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcW = cell._tc.get_or_add_tcPr().get_or_add_tcW(); tcW.set(qn('w:w'), str(widths[idx])); tcW.set(qn('w:type'), 'dxa')
            margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(doc, headers, rows, widths, font=8.6):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
    for i, h in enumerate(headers):
        shade(t.rows[0].cells[i], PALE)
        p=t.rows[0].cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(h), font, True, NAVY)
    set_repeat_header(t.rows[0])
    for ridx, row in enumerate(rows):
        cells=t.add_row().cells
        for i, val in enumerate(row):
            if ridx % 2: shade(cells[i], 'FAFBFC')
            p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.08
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_font(p.add_run(str(val)), font, False, '222222')
    set_table_geometry(t, widths)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t

def add_bullet(doc, text, level=0, checked=None):
    p=doc.add_paragraph(style='List Bullet' if checked is None else 'Normal')
    if checked is not None:
        set_font(p.add_run(('☐ ' if not checked else '☑ ') + text), 10.5)
    else:
        set_font(p.add_run(text), 10.5)
    p.paragraph_format.left_indent=Inches(.375 + .25*level); p.paragraph_format.first_line_indent=Inches(-.188)
    p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.25
    return p

def callout(doc, label, text, fill=CALLOUT, color=NAVY):
    t=doc.add_table(rows=1, cols=1); c=t.cell(0,0); shade(c,fill); margins(c,120,120,160,160)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.15
    set_font(p.add_run(label+'  '),10.5,True,color); set_font(p.add_run(text),10.5,False,'222222')
    set_table_geometry(t,[9360]); doc.add_paragraph().paragraph_format.space_after=Pt(2)

def add_heading(doc, text, level=1):
    p=doc.add_paragraph(text, style=f'Heading {level}'); return p

doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11)
sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
sec.header_distance=sec.footer_distance=Inches(.492)

# Preset: compact_reference_guide; header: memo_masthead
styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11)
normal._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕')
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for name,size,color,before,after in [('Heading 1',16,BLUE,18,10),('Heading 2',13,BLUE,14,7),('Heading 3',12,DARK,10,5)]:
    st=styles[name]; st.font.name='Calibri'; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)
    st._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True

header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.LEFT
set_font(header.add_run('TEAM 2 · HOTEL SEARCH DATA AUGMENTATION'),8.5,True,GRAY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT
set_font(footer.add_run('2팀 데이터 증강 계획서 · 2026.09.02'),8.5,False,GRAY)

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)
set_font(p.add_run('데이터 증강 계획서'),23,True,NAVY)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(14)
set_font(p.add_run('v06 기반 · 세그먼트 및 가설 중심 실행안'),14,False,DARK)
for lab,val in [('기준 문서',SOURCE),('작성일','2026-09-02'),('대상','2팀 제작팀 / 지도 교수님'),('문서 버전','v04')]:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    set_font(p.add_run(lab+': '),10.5,True); set_font(p.add_run(val),10.5)
callout(doc,'핵심 결정','증강 데이터는 “0건 사용자 = 검색 횟수 증가”를 전제로 만들지 않는다. 세션 탐색성향과 실패·회복 경로를 각각 생성한 뒤 연결한다.', 'EAF2F8', NAVY)

add_heading(doc,'1. 목적과 적용 범위',1)
doc.add_paragraph('본 계획서는 v06에서 정리된 결과 없음, 재검색, 사용자 세그먼트 가설을 실제 합성 데이터 생성 규칙으로 전환한다. 관측 데이터의 통계값은 방향과 현실성 검증의 기준으로 사용하며, p-value 재현 자체를 생성 목표로 삼지 않는다.')
add_bullet(doc,'우선 생성 테이블: USER → SEARCH → SEARCH_FILTER → SEARCH_RESULT → EVENT → BOOKING')
add_bullet(doc,'HOTEL·ROOM은 기존 기준정보 풀을 우선 재사용하며, 증강 여부는 교수님 승인 후 확정한다.')
add_bullet(doc,'모든 합성 행에 data_origin, scenario_id, generation_version을 기록한다.')
add_bullet(doc,'실데이터와 합성데이터의 성과 지표를 분리 집계한다.')

add_heading(doc,'2. v06 판정의 증강 반영 원칙',1)
add_table(doc,['판정','증강 반영','금지·주의'],[
('B2 제외','0건 여부와 세션 검색 횟수를 별도 변수로 생성','0건 사용자에게 검색 횟수를 일괄 가산하지 않음'),
('C3 제외','품질 우선형을 현재 핵심 증강 규칙에서 제외','별점·등급 중심 확률을 임의로 강화하지 않음'),
('A3 보류','오타·중간입력 시나리오를 별도 실험군으로 격리','핵심 분포에 자동 혼합하지 않음'),
('B1 채택','0건 이후 후속 검색 가능성을 높게 유지','검색 횟수 자체를 B1의 결과로 해석하지 않음'),
('B3 채택','즉시 회복과 세션 최종 회복을 분리 생성·검증','최종 회복률을 다음 검색 성공률로 대체하지 않음'),
], [1200,3900,4260], 8.8)

add_heading(doc,'3. 세그먼트 설계',1)
add_heading(doc,'3.1 상호배타적 행동·결과 세그먼트',2)
add_table(doc,['세그먼트','관측 기준','생성해야 할 핵심 경로','초기 기준'],[
('직접 성공','첫 검색에서 결과·선택 완료','검색 → 결과 노출 → 클릭/선택','12명 / 17.6%'),
('결과 노출·미선택','결과는 있으나 선택 없음','검색 → 결과 노출 → 이탈 또는 재탐색','18명 / 26.5%'),
('재검색 회복','실패 후 세션 내 성공','0건 → 조건변경 → 결과/선택','26명 / 38.2%'),
('지속 실패','재검색 후에도 결과 없음','0건 → 반복 탐색 → 최종 실패','6명 / 8.8%'),
('0건 즉시 이탈','첫 0건 후 종료','0건 → 후속 행동 없음','6명 / 8.8%'),
], [1450,2050,3800,2060], 8.5)
callout(doc,'운영 원칙','위 비율은 초기 베이스라인이다. 최종 데이터에서 그대로 고정할지, 실패 세그먼트를 과표집할지는 교수님 판단 항목으로 남긴다.')

add_heading(doc,'3.2 다중라벨 검색 성향',2)
for x in ['지역 고정형: 지역 유지 비율이 높고 다른 조건을 조정','가격 민감형: 가격 상한·하한 설정 및 완화가 빈번','테마·숙소유형형: 테마·숙소유형 조건을 중심으로 탐색','호텔 지명형: 특정 호텔명 또는 유사 문자열로 탐색','비교 탐색형: 다회 검색·다수 결과 비교 후 선택','빠른 결정형: 적은 검색과 빠른 선택','입력 탐색형: 짧은 입력·중간 입력·오타 수정 경로']:
    add_bullet(doc,x)
callout(doc,'현재 제외','품질 우선형(C3)은 필요한 데이터가 부족하고 SEARCH·SEARCH_FILTER 검증보다 우선순위가 낮으므로 이번 증강 규칙에서 제외한다.', 'FFF4E5', GOLD)

add_heading(doc,'4. 가설 중심 생성 규칙',1)
rows=[
('A1','필터 제한이 강할수록 0건 증가','편의시설 3개+, 최소평점, 가격 조건의 조합 강도','조건강도별 0건률이 단조 또는 준단조 증가','진행'),
('A2','검색 의도/지역별 0건 차이','지역·의도 카테고리와 결과 수','표본 확보 후 그룹별 차이 확인','교수 판단'),
('A3','오타·중간입력이 0건과 연결','입력 길이, 수정 이벤트, 오타 플래그','별도 실험군에서 0건률 비교','보류'),
('B1','0건 후 후속 검색 증가','0건 직후 search 이벤트','후속 검색률이 비0건보다 높음','진행'),
('B2','0건 사용자의 전체 검색 증가','해당 없음','증강 규칙에 사용하지 않음','제외'),
('B3','세션 내 최종 회복 가능','즉시 성공·최종 성공을 별도 플래그','최종 회복률 > 즉시 회복률','진행'),
('B4','필터 완화가 반복보다 회복에 유리','첫 실패 다음 행동 유형','완화군 성공률 > 동일조건 반복군','진행'),
('B5','조건 강화는 회복에 불리','강화 행동과 다음 결과','강화군 성공률이 완화군보다 낮음','진행'),
('C1/C2','지역·가격 성향별 경로 차이','USER 라벨과 SEARCH_FILTER 패턴','라벨별 조건변경 방향 일치','진행'),
('C4~C6','명칭·비교·빠른결정 성향','검색문·횟수·클릭 시점','라벨별 행동 규칙과 로그 일치','진행'),
('C3','품질 우선 세그먼트','해당 없음','현재 버전에서 생성·판정하지 않음','제외'),
]
add_table(doc,['가설','핵심 주장','생성 변수','판정 체크','상태'],rows,[850,2050,2380,3080,1000],7.9)

add_heading(doc,'5. 생성 로직과 확률 운영',1)
add_heading(doc,'5.1 권장 생성 순서',2)
for i,x in enumerate(['USER에 행동·결과 세그먼트 1개와 검색성향 라벨 0~N개 부여','세그먼트별 탐색 지속성(search persistence)을 먼저 생성','각 SEARCH에 SEARCH_FILTER 1행을 생성하고 제한 강도 계산','조건·시나리오에 따라 SEARCH_RESULT의 0건/비0건 결정','0건이면 후속 행동 유형(반복·완화·강화·혼합·이탈)을 생성','EVENT를 시간순으로 생성하고 클릭 대상이 결과 목록에 존재하는지 검증','BOOKING은 상세/클릭 이후에만 생성하고 합성 표시'],1):
    p=doc.add_paragraph(style='List Number'); set_font(p.add_run(x),10.5); p.paragraph_format.space_after=Pt(4)

add_heading(doc,'5.2 관측값을 사용하는 방법',2)
add_table(doc,['관측 기준','용도','생성 시 처리'],[
('0건 검색 310/608 = 51.0%','전체 난이도 기준','세그먼트·조건별 차이를 둔 뒤 전체값이 허용범위에 드는지 확인'),
('0건 후 후속검색 298/310 = 96.1%','B1 방향 기준','높은 후속검색 확률을 두되 즉시이탈 세그먼트는 예외'),
('다음 검색 성공 67/298 = 22.5%','즉시 회복 기준','행동 유형별로 차등화'),
('세션 최종 회복 30/42 = 71.4%','B3 최종 기준','즉시 회복과 분리하여 세션 종료 시점에 판정'),
('검색 횟수 7.43 vs 7.88, p=.953','B2 제외 근거','0건 여부로 검색 횟수 분포를 이동시키지 않음'),
], [2350,2600,4410], 8.4)
callout(doc,'확률 조정 원칙','관측 비율을 하나의 고정 확률로 복사하지 않고 낮음·기준·높음의 3개 설정값으로 관리한다. 검증은 p-value 일치보다 가설 방향, 논리 제약, 분포 왜곡 여부를 우선한다.')

add_heading(doc,'6. 실험 시나리오',1)
add_table(doc,['ID','시나리오','주요 변경','확인 가설','산출 비교'],[
('S0','관측형 기준군','현재 검색·필터 패턴 유지','A1, B1, B3','0건률·회복률 기준선'),
('S1','인접 지역 제안','지역 고정 실패 시 인접 지역 후보 노출','A2, C1','결과 노출·선택 변화'),
('S2','필터 완화 제안','가격/평점/편의시설 단계적 완화','A1, B4, C2','즉시·최종 회복 변화'),
('S3','통합 제안','인접 지역 + 필터 완화 선택지','A1, A2, B4','회복률 및 선택률 변화'),
('SX','입력 정정 실험','오타·중간입력 → 수정 이벤트','A3','0건률·수정 후 성공률'),
], [700,1800,3200,1500,2160], 8.4)

add_heading(doc,'7. 제작팀 실행 영역',1)
callout(doc,'배정 표기','아래 담당은 권장 배정안이며 팀 내부 확정이 필요하다. 교수님 판단이 필요한 항목은 8장에서 분리한다.')
add_table(doc,['작업묶음','권장 담당','실행 내용','산출물·완료 기준'],[
('스키마·키 설계','권순성','PK/FK, 생성 순서, HOTEL·ROOM 참조 규칙 정의','ERD/DDL; 고아키·중복키 0건'),
('세그먼트·확률표','박지현','5개 결과 세그먼트와 다중라벨별 분포·전이 범위 작성','설정표; 합계·배타성·범위 검증'),
('생성기 구현','손지영','USER~BOOKING 생성, scenario_id·seed·version 적용','재실행 가능한 코드와 샘플 1천 명'),
('검증·리포트','이소이','가설 방향, 분포, 시간순서, 실제/합성 분리 점검','QA 보고서와 오류 목록'),
('통합 리뷰','전원','B2/C3 제외 여부, 가설-변수 매핑, 결과 해석 검토','리뷰 체크리스트 전 항목 확인'),
], [1350,1200,3700,3110], 8.2)

add_heading(doc,'7.1 제작팀 체크리스트',2)
for x in ['B2를 생성 로직·설명·대시보드에서 모두 제외했는가?','C3를 핵심 세그먼트 분포와 확률표에서 제외했는가?','행동·결과 세그먼트가 사용자당 정확히 1개인가?','검색 성향 라벨이 다중 선택 가능하도록 설계되었는가?','SEARCH_FILTER가 모든 SEARCH와 1:1로 연결되는가?','0건 SEARCH에는 SEARCH_RESULT 행이 0개인가?','클릭 호텔이 해당 검색 결과 목록에 존재하는가?','모든 EVENT가 사용자·세션 내 시간순서에 맞는가?','실데이터와 합성데이터가 data_origin으로 분리되는가?','동일 seed·version으로 결과를 재현할 수 있는가?']:
    add_bullet(doc,x,checked=False)

add_heading(doc,'8. 교수님 판단 필요사항',1)
add_table(doc,['결정 항목','선택지','권장안','판단이 필요한 이유','결정'],[
('최종 규모','1천 pilot / 1만 final / 기타','1천 검증 후 1만 확장','오류를 작은 비용으로 선제 확인','☐'),
('세그먼트 비율','관측비율 / 실패군 과표집','관측형+과표집형 2종','희귀 실패경로의 분석력과 현실성 균형','☐'),
('A2 범위','이번 포함 / 보류','조건부 포함','지역·의도 표본 정의가 먼저 필요','☐'),
('A3 범위','핵심 혼합 / 별도 실험 / 제외','별도 SX 실험','오타 정의와 라벨링 기준 미확정','☐'),
('C3 제외 승인','제외 유지 / 보완 후 포함','제외 유지','자료 부족·현 우선순위 낮음','☐'),
('HOTEL·ROOM','기존 풀 / 기준정보도 증강','기존 풀 사용','허구 숙소가 해석을 왜곡할 위험','☐'),
('BOOKING 활용','핵심 KPI / 참고 시나리오 / 제외','참고 시나리오','실제 예약 근거가 부족함','☐'),
('허용 오차','±3%p / ±5%p / 구간별','핵심 ±3%p, 기타 ±5%p','QA 통과 기준을 사전 확정해야 함','☐'),
], [1300,1950,1780,3530,800], 7.9)

add_heading(doc,'9. 품질 검증 및 승인 게이트',1)
add_table(doc,['게이트','필수 검사','통과 기준','담당'],[
('G1 구조','PK/FK, 중복, null, 참조 무결성','치명 오류 0건','제작팀'),
('G2 행동','시간순서, 클릭-결과 일치, 예약 선행행동','논리 위반 0건','제작팀'),
('G3 분포','세그먼트·0건·회복·행동유형 분포','승인된 허용오차 이내','제작팀'),
('G4 가설','A1/B1/B3/B4/B5 및 C계열 방향','활성 가설 방향 충족','제작팀'),
('G5 제외','B2·C3가 규칙에 재유입되지 않음','재유입 0건','전원'),
('G6 승인','8장 의사결정 기록 및 버전 동결','교수님 확인 완료','교수님'),
], [900,3660,3000,1800], 8.3)

add_heading(doc,'10. 실행 순서와 산출물',1)
for title,detail in [
('1단계 · 설계 동결','교수님 판단표 확정 → config_v01.yaml, 스키마 명세, 가설-변수 매핑표'),
('2단계 · 파일럿','1천 사용자 생성 → 구조·행동·분포 QA → 오류 수정'),
('3단계 · 본 생성','승인된 seed/version으로 1만 사용자 생성 → 산출 테이블 저장'),
('4단계 · 분석 비교','실제/합성 분리 분석 → 세그먼트별 0건·재검색·회복 비교'),
('5단계 · 인계','생성 코드, 설정값, 데이터 사전, QA 결과, 변경 이력 패키징')]:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4)
    set_font(p.add_run(title+' — '),10.5,True,DARK); set_font(p.add_run(detail),10.5)

add_heading(doc,'11. 최종 승인 체크',1)
for x in ['☐ 교수님 판단 항목 8개가 모두 결정되었다.','☐ B2·C3 제외가 코드·설정·보고서에 동일하게 반영되었다.','☐ 세그먼트와 가설별 생성 규칙이 설정 파일로 분리되었다.','☐ 파일럿 QA가 G1~G5를 통과했다.','☐ 최종 생성 버전과 seed가 기록되었다.','☐ 실제 데이터와 합성 데이터의 결과가 분리 보고된다.']:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5); set_font(p.add_run(x),10.5)

add_heading(doc,'부록 A. 판정 근거 요약',1)
add_table(doc,['항목','v06 근거','계획서 반영'],[
('A1','편의시설 3개+ 0건 85.9%; 최소평점 70.6%; 가격 66.8%','조건 제한 강도별 0건 확률 차등'),
('B1','0건 후 후속검색 96.1% vs 77.9%; OR 7.07; p<.001','후속검색 전이 강화'),
('B2','7.43 vs 7.88회; p=.953','기각·생성 규칙 제외'),
('B3','즉시 성공 22.5%; 세션 최종 회복 71.4%','두 회복지표 분리'),
('B4/B5','동일반복 즉시 0%; 완화 37.5%; 강화 0%; 혼합 33.3%','행동유형별 회복 확률 차등'),
('C3','필요 데이터 부족, SEARCH·SEARCH_FILTER보다 우선순위 낮음','현 버전 제외'),
], [1200,4380,3780], 8.3)

# keep rows together where practical and mark core properties
props=doc.core_properties; props.title='호텔검색 데이터 증강 계획서'; props.subject='v06 기반 세그먼트·가설 중심 증강 실행안'; props.author='2팀'
OUT.parent.mkdir(parents=True, exist_ok=True); doc.save(OUT)
print(OUT)
