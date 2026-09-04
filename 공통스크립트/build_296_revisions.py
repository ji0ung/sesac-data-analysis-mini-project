from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(__file__).resolve().parents[1]
HYP=ROOT/'04_분석설계'/'팀프로젝트'/'2026'/'09'/'호텔검색_가설문서_결과없음_재검색_세그먼트_20260902_v08_현행본_296건기준.docx'
PLAN=ROOT/'04_분석설계'/'팀프로젝트'/'2026'/'09'/'호텔검색_데이터증강계획서_세그먼트가설중심_20260902_v06_현행본_296건기준.docx'
DB='travel_data_filtered_complete_2026-09-01_v01_원본.sqlite'
BLUE='2E74B5'; DARK='1F4D78'; NAVY='0B2545'; PALE='E8EEF5'; LIGHT='F4F6F9'; GOLD='7A5A00'; RED='9B1C1C'; GREEN='1F5F4A'; GRAY='666666'

def font(r,size=11,bold=False,color='000000'):
    r.font.name='Calibri'; r.font.size=Pt(size); r.bold=bold; r.font.color.rgb=RGBColor.from_string(color)
    rf=r._element.get_or_add_rPr().get_or_add_rFonts(); rf.set(qn('w:ascii'),'Calibri'); rf.set(qn('w:hAnsi'),'Calibri'); rf.set(qn('w:eastAsia'),'맑은 고딕')
def shade(c,fill):
    x=OxmlElement('w:shd'); x.set(qn('w:fill'),fill); c._tc.get_or_add_tcPr().append(x)
def geom(t,widths):
    t.autofit=False; t.alignment=WD_TABLE_ALIGNMENT.LEFT; pr=t._tbl.tblPr
    for tag,val in [('tblW',9360),('tblInd',120)]:
        x=OxmlElement('w:'+tag); x.set(qn('w:w'),str(val)); x.set(qn('w:type'),'dxa'); pr.append(x)
    g=t._tbl.tblGrid
    for x in list(g):g.remove(x)
    for w in widths:
        x=OxmlElement('w:gridCol'); x.set(qn('w:w'),str(w)); g.append(x)
    for row in t.rows:
        for i,c in enumerate(row.cells):
            c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tcw=c._tc.get_or_add_tcPr().get_or_add_tcW(); tcw.set(qn('w:w'),str(widths[i])); tcw.set(qn('w:type'),'dxa')
            mar=OxmlElement('w:tcMar')
            for tag,val in [('top',80),('bottom',80),('start',120),('end',120)]:
                x=OxmlElement('w:'+tag); x.set(qn('w:w'),str(val)); x.set(qn('w:type'),'dxa'); mar.append(x)
            c._tc.get_or_add_tcPr().append(mar)
def table(d,heads,rows,widths,size=8.3):
    t=d.add_table(rows=1,cols=len(heads)); t.style='Table Grid'
    for i,h in enumerate(heads):
        shade(t.rows[0].cells[i],PALE); p=t.rows[0].cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run(h),size,True,NAVY)
    rep=OxmlElement('w:tblHeader'); rep.set(qn('w:val'),'true'); t.rows[0]._tr.get_or_add_trPr().append(rep)
    for ri,row in enumerate(rows):
        cs=t.add_row().cells
        for i,v in enumerate(row):
            if ri%2:shade(cs[i],'FAFBFC')
            p=cs[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.08; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if i==0 else WD_ALIGN_PARAGRAPH.LEFT
            font(p.add_run(str(v)),size,False,'222222')
    geom(t,widths); d.add_paragraph()
def callout(d,label,text,fill=LIGHT,color=NAVY):
    t=d.add_table(rows=1,cols=1); c=t.cell(0,0); shade(c,fill); p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    font(p.add_run(label+'  '),10.3,True,color); font(p.add_run(text),10.3); geom(t,[9360]); d.add_paragraph()
def bullet(d,text):
    p=d.add_paragraph(style='List Bullet'); font(p.add_run(text),10.4); p.paragraph_format.left_indent=Inches(.375); p.paragraph_format.first_line_indent=Inches(-.188); p.paragraph_format.space_after=Pt(4)
def base(title,subtitle,version):
    d=Document(); s=d.sections[0]; s.page_width=Inches(8.5); s.page_height=Inches(11); s.top_margin=s.bottom_margin=s.left_margin=s.right_margin=Inches(1); s.header_distance=s.footer_distance=Inches(.492)
    st=d.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(11); st._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); st.paragraph_format.space_after=Pt(6); st.paragraph_format.line_spacing=1.25
    for name,size,bef,aft,col in [('Heading 1',16,18,10,BLUE),('Heading 2',13,14,7,BLUE),('Heading 3',12,10,5,DARK)]:
        x=d.styles[name]; x.font.name='Calibri'; x.font.size=Pt(size); x.font.bold=True; x.font.color.rgb=RGBColor.from_string(col); x._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); x.paragraph_format.space_before=Pt(bef); x.paragraph_format.space_after=Pt(aft); x.paragraph_format.keep_with_next=True
    font(s.header.paragraphs[0].add_run('TEAM 2 · HOTEL SEARCH ANALYSIS'),8.5,True,GRAY); s.footer.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(s.footer.paragraphs[0].add_run(version+' · 296건 원본 DB 기준'),8.5,False,GRAY)
    p=d.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4); font(p.add_run(title),22,True,NAVY)
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(12); font(p.add_run(subtitle),13,False,DARK)
    for a,b in [('기준 DB',DB),('분석 단위','SEARCH 296건 · 검색 세션 43개 · 검색 사용자 매핑 41명'),('작성일','2026-09-02'),('문서 버전',version)]:
        p=d.add_paragraph(); p.paragraph_format.space_after=Pt(2); font(p.add_run(a+': '),10.3,True); font(p.add_run(b),10.3)
    d.add_paragraph(); return d

# Hypothesis v07
d=base('호텔검색 가설문서','결과 없음 · 재검색 · 세그먼트 — v06 전체 구조 계승, 296건 재판정','v08 통합완성본')
callout(d,'최종 메시지','제한적인 필터는 결과 없음과 강하게 관련되고, 결과 없음 경험 세션은 검색도 더 길다. 다만 이는 인과가 아닌 관찰 결과이므로 증강에서는 사용자 탐색지속성과 필터 제한을 분리해 생성한다.','EAF2F8')
d.add_heading('1. 변경 이유와 데이터 기준',level=1)
bullet(d,'기존 v06의 SEARCH 608건(team_i_analysis_20260901.db) 대신 팀원이 올린 원본 SQLite의 SEARCH 296건으로 전면 재계산했다.')
bullet(d,'서로 다른 DB의 수치가 섞이지 않도록 v06은 변경 이력으로 보존하고 본 문서를 새 기준본으로 사용한다.')
bullet(d,'SEARCH_FILTER는 SEARCH와 296:296으로 1:1이며, 검색 세션은 43개다.')
callout(d,'사용자 매핑 주의','search_submit EVENT로 SEARCH 296건은 모두 연결되지만 고유 user_id는 41명이다. 세션은 43개이므로 사용자 분석과 세션 분석을 섞지 않는다.','FFF4E5',GOLD)
d.add_heading('2. 핵심 수치 재검산',level=1)
table(d,['지표','분자/분모','결과','이전 608건 값'],[
('전체 0건률','147/296','49.7%','310/608=51.0%'),
('0건 후 후속검색률','140/147','95.2%','298/310=96.1%'),
('다음 검색 즉시 회복률','24/140','17.1%','67/298=22.5%'),
('세션 최종 회복률','21/28','75.0%','30/42=71.4%'),
('검색 세션','43','43개','78개'),
],[2100,1850,1850,3560],8.5)
d.add_heading('3. A계열 — 결과 없음 가설',level=1)
d.add_heading('3.1 A1 제한 조건 가설',level=2)
table(d,['조건','설정군 0건률','비교군 0건률','OR','Fisher p-value','판정'],[
('편의시설 3개 이상','120/136=88.2%','27/160=16.9%','36.94','2.17×10⁻³⁷','부분 채택'),
('최소평점 설정','115/152=75.7%','32/144=22.2%','10.88','7.16×10⁻²¹','부분 채택'),
('가격 설정','106/146=72.6%','41/150=27.3%','7.05','5.80×10⁻¹⁵','부분 채택'),
],[1600,1900,1900,1000,1700,1260],8.0)
callout(d,'A1 통일 판정','세 조건 모두 방향과 통계적 차이는 뚜렷하지만 오타·중간입력 여부를 확정할 컬럼이 없으므로 “부분 채택”으로 통일한다. 문서 전 구간에서 상위 H1을 채택으로 올려 쓰지 않는다.')
d.add_heading('3.2 A3 오타형 선결조건',level=2)
bullet(d,'현재 DB에는 typo_flag 또는 입력완료 상태 컬럼이 없다. 빈 query_text 108건은 destination으로 검색한 정상 검색일 수 있어 오타로 간주하지 않는다.')
bullet(d,'keyword_search_error_exposure 5건은 모두 결과가 1건 이상이어서 오타 0건 집단의 대리값으로 사용할 수 없다.')
bullet(d,'따라서 A3는 “보류”가 아니라 A1 해석의 미해결 전제다. 원본 입력상태 로그를 확보하거나 명시적 판정 규칙을 승인받은 뒤 제외 전후를 재계산한다.')
d.add_heading('3.3 A2 지역·의도 가설',level=2)
callout(d,'상태','가설 포함 승인. 단, 지역·의도 표본 정의를 먼저 제출한 뒤 최종 확정한다. 아래 분류는 현재 DB에 적용 가능한 1차 정의다.','E2F0D9',GREEN)
table(d,['지역','검색 n','0건 n','0건률'],[
('Tokyo','118','59','50.0%'),('Osaka','80','49','61.3%'),('Kyoto','17','6','35.3%'),('Sapporo','21','3','14.3%'),('Fukuoka','15','3','20.0%'),('UNKNOWN','45','27','60.0%')],[2200,1800,1800,3560],8.5)
table(d,['의도 코드','정의','검색 n','0건률'],[
('LOCATION_ONLY','가격·평점·편의시설 조건 없음','84','7/84=8.3%'),('PRICE','가격만 설정','9','1/9=11.1%'),('QUALITY_FILTER','최소평점만 설정 — C3 세그먼트와는 별개','3','1/3=33.3%'),('AMENITY','편의시설만 설정','31','13/31=41.9%'),('MIXED','가격·평점·편의시설 중 2개 이상','169','125/169=74.0%')],[2000,4300,1200,1860],8.3)
callout(d,'A2 해석 제한','지역별 차이에는 의도 구성 차이가 섞여 있다. 지역×의도 셀 중 표본이 매우 작은 곳이 있으므로 지금 단계에서는 기술통계만 제시하고 인과·우열을 단정하지 않는다.')
d.add_heading('4. B계열 — 재검색 가설',level=1)
table(d,['가설','296건 결과','검정','판정','증강 반영'],[
('B1 0건 후 후속검색','95.2% vs 비0건 75.8%','OR 6.37; p=1.81×10⁻⁶','채택','0건 후 전이확률 차등'),
('B2 0건 경험 세션 검색횟수','평균 9.00회 vs 2.93회','Mann–Whitney U=382; p=1.12×10⁻⁵','채택·주의','탐색지속성 매개로만 반영'),
('B3 즉시/최종 회복','17.1% vs 75.0%','분모 140 / 세션 28','채택','두 지표 분리'),
],[2200,2700,2300,1200,960],8.0)
callout(d,'B2 변경','608건 DB에서는 기각(p=.953)이었으나 296건 DB에서는 지지된다. 데이터 버전에 따라 결론이 바뀌었으므로 “0건이 검색을 늘린다”는 인과 문장은 금지하고, 탐색 의지가 높은 사용자가 0건도 재검색도 더 경험했을 가능성을 함께 적는다.','FFF4E5',GOLD)
d.add_heading('5. 클릭·상세조회 검증',level=1)
bullet(d,'hotel_click 231건과 hotel_detail_view 231건의 (session_id, search_id, hotel_id) 조합이 완전히 동일하다.')
bullet(d,'따라서 기존처럼 클릭률과 상세조회율을 별도 성과로 제시하면 중복 지표가 된다.')
callout(d,'조치','현재 데이터에서는 두 이벤트를 “호텔 상세진입” 1개 지표로 합친다. 퍼널에는 hotel_click을 대표 이벤트로 사용하고 hotel_detail_view는 동반 로그로 주석 처리한다.')
d.add_heading('6. 세션 결과 세그먼트 재산정',level=1)
table(d,['세그먼트','세션 n','비율','296건 기준 정의'],[
('직접 성공','27','62.8%','첫 검색 비0건이며 세션 내 호텔 클릭'),('결과 노출·미선택','10','23.3%','첫 검색 비0건이나 호텔 클릭 없음'),('재검색 회복','4','9.3%','첫 검색 0건 후 이후 검색에서 결과 발생'),('지속 실패','2','4.7%','첫 검색 0건, 재검색 후에도 결과 없음'),('0건 즉시 이탈','0','0.0%','첫 검색 0건 후 검색 종료')],[2200,1200,1300,4660],8.3)
callout(d,'비율 변경','이전 실패군 17.6%는 608건 DB의 값이다. 296건 DB에서는 지속 실패 4.7%, 0건 즉시 이탈 0%이므로 이전 비율을 증강 목표로 사용하지 않는다. 세션 표본 n=43으로 작아 파일럿 후 재검증한다.')
d.add_heading('7. 사용자 검색 성향 분류',level=1)
table(d,['검색 성향','판정 기준','관찰 행동','분석·증강 활용'],[
('지역 고정형','동일 도시·지역을 유지하며 조건 변경','destination/region 유지','지역 유지 확률'),('가격 민감형','가격 설정·변경·가격정렬 사용','price·sort_option 변화','가격 완화/강화 확률'),('품질 우선형','최소평점 중심 탐색','user_rating_min 사용','C3 제외: 기술통계만'),('테마·편의형','편의시설 조건 중심','amenity_count 변화','편의시설 조건 전이'),('호텔 지명형','특정 숙소명을 검색','query_text와 호텔명 매칭','사전 구축 후 적용'),('비교 탐색형','다회 검색·다수 결과 비교','검색 횟수·클릭 시점','탐색지속성 분포'),('빠른 결정형','적은 검색 후 클릭','첫 검색·초기 클릭','짧은 경로 생성'),('입력 탐색형','검색어 수정·중간입력','query_text 변화','A3 승인 후 별도 적용')],[1700,2500,2400,2760],7.8)
callout(d,'분류 원칙','검색 성향은 다중 라벨이며 결과 세그먼트와 다르다. 품질 우선형 C3는 현재 사용자 세그먼트 가설에서 제외하지만 QUALITY_FILTER라는 검색 1회 의도 코드는 A2 기술통계에만 유지한다.')
d.add_heading('8. 첫 실패 이후 재검색 유형',level=1)
table(d,['재검색 유형','사례 n','다음 검색 성공','상세진입','판정'],[
('동일조건 반복','53','0/53=0.0%','0/53=0.0%','회복에 불리'),('조건 완화','41','11/41=26.8%','8/41=19.5%','회복 가능'),('검색어 수정','10','3/10=30.0%','3/10=30.0%','소표본'),('지역 변경','24','10/24=41.7%','3/24=12.5%','성공률 최고'),('조건 강화','10','0/10=0.0%','0/10=0.0%','회복에 불리'),('혼합 변경','2','0/2=0.0%','0/2=0.0%','판단 불가')],[1800,1200,2300,2200,1860],8.0)
callout(d,'집계 범위','0건 뒤 같은 세션의 바로 다음 검색이 있는 140건을 중복 없이 분류했다. 이는 사용자 구성비가 아니라 “0건 검색 1행 → 다음 검색 1행” 전이 단위다.')
d.add_heading('9. 상위가설 H3 — 재검색 방식과 선택',level=1)
table(d,['하위가설','현재 결과','판정','해석'],[
('H3-1 재검색 방식별 결과 회복 차이','지역 변경 41.7%, 검색어 수정 30.0%, 완화 26.8%, 반복·강화 0%','탐색적 채택','유형별 n과 규칙 정의에 민감'),('H3-2 재검색 방식별 상세진입 차이','검색어 수정 30.0%, 완화 19.5%, 지역 변경 12.5%','탐색적 채택','클릭·상세조회는 한 지표'),('H3-3 재검색이 선택을 높인다','자기선택·탐색의지 통제 불가','인과 보류','제품 효과로 단정 금지'),('H3-4 예약완료 차이','BOOKING 36건, 실측 근거 제한','참고','핵심 성과에서 제외')],[2400,3300,1300,2360],8.0)
d.add_heading('10. 재검색 전환 단위별 상세 결과',level=1)
table(d,['재검색 유형','사례 수','결과 발생률','호텔 상세진입률','발표 주의'],[
('지역 변경','24','41.7%','12.5%','n=24 명시'),('검색어 수정','10','30.0%','30.0%','n=10 소표본'),('조건 완화','41','26.8%','19.5%','단계적 완화 후보'),('조건 강화','10','0.0%','0.0%','조건 정의 확인'),('동일조건 반복','53','0.0%','0.0%','구조적으로 자명할 수 있음'),('혼합 변경','2','0.0%','0.0%','해석 금지')],[1800,1100,1800,2100,2560],8.0)
callout(d,'보조 발견 처리','“동일조건 반복은 또 0건”은 정보량이 낮아 제품 가설의 핵심 근거로 과장하지 않는다. 비교 기준선으로만 보존한다.')
d.add_heading('11. 가설 검증 종합',level=1)
table(d,['계층','판정','핵심 결과','발표 메시지'],[
('H1 제한조건','부분 채택','세 조건 OR 7.05~36.94; 오타 기준 미확정','조건 제한과 0건은 강하게 관련'),('H2 결과없음·재검색','채택·주의','후속검색 OR 6.37; B2 p<.001','연관은 있으나 인과 아님'),('H3 재검색 방식','탐색적 채택','지역변경·검색어수정·완화에서 회복','유형별 n과 정의를 함께 제시'),('C3 품질 세그먼트','제외','사용자 세그먼트 근거 부족','현 단계 미진행')],[1900,1600,3400,2460],8.0)
d.add_heading('12. 증강 데이터 생성 규칙과의 연결',level=1)
table(d,['가설 결과','증강 규칙','적용 여부'],[
('A1 부분 채택','필터 제한 수준별 0건 확률 차등','조건부 적용'),('A2 조건부 진행','지역×의도 표본층 생성','교수 승인 후'),('A3 미해결','typo_flag 확보 전 핵심 분포에서 분리','보류'),('B1 채택','0건 후 후속검색 전이 강화','적용'),('B2 채택·주의','탐색지속성 선생성; 0건 후 횟수 강제가산 금지','간접 적용'),('B3 채택','즉시 회복과 최종 회복 분리','적용'),('H3 탐색적','재검색 유형별 회복·상세진입 확률 차등','민감도 적용'),('C3 제외','품질 우선 사용자 세그먼트 미생성','적용 제외')],[2500,4500,2360],8.0)
d.add_heading('13. 최종 제품 가설',level=1)
callout(d,'제품 가설','검색 결과가 0건인 사용자에게 결과가 존재하는 같은 도시의 인접 지역과 단계적으로 완화 가능한 조건을 제시하면, 동일조건 반복을 줄이고 다음 검색의 결과 회복 및 호텔 상세진입을 높일 수 있을 것이다.','E2F0D9',GREEN)
table(d,['우선순위','성공지표','정의'],[
('1','다음 검색 회복률','0건 뒤 다음 검색 total_result_count>0'),('2','세션 최종 회복률','첫 0건 이후 세션 내 결과 발생'),('3','호텔 상세진입률','hotel_click 대표 이벤트'),('보조','예약완료율','실제/합성 분리된 참고지표')],[1400,2600,5360],8.2)
d.add_heading('14. 해석 및 발표 시 주의사항',level=1)
for x in ['296건 DB의 검색 세션은 43개로 작으므로 세션 비율에는 n을 함께 표시한다.','지역×의도 희소셀은 기술통계이며 우열·인과로 말하지 않는다.','빈 query_text는 destination 검색일 수 있어 오타로 분류하지 않는다.','hotel_click과 hotel_detail_view는 같은 대상의 동반 이벤트이므로 이중 성과로 제시하지 않는다.','B2는 데이터 버전에 따라 판정이 바뀌었으므로 탐색성향이라는 교란 가능성을 반드시 언급한다.','BOOKING은 참고지표이며 실제 예약 성과처럼 발표하지 않는다.']:
    bullet(d,x)
d.add_heading('15. 최종 판정표',level=1)
table(d,['가설','상태','핵심 근거','발표 문장'],[
('A1','부분 채택','필터별 OR 7.05~36.94; 오타 기준 미확정','제한 조건과 0건은 강하게 관련'),('A2','조건부 진행','지역·의도 정의 제출 선행','지역·의도별 차이를 탐색'),('A3','선결조건 미해결','명시적 오타 컬럼 없음','오타 정의 후 재검산 필요'),('B1','채택','OR 6.37; p<.001','0건 뒤 후속검색 가능성이 높음'),('B2','채택·주의','9.00 vs 2.93; p<.001','연관은 있으나 인과 아님'),('B3','채택','즉시 17.1%, 최종 75.0%','즉시와 최종 회복을 분리'),('C3','제외 유지','품질 사용자 세그먼트 근거 부족','현 증강 범위 제외')],[1200,1400,3400,3360],8.1)
d.core_properties.title='호텔검색 가설문서 296건 기준 v08 통합완성본'; d.core_properties.author='2팀'; HYP.parent.mkdir(parents=True,exist_ok=True); d.save(HYP)

# Augmentation plan v05
d=base('호텔검색 데이터증강계획서','v04 전체 실행구조 계승 · 296건 원본 DB 및 교수 첨삭 통합','v06 통합완성본')
callout(d,'승인 방향','관측형 1천 명으로 파일럿 후 1만 명까지 확장한다. 과표집형은 필요할 때 별도 세트로 만들며 관측형과 섞지 않는다. A2는 포함하되 지역×의도 표본 정의를 먼저 제출한다.','EAF2F8')
d.add_heading('1. 기준과 추적 컬럼',level=1)
bullet(d,'기준 DB: 03_데이터모델링/이전버전/데이터셋/2026-09-03_v02/'+DB+' — USER 89, SEARCH 296, SEARCH_FILTER 296, 검색 세션 43.')
bullet(d,'기존 v04의 608건 기반 비율은 히스토리로만 보존하고 생성 파라미터에서 제거한다.')
table(d,['구분','필수 컬럼','목적'],[
('출처·시나리오','data_origin, scenario_id, generation_version','실제/합성 및 생성 버전 분리'),('표본설계','sample_design_id, sample_set_type, sample_stratum','관측형/과표집형과 층 구분'),('가중치','target_population_share, sample_share, selection_probability, sample_weight, weight_version','과표집 결과의 실제비율 오해 방지'),('재현성','random_seed, config_version','동일 조건 재생성')],[1700,4300,3360],8.1)
d.add_heading('2. 생성 규모와 표본 원칙',level=1)
table(d,['단계','규모','세트','규칙','완료 기준'],[
('파일럿','1,000명','관측형 1종','296건 관측분포를 기본값으로 사용','논리·분포 QA 통과'),('본 생성','10,000명','관측형 기본','파일럿 승인 설정을 고정해 확장','버전·seed 기록'),('선택 실험','별도','과표집형','관측형과 파일·지표·발표를 분리','가중치와 과표집 표시')],[1100,1000,1500,3600,2160],8.2)
callout(d,'실패군 최소 50명 문제','296건 기준 지속 실패는 2/43=4.7%이므로 1천 명에서 약 47명이다. 최소 50명선은 소폭 조정으로 충족할 수 있다. 그러나 0건 즉시 이탈은 관측 0명이므로 관측형에 50명을 임의 생성하지 않는다. 필요하면 별도 스트레스 세트로 만든 뒤 교수님 승인을 받는다.','FFF4E5',GOLD)
d.add_heading('3. 세그먼트 생성 목표',level=1)
table(d,['세그먼트','관측 n/43','목표 비율','파일럿 기대','처리'],[
('직접 성공','27','62.8%','약 628명','관측형'),('결과 노출·미선택','10','23.3%','약 233명','관측형'),('재검색 회복','4','9.3%','약 93명','관측형'),('지속 실패','2','4.7%','약 47명 → 최소 50 검토','관측형 소폭 보정'),('0건 즉시 이탈','0','0.0%','0명','별도 승인 전 생성 금지')],[2000,1350,1350,2600,2060],8.2)
d.add_heading('4. A2 지역·의도 표본 정의안',level=1)
bullet(d,'지역은 destination의 앞부분을 Tokyo, Osaka, Kyoto, Sapporo, Fukuoka, UNKNOWN으로 표준화한다.')
bullet(d,'의도는 SEARCH_FILTER에서 가격·최소평점·편의시설 설정 여부를 이용해 LOCATION_ONLY, PRICE, QUALITY_FILTER, AMENITY, MIXED로 분류한다.')
bullet(d,'QUALITY_FILTER는 검색 1회 의도 코드일 뿐, 제외된 C3 “품질 우선 사용자 세그먼트”를 부활시키지 않는다.')
table(d,['지역','관측 검색 n','파일럿 검색 할당 원칙','본 생성 원칙'],[
('Tokyo','118 / 39.9%','전체 생성 검색의 39.9%','동일 비율'),('Osaka','80 / 27.0%','27.0%','동일 비율'),('Kyoto','17 / 5.7%','5.7%','동일 비율'),('Sapporo','21 / 7.1%','7.1%','동일 비율'),('Fukuoka','15 / 5.1%','5.1%','동일 비율'),('UNKNOWN','45 / 15.2%','15.2%; 결측 원인 표시','동일 비율·별도 보고')],[1800,2000,3000,2560],8.2)
table(d,['의도','관측 n','비율','파일럿 규칙'],[
('LOCATION_ONLY','84','28.4%','관측비율'),('PRICE','9','3.0%','최소 셀 수 확인'),('QUALITY_FILTER','3','1.0%','희소셀; 기술통계만'),('AMENITY','31','10.5%','관측비율'),('MIXED','169','57.1%','관측비율')],[2200,1400,1400,4360],8.3)
callout(d,'교수님 제출 항목','지역×의도 교차표는 원본 n/0건 n을 함께 제출한다. PRICE·QUALITY_FILTER처럼 n이 작은 셀은 병합 또는 별도 과표집 여부를 교수님이 결정한다.')
d.add_heading('5. 가설별 생성 규칙',level=1)
table(d,['가설','296건 판정','생성 규칙','금지·주의'],[
('A1','부분 채택','필터 제한이 강할수록 0건 확률 상승','오타 제거 전 확정값으로 단정 금지'),('A2','조건부 포함','지역×의도 층별 분포와 0건률 적용','희소셀 수치 고정 금지'),('A3','선결조건 미해결','별도 typo_flag 확보 전 핵심 생성에서 제외','빈 검색어를 오타로 간주 금지'),('B1','채택','0건 후 후속검색 전이 강화','세션 종료 가능성은 별도'),('B2','채택·주의','사용자 탐색지속성을 먼저 생성하고 0건 경로와 연결','0건 발생 후 검색횟수 강제 가산 금지'),('B3','채택','즉시 회복 17.1%, 최종 회복 75.0%를 분리','서로 같은 지표로 사용 금지'),('C3','제외','사용자 품질우선 세그먼트 생성 안 함','QUALITY_FILTER 검색 의도와 혼동 금지')],[1100,1500,3800,2960],8.0)
d.add_heading('6. 이벤트·성과지표 정리',level=1)
table(d,['항목','원본 확인','생성·분석 처리'],[
('hotel_click','231건','상세진입 대표 이벤트'),('hotel_detail_view','231건; 클릭과 키 조합 완전 동일','동반로그로 생성하되 별도 전환율로 발표하지 않음'),('BOOKING','36건','참고 시나리오; 실제/합성 분리'),('0건 SEARCH','147건','SEARCH_RESULT 0행 유지')],[2200,3200,3960],8.3)
d.add_heading('7. 제작팀 실행',level=1)
table(d,['작업','권장 담당','산출물','완료 기준'],[
('지역·의도 코드 SQL','박지현','A2 교차표·코드 사전','296건 전부 단일 코드'),('표본설계·가중치','권순성','sample design 설정표','비율·가중치 검산'),('생성기 수정','손지영','v05 설정 기반 파일럿','seed 재현 가능'),('QA·문서 정합성','이소이','가설-계획-결과 검증표','608 수치 잔존 0건'),('공동 검토','전원','교수 판단표','미결정 항목 명시')],[1700,1400,3000,3260],8.2)
d.add_heading('8. 교수님 판단 필요사항',level=1)
table(d,['결정','현 상태','권장안','확인'],[
('A2 지역×의도 희소셀','PRICE 9, QUALITY_FILTER 3','기술통계 유지; 과표집은 별도 세트','☐'),('A3 오타 정의','명시적 컬럼 없음','typo_flag 수집 전 미판정','☐'),('지속 실패 최소 50','기대 약 47명','50명으로 소폭 보정 승인','☐'),('0건 즉시 이탈 최소 50','관측 0명','관측형 생성 금지; 별도 스트레스 세트만','☐'),('허용오차','기존 권장 핵심 ±3%p','그대로 승인','☐'),('BOOKING','실측 근거 제한','참고 시나리오 유지','☐')],[2400,2700,3460,800],8.1)
d.add_heading('9. QA 체크리스트',level=1)
for x in ['296건 DB 외 수치가 생성 설정에 남지 않았는가?','SEARCH와 SEARCH_FILTER가 1:1인가?','표본설계·가중치 컬럼이 모든 합성 사용자에 있는가?','관측형과 과표집형 파일이 물리적으로 분리됐는가?','B2를 인과 규칙으로 구현하지 않았는가?','hotel_click과 hotel_detail_view를 이중 성과로 집계하지 않는가?','A3 미확정 검색을 오타로 임의 분류하지 않는가?','지역×의도 희소셀을 교수 승인 없이 확대하지 않았는가?']:
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(5); font(p.add_run('☐ '+x),10.4)
d.add_heading('10. 실행 순서',level=1)
for x in ['1) 교수 판단표 확정 및 A2 교차표 제출','2) 설정 파일·스키마·가중치 계산 동결','3) 관측형 1천 명 생성','4) 구조·분포·가설방향 QA','5) 승인 후 관측형 1만 명 확장','6) 필요 시 과표집형 별도 생성·가중 분석']:
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(5); font(p.add_run(x),10.4)
d.add_heading('11. 생성 파이프라인 상세',level=1)
table(d,['순서','테이블','생성 내용','검증'],[
('1','USER','결과 세그먼트 1개·검색성향 다중라벨·표본설계·가중치','user_id·세트 구분'),('2','SEARCH','탐색지속성에 따라 세션별 검색 횟수·시간·의도 생성','B2 강제가산 없음'),('3','SEARCH_FILTER','검색당 1행, 가격·평점·편의시설·지역 생성','SEARCH 1:1'),('4','SEARCH_RESULT','조건·시나리오별 0건 또는 기존 HOTEL·ROOM 참조','0건은 0행'),('5','EVENT','검색·노출·상세진입·후속행동을 시간순 생성','클릭 대상 결과 존재'),('6','BOOKING','상세진입 이후 참고 시나리오만 생성','실제/합성 분리')],[700,1400,4700,2560],8.0)
d.add_heading('12. 실험 시나리오',level=1)
table(d,['ID','시나리오','주요 변경','확인 가설','비교 지표'],[
('S0','관측형 기준군','296건 분포·전이 유지','A1,B1,B2,B3','0건·회복·상세진입'),('S1','인접 지역 제안','지역 고정 실패 시 인접 후보','A2,H3','다음 검색 회복'),('S2','필터 완화 제안','가격·평점·편의시설 단계 완화','A1,H3','즉시·최종 회복'),('S3','통합 제안','인접 지역+필터 완화','A1,A2,H3','회복·상세진입'),('SX','입력 정정','명시적 typo_flag 대상만','A3','수정 후 회복')],[700,1700,3000,1500,2460],8.0)
d.add_heading('13. 산출물 및 인계 목록',level=1)
for x in ['표본설계표와 지역×의도 교차표','생성 설정 파일(config/version/seed)','USER~BOOKING 합성 테이블','데이터 사전과 intent_code 정의서','가중치 산식 및 관측형·과표집형 구분표','구조·분포·가설방향 QA 결과','변경 이력과 교수 승인 기록']:
    bullet(d,x)
d.add_heading('부록 A. 296건 판정 근거',level=1)
table(d,['항목','근거','계획 반영'],[
('A1','편의 3+ 88.2%, 평점 75.7%, 가격 72.6%; 모두 p<.001','조건 제한별 확률 차등'),('B1','95.2% vs 75.8%; OR 6.37, p<.001','후속검색 전이 차등'),('B2','9.00 vs 2.93회; p<.001','탐색지속성 매개'),('B3','즉시 17.1%, 최종 75.0%','회복지표 분리'),('H3','지역변경 41.7%, 검색어수정 30.0%, 완화 26.8%','유형별 민감도'),('클릭/상세','231건씩 키 조합 동일','상세진입 1개 지표'),('C3','사용자 세그먼트 근거 부족','제외 유지')],[1600,4300,3460],8.1)
d.core_properties.title='호텔검색 데이터증강계획서 296건 기준 v06 통합완성본'; d.core_properties.author='2팀'; d.save(PLAN)
print(HYP); print(PLAN)
