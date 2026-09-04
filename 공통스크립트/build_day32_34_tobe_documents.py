from pathlib import Path
from shutil import copy2
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SRC32 = ROOT / "07_교육과제/팀과제/2026/09/호텔검색_32일차_2팀_실행체크리스트_20260902_v03_현행본_데이터증강반영.docx"
SRC33 = ROOT / "04_분석설계/팀프로젝트/2026/09/BI시각화_33일차_2팀프로젝트실습_해답_20260903_v03_제출본.docx"
SRC34 = ROOT / "07_교육과제/팀과제/2026/09/BI시각화_34일차_2팀프로젝트실습_해답_20260904_v01_제출본.docx"
OUT32 = SRC32.with_name(SRC32.stem + "_[TO-BE].docx")
OUT33 = SRC33.with_name(SRC33.stem + "_[TO-BE].docx")
OUT34 = SRC34.with_name(SRC34.stem + "_[TO-BE].docx")
BRIEF_OUT = ROOT / "04_분석설계/팀프로젝트/2026/09/BI시각화_32-34일차_TO-BE_슬랙공유_브리핑_20260904.docx"

BLUE="2E74B5"; NAVY="0B2545"; PALE="E8EEF5"; LIGHT="F2F4F7"; RED="9B1C1C"; GRAY="666666"

def set_font(run,size=10,bold=False,color="222222"):
    run.font.name="Calibri"; rpr=run._element.get_or_add_rPr(); rf=rpr.get_or_add_rFonts()
    rf.set(qn("w:ascii"),"Calibri"); rf.set(qn("w:hAnsi"),"Calibri"); rf.set(qn("w:eastAsia"),"맑은 고딕")
    run.font.size=Pt(size); run.bold=bold; run.font.color.rgb=RGBColor.from_string(color)

def shade(cell,fill):
    pr=cell._tc.get_or_add_tcPr(); n=pr.find(qn("w:shd"))
    if n is None: n=OxmlElement("w:shd"); pr.append(n)
    n.set(qn("w:fill"),fill)

def margins(cell,top=90,bottom=90,start=120,end=120):
    pr=cell._tc.get_or_add_tcPr(); mar=pr.first_child_found_in("w:tcMar")
    if mar is None: mar=OxmlElement("w:tcMar"); pr.append(mar)
    for tag,val in [("top",top),("bottom",bottom),("start",start),("end",end)]:
        n=mar.find(qn("w:"+tag))
        if n is None: n=OxmlElement("w:"+tag); mar.append(n)
        n.set(qn("w:w"),str(val)); n.set(qn("w:type"),"dxa")

def geometry(t,widths):
    t.autofit=False; t.alignment=WD_TABLE_ALIGNMENT.LEFT; pr=t._tbl.tblPr
    for tag,val in [("tblW",sum(widths)),("tblInd",120)]:
        n=pr.find(qn("w:"+tag))
        if n is None: n=OxmlElement("w:"+tag); pr.append(n)
        n.set(qn("w:w"),str(val)); n.set(qn("w:type"),"dxa")
    grid=t._tbl.tblGrid
    for n in list(grid): grid.remove(n)
    for w in widths:
        n=OxmlElement("w:gridCol"); n.set(qn("w:w"),str(w)); grid.append(n)
    for row in t.rows:
        for i,c in enumerate(row.cells):
            tw=c._tc.get_or_add_tcPr().get_or_add_tcW(); tw.set(qn("w:w"),str(widths[i])); tw.set(qn("w:type"),"dxa")
            margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(doc,headers,rows,widths,size=8.3,center=(0,)):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Table Grid"
    for i,h in enumerate(headers):
        shade(t.rows[0].cells[i],PALE); p=t.rows[0].cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after=Pt(0); set_font(p.add_run(h),size,True,NAVY)
    trpr=t.rows[0]._tr.get_or_add_trPr(); rep=OxmlElement("w:tblHeader"); rep.set(qn("w:val"),"true"); trpr.append(rep)
    for ri,row in enumerate(rows):
        cs=t.add_row().cells
        for i,v in enumerate(row):
            if ri%2: shade(cs[i],"FAFBFC")
            p=cs[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if i in center else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.06; set_font(p.add_run(str(v)),size)
    geometry(t,widths); doc.add_paragraph().paragraph_format.space_after=Pt(1)

def heading(doc,text,level=1):
    p=doc.add_paragraph(text,style=f"Heading {level}"); p.paragraph_format.keep_with_next=True; return p

def para(doc,text,size=10.2,color="222222",bold=False):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.15
    set_font(p.add_run(text),size,bold,color); return p

def bullet(doc,text):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.15; set_font(p.add_run(text),10)

def callout(doc,label,text,fill=LIGHT,color=NAVY):
    t=doc.add_table(rows=1,cols=1); c=t.cell(0,0); shade(c,fill); margins(c,130,130,160,160)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); set_font(p.add_run(label+"  "),10.4,True,color); set_font(p.add_run(text),10.4)
    geometry(t,[9360]); doc.add_paragraph().paragraph_format.space_after=Pt(2)

def prepare(src,out,label):
    copy2(src,out); d=Document(out)
    if d.paragraphs:
        p=d.paragraphs[0]
        if "[TO-BE]" not in p.text:
            original=p.text; p.clear(); set_font(p.add_run("[TO-BE] "+original),16,True,NAVY)
    heading(d,"TO-BE 적용 안내",1)
    callout(d,"적용 우선순위",f"이 절 이후의 TO-BE 정의가 기존 본문의 증강·세그먼트·BOOKING 해석보다 우선한다. 기존 본문은 {label} 실습 기록으로 보존한다.",PALE)
    return d

def common_baseline(doc):
    heading(doc,"원본 v03 기준선",2)
    add_table(doc,["항목","원본 수치","용도"],[
        ("검색 사용자","41명","10,000명 증강 모수"),
        ("검색 세션","43개","세션 단위 QA"),
        ("전체/0건 검색","296/147건","검색 기준선"),
        ("0건 경험 사용자","27/41명=65.9%","실험 진입 확률"),
        ("0건 경험 세션","28/43개=65.1%","세션 기준선"),
    ],[2500,2500,4360],8.7,center=(0,1))
    para(doc,"합성 1,000명은 원본 구조 재현 QA에만 사용하고, 세그먼트 비중·생성 확률은 원본 v03을 기준으로 한다.",9.4,RED)

def dual_segments(doc):
    heading(doc,"이중 세그먼트 구조",2)
    add_table(doc,["시점","분류","역할","저장 컬럼"],[
        ("실험 전","필터 상태 + 추정 검색 의도","누구에게 어떤 개선안을 보여줄지 결정","filter_state_segment / intent_segment"),
        ("실험 후","SG1~SG4 검색·전환 성과 결과군","개선안 적용 후 최종 성과 판정","outcome_segment"),
    ],[1400,2900,3300,1760],8.45,center=(0,))
    add_table(doc,["코드","결과군","판정 기준"],[
        ("SG1","직접 성공형","첫 검색 결과≥1, 세션 내 상세진입"),
        ("SG2","결과 노출·미선택형","첫 검색 결과≥1, 상세진입 없음"),
        ("SG3","재검색 회복형","첫 검색 0건 후 후속 검색에서 결과 발생"),
        ("SG4","지속 실패형","첫 검색 0건, 세션 종료까지 미회복"),
    ],[1000,2500,5860],8.55,center=(0,1))
    para(doc,"핵심 판정은 treatment에서 SG3가 증가하고 SG4가 감소하는지이다. SG1~SG4는 실험 전 배정에 사용하지 않는다.",9.4,RED)

def filter_rules(doc):
    heading(doc,"필터 상태·완화 규칙",2)
    add_table(doc,["첫 0건 필터 상태","원본 사용자","비중","개선안"],[
        ("가격+옵션 수 복합 설정","17명","63.0%","가격 확대·옵션 수 감소 순서 검증"),
        ("가격·옵션 수 제한 없음","7명","25.9%","지역 확대·검색어 수정"),
        ("옵션 수만 설정","2명","7.4%","요구 옵션 수 줄이기"),
        ("가격만 설정","1명","3.7%","가격 범위 확대"),
    ],[3400,1500,1200,3260],8.5,center=(0,1,2))
    para(doc,"평점·구체적 편의시설은 증강 세그먼트에서 제외한다. 현재 DB에서 사용할 수 있는 완화 조건은 가격 범위 확대와 요구 옵션 수 감소다.",9.4,RED)

def booking_rules(doc):
    heading(doc,"예약 성과 KPI 준비",2)
    add_table(doc,["확인 결과","처리"],[
        ("BOOKING 36건, booking_complete 36건, 고유 연결 검색 30건","상세진입→예약시작→예약완료 보조 KPI 후보"),
        ("booking.data_origin='전체 가상 예약 시뮬레이션'","팀이 실제 예약으로 확인했다면 출처 값·메타데이터 정정 후 사용"),
        ("중복 의심 6건, BOOKING–ROOM 호텔 불일치 2건","정답 키·중복 규칙 확정 전에는 조건부 KPI"),
        ("BOOKING에 search_id 없음","booking_complete 이벤트의 search_id로 예약 연결 브리지 생성"),
    ],[3900,5460],8.45)

SLACK_BRIEFS=[]

def slack_brief(doc, day, title, summary, points, next_actions):
    """Keep Slack copy outside the three practice documents."""
    SLACK_BRIEFS.append((day,title,summary,points,next_actions))

def save_slack_briefs():
    doc=Document()
    p=doc.add_paragraph()
    set_font(p.add_run("32·33·34일차 TO-BE 슬랙 공유 브리핑"),18,True,NAVY)
    para(doc,"팀 공유 시 일차별 블록을 그대로 복사해 사용하면 됩니다.",10.2,GRAY)
    for day,title,summary,points,next_actions in SLACK_BRIEFS:
        heading(doc,f"{day}일차 공유안",1)
        callout(doc,"공유 주제",title,PALE)
        para(doc,summary,10.0)
        heading(doc,"핵심 내용",2)
        for item in points:
            bullet(doc,item)
        heading(doc,"다음 행동",2)
        for item in next_actions:
            bullet(doc,item)
    doc.save(BRIEF_OUT)

# Day 32
d=prepare(SRC32,OUT32,"32일차")
common_baseline(d); dual_segments(d); filter_rules(d)
heading(d,"TO-BE 가설",2)
add_table(d,["ID","가설","판정 KPI"],[
    ("H1","맞춤 개선안을 받은 실험군은 대조군보다 결과 회복률이 높다.","SG3 비율↑, SG4 비율↓"),
    ("H2","가격 설정 사용자에게 가격 범위 확대를 제안하면 회복률이 높다.","가격 완화 후 회복률"),
    ("H3","옵션 수 설정 사용자에게 요구 옵션 수 감소를 제안하면 회복률이 높다.","옵션 수 완화 후 회복률"),
    ("H4","복합 설정 사용자에게 단계적 완화를 제공하면 잔여 0건률이 낮다.","최종 회복률"),
    ("H5","맞춤 제안은 동일 조건 반복률을 낮춘다.","RT1 비율↓"),
    ("H6","결과 회복은 상세진입·예약완료 증가로 이어진다.","상세진입률·예약완료율"),
],[900,6100,2360],8.25,center=(0,))
booking_rules(d)
slack_brief(d,"32","원본 DB 기준의 10,000명 증강·세그먼트 가설 재설계",
    "32일차 문서를 원본 v03 수치와 이중 세그먼트 구조에 맞춰 TO-BE로 갱신했습니다. 핵심은 실험 전에는 검색 의도·필터 상태로 개선안을 배정하고, 실험 후에는 SG1~SG4로 성과를 판정하는 것입니다.",[
        "원본 검색 사용자 41명 중 0건 경험자는 27명(65.9%)입니다.",
        "첫 0건 시 가격+옵션 수 복합 설정이 17명(63.0%)으로 가장 크므로, 가격 확대와 요구 옵션 수 감소를 주요 가설로 둡니다.",
        "1,000명 합성 DB는 구조 QA용이고, 10,000명 증강 비중과 생성 확률은 원본 v03에서 가져옵니다.",
        "성공 기준은 재검색 회복형 SG3 증가와 지속 실패형 SG4 감소입니다."
    ],[
        "증강 스크립트에 filter_state_segment, intent_segment, outcome_segment 컬럼을 반영합니다.",
        "H1~H6의 판정식·모수·허용 오차를 팀이 함께 확정합니다.",
        "BOOKING은 출처·중복·호텔 연결을 검증한 뒤 보조 KPI로 사용합니다."
    ]); d.save(OUT32)

# Day 33
d=prepare(SRC33,OUT33,"33일차")
common_baseline(d); dual_segments(d); filter_rules(d)
heading(d,"TO-BE 인사이트 카드 역할",2)
add_table(d,["기존 카드","보존할 관측 근거","TO-BE 역할"],[
    ("A","0건이 제한 조건에 집중","가격·옵션 수 완화 가설의 기준선"),
    ("B","0건 후 후속 검색 140/147=95.2%","이탈 로그 보강 게이트"),
    ("C","즉시 회복 17.1%, 세션 최종 회복 75.0%","회복 시점별 KPI 분리"),
    ("D","지역·검색어·조건 완화별 차이","제안 유형별 효과 가정 근거"),
    ("F","동일 조건 0/53, 조건 강화 0/10 회복","반복 방지 가설 근거"),
    ("G","직접·재검색 후·미상호작 구조","성공 경로 QA; 이탈률로 사용 금지"),
    ("H","고유 1위 클릭 29/149=19.5%","결과 품질·순위 보정 지표"),
],[1200,3900,4260],8.2,center=(0,))
para(d,"기존 카드는 삭제하지 않고 '원본 관측 기준선'으로 보존한다. TO-BE 카드에는 실험 전 의도·필터 상태, 제공 시나리오, SG1~SG4 결과, 95% 불확실성 구간을 추가한다.",9.4,RED)
booking_rules(d)
slack_brief(d,"33","기존 인사이트 카드를 세그먼트 A/B 실험 근거로 전환",
    "33일차의 기존 카드는 삭제하지 않고 AS-IS 관측 근거로 보존했습니다. TO-BE에서는 각 카드가 어떤 세그먼트·개선안·KPI를 설계하는 근거인지 연결했습니다.",[
        "카드 A는 가격·옵션 수 완화, C는 즉시/최종 회복 KPI, F는 동일 조건 반복 방지의 근거입니다.",
        "카드 G는 성공 경로 QA에만 사용하고 이탈률로 해석하지 않습니다.",
        "카드 H의 1위 클릭 29/149(19.5%)는 검색 결과 품질과 순위 보정 지표로 유지합니다.",
        "TO-BE 카드에는 실험 전 의도·필터 상태, 제공 시나리오, SG1~SG4 결과, 95% 불확실성 구간을 함께 보여줍니다."
    ],[
        "각 카드에 대조군·실험군과 핵심 세그먼트 필터를 추가합니다.",
        "소표본 카드는 단정하지 않고 n값과 95% 구간을 같이 표시합니다.",
        "32일차 H1~H6와 카드별 판정 KPI가 일치하는지 교차 검수합니다."
    ]); d.save(OUT33)

# Day 34
d=prepare(SRC34,OUT34,"34일차")
common_baseline(d); dual_segments(d); filter_rules(d)
heading(d,"TO-BE 개선안·우선순위",2)
add_table(d,["순위","개선안","실험 전 대상","성과 판정"],[
    ("1","가격 범위 확대","price 설정 0건 사용자","SG3↑, SG4↓, 상세진입↑"),
    ("2","요구 옵션 수 줄이기","option_count>0인 0건 사용자","SG3↑, SG4↓"),
    ("3","인접지역 확대","가격·옵션 수 제한이 없거나 지역 유연 신호","SG3↑, 지역 변경 회복률"),
    ("4","동일 조건 반복 방지","조건 고수 신호","RT1↓, SG4↓"),
    ("5","복합 제약 단계 완화","가격+옵션 수 복합 설정","즉시·최종 회복률"),
    ("6","연관 검색어·자동완성","검색어 수정 신호","수정률·회복·상세진입"),
],[850,2800,3350,2360],8.15,center=(0,))
heading(d,"A/B 적용 규칙",2)
bullet(d,"전체 10,000명은 대조군 5,000명·실험군 5,000명으로 나누고 실험 전 세그먼트 비중을 균형화한다.")
bullet(d,"대조군은 기존 0건 화면, 실험군은 필터 상태·추정 의도에 맞는 제안을 받는다.")
bullet(d,"효과는 보수·기준·낙관 시나리오와 95% 불확실성 구간으로 표시한다.")
bullet(d,"실험군에서 SG3 증가·SG4 감소를 1차, 상세진입·예약완료를 후속 성과로 판정한다.")
booking_rules(d)
slack_brief(d,"34","세그먼트 맞춤형 0건 회복 개선안과 A/B 실행안 확정",
    "34일차 TO-BE는 0건 화면의 모든 제안을 한번에 비교하는 단순 실험이 아닙니다. 사용자의 필터 상태와 추정 의도에 맞는 제안을 제공한 뒤, 결과 회복과 전환 성과를 SG1~SG4로 비교하는 세그먼트 기반 A/B 시뮬레이션입니다.",[
        "10,000명을 대조군 5,000명과 실험군 5,000명으로 나누고, 각 집단 안의 세그먼트 비중을 맞춥니다.",
        "1차 우선순위는 가격 범위 확대, 2차는 요구 옵션 수 감소, 3차는 인접지역 확대입니다.",
        "초기 2주는 선택지별 노출·선택·조건 변경·결과 회복을 관찰하고, 차이가 확인되면 후속 A/B로 노출 순서와 표현을 검증합니다.",
        "1차 성공은 SG3 증가·SG4 감소, 2차 성과는 호텔 상세진입과 예약완료입니다."
    ],[
        "원본 v03 비중을 유지하며 5,000/5,000 배정과 세그먼트 균형을 검증합니다.",
        "Power BI에 전체 오버뷰, 세그먼트 A/B, 0건 회복 퍼널, 제안별 성과, 불확실성 페이지를 구성합니다.",
        "BOOKING KPI는 데이터 출처와 연결 무결성을 확정한 뒤 최종 성과에 포함합니다."
    ]); d.save(OUT34)

for p in [OUT32,OUT33,OUT34]: print(p)
save_slack_briefs()
print(BRIEF_OUT)
