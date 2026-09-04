from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "04_분석설계" / "팀프로젝트" / "2026" / "09" / "호텔검색_1000명_10000명_세그먼트AB시뮬레이션_증강계획서_20260904_v01_작업본.docx"

BLUE = "2E74B5"; DARK = "1F4D78"; NAVY = "0B2545"; PALE = "E8EEF5"
LIGHT = "F2F4F7"; CALLOUT = "F4F6F9"; GOLD = "7A5A00"; RED = "9B1C1C"
GREEN = "1F5F4A"; GRAY = "666666"; WHITE = "FFFFFF"


def set_font(run, size=11, bold=False, color="000000", name="Calibri", italic=False):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name); fonts.set(qn("w:hAnsi"), name); fonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = tc_pr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=90, bottom=90, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr(); tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None: tc_mar = OxmlElement("w:tcMar"); tc_pr.append(tc_mar)
    for tag, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn("w:" + tag))
        if node is None: node = OxmlElement("w:" + tag); tc_mar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr(); el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true"); tr_pr.append(el)


def set_table_geometry(table, widths):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag, val in [("tblW", sum(widths)), ("tblInd", 120)]:
        el = tbl_pr.find(qn("w:" + tag))
        if el is None: el = OxmlElement("w:" + tag); tbl_pr.append(el)
        el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col in list(grid): grid.remove(col)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(widths[idx])); tc_w.set(qn("w:type"), "dxa")
            margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths, font=8.7, align_center_cols=(0,)):
    table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"
    for i, header in enumerate(headers):
        shade(table.rows[0].cells[i], PALE)
        p = table.rows[0].cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.05
        set_font(p.add_run(header), font, True, NAVY)
    set_repeat_header(table.rows[0])
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if ridx % 2: shade(cells[i], "FAFBFC")
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in align_center_cols else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.08
            set_font(p.add_run(str(value)), font, False, "222222")
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_para(doc, text, size=10.5, bold=False, color="222222", after=6, italic=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(text), size, bold, color, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.5 + .25 * level)
    p.paragraph_format.first_line_indent = Inches(-.25)
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.167
    set_font(p.add_run(text), 10.3)
    return p


def callout(doc, label, text, fill=CALLOUT, color=NAVY):
    table = doc.add_table(rows=1, cols=1); cell = table.cell(0, 0)
    shade(cell, fill); margins(cell, 130, 130, 170, 170)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(label + "  "), 10.5, True, color)
    set_font(p.add_run(text), 10.5, False, "222222")
    set_table_geometry(table, [9360]); doc.add_paragraph().paragraph_format.space_after = Pt(2)


def page_break(doc):
    # Let Word paginate naturally. Explicit break runs render as a square glyph
    # in macOS Quick Look for this Korean-font document.
    return None


doc = Document(); sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.492)

# Preset: standard_business_brief / header: memo_masthead
styles = doc.styles
normal = styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8), ("Heading 2", 13, BLUE, 12, 6), ("Heading 3", 12, DARK, 8, 4)
]:
    st = styles[name]; st.font.name = "Calibri"; st.font.size = Pt(size); st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)

header = sec.header.paragraphs[0]; header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("2팀 · 일본 호텔 검색 데이터 확장 설계"), 9, False, GRAY)
footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("작업본 · 2026-09-04"), 8.5, False, GRAY)

# Cover / executive brief
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(4)
set_font(p.add_run("데이터 증강 계획서"), 24, True, NAVY)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(18)
set_font(p.add_run("1,000명 → 10,000명 의도 세그먼트 기반 A/B 시뮬레이션 재설계"), 14, True, DARK)

add_table(doc, ["항목", "내용"], [
    ("프로젝트", "일본 여행객 호텔 검색 의도·결과 비교 및 이탈 요인 대시보드"),
    ("기준", "2026-09-03 v03 비식별 원본 + 관측형 합성 1,000명 QA 교정본"),
    ("확장 규모", "10,000명: 대조군 5,000명 + 실험군 5,000명"),
    ("설계 단위", "사용자 의도 세그먼트 → 검색 행동 → 0건 반응 → 회복·상세진입"),
    ("문서 상태", "통합 재설계 작업본 v01 / 실제 생성 전 게이트 승인 필요"),
], [1700, 7660], font=9.2)

callout(doc, "핵심 결정", "이번 10,000명 확장은 단순 복제가 아니라, 같은 의도 세그먼트 구성을 가진 대조군과 실험군을 비교하는 검색 행동 시뮬레이션으로 설계한다.", PALE)
add_para(doc, "※ 본 문서는 2026-09-02 증강계획, 09-03 인사이트 실습, 09-04 2팀 개선안·Power BI 구성안과 교수 피드백을 통합한다.", 9.3, False, GRAY, 0, True)

add_heading(doc, "1. 재설계 목적과 범위", 1)
add_para(doc, "1,000명 파일럿은 원본 세션 분포를 유지하는 관측형 합성 가능성을 검증했다. 10,000명 단계에서는 관측형 데이터만 늘리지 않고, 34일차 개선안이 어떤 의도의 사용자에게 효과적일지를 비교할 수 있는 대조군·실험군 시뮬레이션으로 확장한다.")
add_bullet(doc, "1차 목표: 0건 검색 후 결과가 1건 이상 나오는 결과 회복")
add_bullet(doc, "2차 목표: 회복된 검색에서 호텔 상세 진입")
add_bullet(doc, "분석 목표: 전체 평균이 아니라 의도 세그먼트별 개선 효과와 불확실성 확인")
add_bullet(doc, "비목표: 시뮬레이션을 실제 서비스 A/B 실험 결과로 표현하지 않음")

add_heading(doc, "2. 현재 기준선과 승인 조건", 1)
add_table(doc, ["구분", "현재 기준", "판정"], [
    ("원본 분석", "SEARCH 296건, 0건 147건, 세션 43개, 연결 사용자 41명", "기준선"),
    ("1,000명 파일럿", "1,000세션·6900검색; 핵심 지표 절대차 최대 0.106%p", "QA PASS"),
    ("1,000명 DB 크기", "현재 약 93MB", "미충족"),
    ("선행 경량화", "VACUUM 포함 20MB 이하 목표", "10,000명 전 필수"),
    ("10,000명 DB", "150MB 이하 목표", "생성 후 게이트"),
    ("표본 설계", "control 5,000 + treatment 5,000; sample_set_type 필수", "교수 조건"),
    ("희소 표본", "지속 실패 4.7%, 검색어 수정 상세진입 n=10", "불확실성 관리"),
], [1550, 5520, 2290], font=8.8, align_center_cols=(0,2))
callout(doc, "중단 게이트 G0", "1,000명 DB를 20MB 이하로 경량화하고 재생성·재QA하기 전에는 10,000명 생성을 시작하지 않는다.", "FFF3CD", RED)

add_heading(doc, "2.1 원본 v03 DB 기준", 2)
add_para(doc, "증강 확률의 출발점은 합성 1,000명이 아니라 2026-09-03 v03 원본 DB다. USER 89명 중 검색에 연결된 활성 사용자는 41명, 검색 세션은 43개, 검색은 296건이다. 0건을 한 번 이상 경험한 사용자는 27명, 세션은 28개다.")
add_table(doc, ["기준", "분자/분모", "비율", "10,000명 설계 용도"], [
    ("0건 검색", "147/296 검색", "49.7%", "검색 단위 기준선"),
    ("0건 경험 사용자", "27/41명", "65.9%", "사용자 단위 실험 진입 확률"),
    ("0건 경험 세션", "28/43세션", "65.1%", "세션 단위 QA"),
], [2200, 2100, 1300, 3760], font=8.9, align_center_cols=(0,1,2))
add_para(doc, "합성 1,000명은 원본 구조를 재현했는지 확인하는 QA 표본으로만 사용한다. 10,000명의 세그먼트 비중·행동 확률은 원본 41명·43세션·296검색을 기준으로 추정하고 소표본 불확실성을 함께 반영한다.", 9.25, False, RED)

add_heading(doc, "2.2 원본 0건 사용자의 필터 상태", 2)
add_table(doc, ["첫 0건 필터 상태", "사용자", "27명 내 비중", "설계 해석"], [
    ("가격+옵션 수 복합 설정", "17명", "63.0%", "두 완화안의 순서·조합 검증"),
    ("가격·옵션 수 제한 없음", "7명", "25.9%", "지역 확대·검색어 수정 후보"),
    ("옵션 수만 설정", "2명", "7.4%", "요구 옵션 수 줄이기"),
    ("가격만 설정", "1명", "3.7%", "가격 범위 확대"),
], [3400, 1200, 1600, 3160], font=8.7, align_center_cols=(0,1,2))
callout(doc, "구분 필수", "필터 상태와 추정 의도는 실험 전 분류이고, SG1~SG4는 실험 후 판정하는 검색·전환 성과 결과군이다.", PALE)

page_break(doc)
add_heading(doc, "3. 전체 실험 구조", 1)
add_para(doc, "10,000명은 세그먼트마다 5,000명을 만드는 구조가 아니다. 전체를 5,000명씩 나누고, 두 집단 안에 같은 세그먼트 비율을 배치한다.")
add_table(doc, ["단계", "대조군 5,000명", "실험군 5,000명"], [
    ("기본 프로필", "동일한 의도·지역·검색 여건 분포", "동일한 의도·지역·검색 여건 분포"),
    ("0건 화면", "기존 관측형 반응", "의도에 맞는 개선안 노출"),
    ("사용자 행동", "자연 재검색·반복·이탈", "제안 선택·조건 변경·재검색"),
    ("결과", "회복·상세진입", "회복·상세진입"),
    ("식별 컬럼", "sample_set_type=control", "sample_set_type=treatment"),
], [1600, 3880, 3880], font=9)
callout(doc, "공정한 비교", "의도 세그먼트·지역·초기 필터 조건은 두 군에서 같게 맞추고, 실험군에서만 개선안 노출과 선택 확률이 달라지게 한다.")

add_heading(doc, "4. 의도 세그먼트 정의", 1)
add_para(doc, "세그먼트는 실험 후 발생한 행동으로 정의하지 않는다. 생성 시점에 사용자의 탐색 성향과 우선순위를 잠재 의도로 부여하고, 조건 변경은 그 뒤에 나오는 관측 행동으로 생성한다.")
add_table(doc, ["intent_segment", "의도 정의", "생성할 행동 경향", "실험군 주 제안"], [
    ("condition_keeper", "원하는 조건을 쉽게 포기하지 않음", "동일 조건 반복·조건 강화 가능성 상승", "반복 경고 + 대안 비교"),
    ("location_flexible", "위치를 넓혀서라도 숙소를 찾음", "인접지역·대체지역 선택", "지역 확대"),
    ("budget_flexible", "예산을 조정해 결과를 찾음", "최대가격 상향·가격 필터 해제", "가격 범위 확대"),
    ("option_count_flexible", "요구 옵션 수를 줄일 수 있음", "옵션 수 감소", "요구 옵션 수 줄이기"),
    ("query_reframer", "검색 표현이 문제라고 판단함", "검색어 수정·연관어 선택", "연관 검색어·자동완성"),
    ("rapid_resolver", "여러 조건을 빠르게 바꿔 결과를 찾음", "지역·가격·필터 복합 변경", "회복 가능성 상위 제안"),
], [1700, 2480, 2700, 2480], font=8.15)
add_para(doc, "중요: 의도는 행동을 강제하는 답이 아니라 행동 확률을 바꾸는 잠재 요인이다. 현재 DB에는 구체적 편의시설 종류가 없고 옵션 수만 있으므로, '편의시설 해제'가 아니라 '요구 옵션 수 감소'로 생성한다.", 9.5, False, RED)

add_heading(doc, "4.1 실험 후 검색·전환 성과 결과군", 2)
add_table(doc, ["코드", "결과군", "판정 기준", "실험 해석"], [
    ("SG1", "직접 성공형", "첫 검색 결과≥1, 세션 내 상세진입", "0건 처치 전 기준선"),
    ("SG2", "결과 노출·미선택형", "첫 검색 결과≥1, 상세진입 없음", "정렬·추천 개선 후보"),
    ("SG3", "재검색 회복형", "첫 검색 0건 후 후속 검색에서 결과 발생", "실험군에서 증가 목표"),
    ("SG4", "지속 실패형", "첫 검색 0건, 세션 종료까지 미회복", "실험군에서 감소 목표"),
], [900,2200,3880,2380], font=8.55, align_center_cols=(0,1))

add_heading(doc, "5. 세그먼트 비율 산정 방법", 1)
add_para(doc, "기존 DB에는 추정 의도가 직접 기록되지 않았다. 원본 검색 사용자 41명의 초기 필터 상태·이전 행동만을 신호로 사용해 추정 의도를 확률적으로 부여한다. SG1~SG4는 의도 부여에 사용하지 않고 처치 후 결과로만 판정한다.")
add_table(doc, ["입력 신호", "활용 방법", "제약"], [
    ("초기 검색 의도", "LOCATION_ONLY 28.4%, PRICE 3.0%, QUALITY_FILTER 1.0%, AMENITY 10.5%, MIXED 57.1% 분포를 교정 근거로 사용", "이 코드와 사용자 의도는 동일하지 않음"),
    ("0건 후 첫 행동", "동일 조건, 지역 변경, 검색어 변경, 조건 완화, 강화, 복합 신호", "사후 행동을 그대로 세그먼트명으로 사용 금지"),
    ("세션 탐색성", "검색 횟수, 변경 다양성, 최종 회복 여부", "0건 발생 후 검색 횟수를 인과로 강제하지 않음"),
    ("지역·조건", "지역, 가격, 옵션 수 강도를 함께 보정", "희소 셀은 인접 범주 병합 또는 불확실성 확대"),
], [1750, 4800, 2810], font=8.4)
add_bullet(doc, "1차: 행동 신호 기반 확률 점수를 계산한다.")
add_bullet(doc, "2차: 사용자당 intent_segment 1개를 seed 기반으로 추출한다.")
add_bullet(doc, "3차: 산출된 비율을 control·treatment에 동일하게 층화 배정한다.")
add_bullet(doc, "4차: 세그먼트 비율을 ±5%p 변경해 전체 결론의 민감도를 확인한다.")

page_break(doc)
add_heading(doc, "6. 세그먼트별 A/B 시나리오", 1)
add_table(doc, ["세그먼트", "대조군", "실험군 처치", "주 KPI", "34일차 근거"], [
    ("조건 고수형", "기존 재검색", "동일 조건 반복 감지 + 완화 선택지", "반복률↓, 회복률↑", "동일 조건 53건 회복 0%"),
    ("위치 유연형", "자연 지역 변경", "인접지역 선택지 우선 노출", "지역 선택률, 회복률", "지역 변경 회복 10/24=41.7%"),
    ("예산 유연형", "자연 가격 변경", "가격 범위 확대 선택지", "선택률, 조건변경률, 회복률", "가격 설정 검색 0건률 72.6%"),
    ("옵션 수 유연형", "자연 옵션 수 감소", "요구 옵션 수 줄이기", "선택률, 회복률", "옵션 수 조건별 재계산"),
    ("표현 수정형", "자연 검색어 수정", "연관 검색어 10개·자동완성", "수정률, 회복률, 상세진입률", "상세진입 3/10=30%; 고불확실"),
    ("빠른 해결형", "자연 복합 변경", "예상 회복 순위 1개씩 단계 제안", "첫 회복, 최종 회복", "즉시 17.1%, 세션 최종 75.0%"),
], [1450, 1550, 2500, 1800, 2060], font=7.85)
add_para(doc, "※ 17.1%는 0건→다음 검색 전이 140건, 75.0%는 0건 경험 세션 28개가 분모다. 두 비율을 17.1%→75.0% 상승으로 표현하지 않는다.", 9.3, False, RED)

add_heading(doc, "7. 효과 가정과 불확실성", 1)
add_para(doc, "실험군의 성과를 하나의 고정값으로 만들지 않는다. 각 세그먼트에서 기준 회복확률을 추정하고, 개선안의 상승폭을 보수·기준·낙관 세 시나리오로 나눈다.")
add_table(doc, ["구분", "효과 가정", "해석"], [
    ("보수", "세그먼트별 회복률 +2~3%p", "제안을 봐도 반응이 작은 경우"),
    ("기준", "세그먼트별 회복률 +5~8%p", "합리적인 기대 효과"),
    ("낙관", "세그먼트별 회복률 +10~15%p", "맞춤 제안 반응이 큰 경우"),
], [1500, 3200, 4660], font=9)
add_para(doc, "위 범위는 확정 성과가 아니라 시뮬레이션 시작값이다. 세그먼트별 표본수와 기준율을 보고 파라미터 표에서 별도 승인한다.", 9.4, False, RED)
add_table(doc, ["불확실성 원인", "처리", "대시보드 표시"], [
    ("기준 표본이 작음", "비율에 Beta(1,1) 사전분포를 적용해 후험분포 추정", "평균 + 95% 불확실성 구간"),
    ("검색어 수정 n=10", "효과 분포를 넓게 설정; 우선순위 확정 금지", "고불확실 배지"),
    ("지속 실패 4.7%", "복제로 성공 케이스를 인위 생성하지 않음", "표본 n·신뢰구간 병기"),
    ("의도 비율 미지", "비율 ±5%p 민감도 분석", "세그먼트 구성 슬라이서"),
    ("난수 편동", "시나리오별 최소 1,000회 Monte Carlo 반복", "승산확률·효과 범위"),
], [2200, 4300, 2860], font=8.5)

add_heading(doc, "8. 시뮬레이션 판정 지표", 1)
add_table(doc, ["KPI", "계산식", "grain", "판정"], [
    ("0건 후 결과 회복률", "결과≥1 후속 검색 / 0건 후 후속 검색", "전이", "1차 성공"),
    ("제안 선택률", "제안 선택 / 제안 노출", "노출", "중간 반응"),
    ("조건 변경률", "조건을 넓힌 후속 검색 / 제안 노출 0건 검색", "노출/검색", "행동 변화"),
    ("동일 조건 반복률", "동일 조건 후속 검색 / 0건 후 후속 검색", "전이", "감소 목표"),
    ("호텔 상세진입률", "회복 후 detail 존재 전이 / 회복 전이", "전이", "2차 성공"),
    ("예약 완료율", "booking_complete 검색 / 상세진입 검색", "검색/예약", "최종 성과·출처 확정 필요"),
    ("세션 최종 회복률", "마지막 검색 결과≥1 세션 / 0건 경험 세션", "세션", "최종 결과"),
], [1850, 4480, 1050, 1980], font=8.35, align_center_cols=(0,2,3))
add_bullet(doc, "전체 A/B 효과와 세그먼트별 A/B 효과를 모두 계산한다.")
add_bullet(doc, "차이(%p), 상대증감(%), 95% 구간, treatment 우세 확률을 함께 표시한다.")
add_bullet(doc, "세그먼트별 여러 비교는 탐색적 결과로 표시하고, 여러 비교 보정 여부를 주석으로 남긴다.")

page_break(doc)
add_heading(doc, "9. 데이터 모델 확장", 1)
add_table(doc, ["컬럼", "값 예시", "역할"], [
    ("sample_set_type", "control / treatment", "전체 A/B 구분"),
    ("intent_segment", "condition_keeper / location_flexible / …", "사전 의도 세그먼트"),
    ("intent_assignment_prob", "0~1", "의도 부여 확률 추적"),
    ("observed_behavior", "same / region / price / rating / amenity / query / mixed", "실제 재검색 행동"),
    ("treatment_policy", "none / personalized_zero_result", "적용 정책"),
    ("suggestion_type", "region_expand / price_relax / rating_relax / …", "노출 선택지"),
    ("suggestion_exposed", "0 / 1", "노출 여부"),
    ("suggestion_selected", "0 / 1", "선택 여부"),
    ("result_recovered", "0 / 1", "후속 검색 결과≥1"),
    ("hotel_detail_entered", "0 / 1", "회복 후 상세진입"),
    ("scenario_type", "conservative / base / optimistic", "효과 가정"),
    ("simulation_run_id", "SIM_BASE_0001", "반복 실행 구분"),
    ("random_seed", "20260904", "재현성"),
], [2200, 3400, 3760], font=8.35)
add_para(doc, "권장 구조: USER에 intent_segment, 실험 배정 테이블에 sample_set_type, ActionEvent에 노출·선택, SearchTransition에 observed_behavior·result_recovered, SessionSummary에 outcome_segment(SG1~SG4)를 저장한다. 예약은 booking_complete 이벤트의 search_id로 연결한 브리지를 별도로 둔다.", 9.5)

add_heading(doc, "10. 생성 파이프라인", 1)
add_table(doc, ["순서", "작업", "완료 기준"], [
    ("G0", "1,000명 DB 구조 점검 → 인덱스·중복 결과 최소화 → VACUUM", "20MB 이하, QA 재PASS"),
    ("G1", "1,000명 사용자에 의도 확률 점수·세그먼트 부여", "전수 1개 의도, 신호·seed 기록"),
    ("G2", "10,000명 기본 프로필 생성", "단순 ID 복제 없음, 분포 QA"),
    ("G3", "세그먼트 층화 후 5,000/5,000 배정", "군 간 세그먼트·지역·조건 균형"),
    ("G4", "대조군 관측형, 실험군 맞춤 제안 행동 생성", "처치 외 파라미터 동일"),
    ("G5", "3개 시나리오·각 1,000회 반복 시뮬레이션", "run_id·seed·가정 추적 가능"),
    ("G6", "PK/FK·시간순서·플래그·지표·용량 QA", "DB 150MB 이하, 오류 0건"),
    ("G7", "Power BI 분석 마트·매니페스트 생성", "분모·grain·가정 주석 완료"),
], [850, 5700, 2810], font=8.45, align_center_cols=(0,))

add_heading(doc, "11. 단순 복제 방지 규칙", 1)
add_bullet(doc, "원본 세션 하나를 그대로 반복하지 않고, 사용자 속성·의도·지역·필터·행동을 조건부 확률로 새로 추출한다.")
add_bullet(doc, "원본의 희소 결과를 단순히 10배 복사해 신뢰도가 높아졌다고 표현하지 않는다.")
add_bullet(doc, "실험군의 성공을 임의로 확정하지 않고, 노출→선택→변경→회복→상세진입 각 단계를 확률적으로 생성한다.")
add_bullet(doc, "모든 시나리오는 설정 파일, 버전, seed, 소스 해시, 실행 시간을 매니페스트에 남긴다.")

add_heading(doc, "12. 데이터 품질·DB 게이트", 1)
add_table(doc, ["게이트", "검증 항목", "통과 기준"], [
    ("용량", "1,000명 VACUUM 후 / 10,000명 생성 후", "≤20MB / ≤150MB"),
    ("무결성", "USER→SEARCH→FILTER→RESULT→EVENT PK/FK", "고아키·중복키 0건"),
    ("균형", "control/treatment 세그먼트·지역·초기조건", "사전 정의 허용차 이내"),
    ("품질 플래그", "invalid_stay_date_flag, click_in_result_flag", "도메인·필터 규칙 준수"),
    ("행동 논리", "노출 전 선택, 회복 전 상세진입, 시간 역전", "논리 위반 0건"),
    ("재현성", "동일 config·seed 재실행", "핵심 지표·해시 일치"),
], [1300, 4700, 3360], font=8.7, align_center_cols=(0,))
add_para(doc, "v03 품질 규칙: 숙박 날짜 지표는 invalid_stay_date_flag=0만 사용하고, 노출 기반 클릭률은 hotel_click 중 click_in_result_flag=1만 사용한다. 행동 존재 분석에서는 예외 행을 삭제하지 않고 한계를 표시한다.", 9.3, False, RED)

page_break(doc)
add_heading(doc, "13. Power BI 대시보드 연결", 1)
add_table(doc, ["페이지", "핵심 질문", "주요 시각화"], [
    ("01 Overview", "전체 A/B 결과와 사업적 판단은?", "회복률·상세진입률 차이, 95% 구간, 승산확률"),
    ("02 세그먼트 현황", "10,000명은 어떤 의도로 구성됐나?", "세그먼트 비율, 지역·조건 구성, A/B 균형"),
    ("03 세그먼트별 A/B", "어떤 의도군에서 개선 가능성이 높은가?", "효과 크기 포레스트 플롯, 분자·분모 툴팁"),
    ("04 행동 퍼널", "노출부터 상세진입까지 어디서 줄었나?", "노출→선택→변경→회복→상세짅입"),
    ("05 시나리오·불확실성", "가정이 바뀌어도 결론이 유지되나?", "보수·기준·낙관 선택, 분포·구간, 민감도"),
    ("06 QA·한계", "결과를 믿을 수 있는가?", "DB 용량, 균형, 무결성, 희소표본, 가정 배지"),
], [1800, 3440, 4120], font=8.55)
add_bullet(doc, "공통 슬라이서: sample_set_type, intent_segment, scenario_type, 지역, 기간, simulation_run_id")
add_bullet(doc, "모든 비율 툴팁: 분자, 분모, grain, 세그먼트 n, 95% 구간, 시뮬레이션 가정")
add_bullet(doc, "실제 관측값과 시뮬레이션 결과를 색·배지로 구분")

add_heading(doc, "14. 실행 일정과 산출물", 1)
add_table(doc, ["단계", "예상 기간", "주요 산출물", "승인 포인트"], [
    ("1. 경량화", "0.5~1일", "VACUUM 포함 생성기, 20MB 이하 DB", "G0"),
    ("2. 의도 설계", "0.5~1일", "의도 사전, 행동 확률표, 비율 민감도", "G1"),
    ("3. 10,000명 생성", "1~2일", "control/treatment DB, config, manifest", "G2~G4"),
    ("4. 반복 시뮬레이션", "1일", "시나리오별 run 결과, 신뢰구간", "G5"),
    ("5. QA·Power BI", "1~2일", "QA 보고서, 분석 마트, 대시보드", "G6~G7"),
], [1700, 1400, 4000, 2260], font=8.7, align_center_cols=(0,1,3))

add_heading(doc, "15. 최종 의사결정 규칙", 1)
add_table(doc, ["결과", "판정", "후속 행동"], [
    ("전체·주요 세그먼트에서 회복 가능성 높음", "후속 실제 A/B 후보", "로깅 보강 후 소규모 라이브 실험"),
    ("일부 세그먼트만 개선", "선택적 적용", "해당 의도군 대상 개인화 검토"),
    ("평균은 개선, 세그먼트별 편차 큰 경우", "전체 적용 보류", "부정 효과 세그먼트 제외 규칙 설계"),
    ("95% 구간이 0을 넓게 포함", "결론 보류", "표본·로그 확대 후 재검증"),
    ("가정 변경에 따라 결론 역전", "민감한 시뮬레이션", "확정적 추천 금지, 가정 보완"),
], [3000, 1900, 4460], font=8.65)
callout(doc, "발표용 한 문장", "1,000명 관측형 파일럿을 통과한 후, 동일한 의도 세그먼트 구성의 대조군 5,000명과 실험군 5,000명을 생성해 34일차 개선안의 세그먼트별 효과와 불확실성을 비교한다.", PALE)

add_heading(doc, "부록 A. 반드시 남길 실행 파라미터", 1)
add_table(doc, ["분류", "필수 항목"], [
    ("출처", "source_db_sha256, parent_manifest_sha256, dataset_version"),
    ("생성", "generation_version, config_version, random_seed, generated_at"),
    ("표본", "n_users, n_sessions, sample_set_type, intent_segment_share"),
    ("실험", "treatment_policy, scenario_type, uplift_assumption, simulation_runs"),
    ("QA", "DB size, row counts, key violations, balance checks, metric checks"),
], [1800, 7560], font=8.8)

add_heading(doc, "부록 B. 참고 문서", 1)
for source in [
    "호텔검색_데이터증강계획서_세그먼트가설중심_20260902_v06_현행본_296건기준.docx",
    "BI시각화_33일차_2팀프로젝트실습_해답_20260903_v03_제출본.docx",
    "BI시각화_34일차_2팀_기초실습_문제-2_20260904_v01_제출본.docx",
    "BI시각화_34일차_2팀프로젝트실습_해답_20260904_v01_제출본.docx",
    "일본호텔검색_PowerBI대시보드_구성안_20260904_v01_작업본.docx",
    "호텔검색_관측형합성1000명_QA교정결과_260903_1745_02.xlsx 및 교정 매니페스트",
    "호텔검색_데이터품질갱신안내_20260903_v01_현행본.md / data_quality_manifest.json",
]: add_bullet(doc, source)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
